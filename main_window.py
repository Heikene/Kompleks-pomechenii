from __future__ import annotations

import sys
import os
import re
import tempfile
import uuid
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime, date, timedelta
from decimal import Decimal, InvalidOperation

from PySide6.QtCore import Qt, QThread, Signal
from PySide6.QtGui import QDragEnterEvent, QDropEvent, QGuiApplication, QKeySequence, QShortcut
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QListWidget,
    QListWidgetItem, QPushButton, QMessageBox, QFileDialog,
    QGridLayout, QDateEdit, QDialog, QVBoxLayout, QHBoxLayout,
    QSpinBox, QTableWidget, QTableWidgetItem, QProgressBar, QComboBox,
    QTabWidget, QAbstractItemView
)

from docx import Document
from docx.shared import Mm
from docx.image.exceptions import UnrecognizedImageError
from docxtpl import DocxTemplate, RichText, InlineImage
from pymorphy3 import MorphAnalyzer

from file_utils import temp_docx
import io_manager
import template_renderer
import table_processor
from logger import logger
import word_table5_splitter


# ===================== Spreadsheet-like table =====================
class SpreadsheetTable(QTableWidget):
    """QTableWidget с excel-копипастой."""
    def keyPressEvent(self, e):
        ctrl = bool(e.modifiers() & Qt.ControlModifier)
        if ctrl and e.key() == Qt.Key_C:
            self._copy_selection_to_clipboard()
            return
        if ctrl and e.key() == Qt.Key_V:
            self._paste_from_clipboard()
            return
        super().keyPressEvent(e)

    def _copy_selection_to_clipboard(self):
        rngs = self.selectedRanges()
        if rngs:
            r = rngs[0]
            rows = range(r.topRow(), r.bottomRow() + 1)
            cols = range(r.leftColumn(), r.rightColumn() + 1)
        else:
            rows = range(0, self.rowCount())
            cols = range(0, self.columnCount())

        lines = []
        for i in rows:
            vals = []
            for j in cols:
                it = self.item(i, j)
                vals.append("" if it is None else it.text())
            lines.append("\t".join(vals))
        QApplication.clipboard().setText("\n".join(lines))

    def _paste_from_clipboard(self):
        text = QApplication.clipboard().text()
        if not text:
            return
        rows_data = [
            [c.strip() for c in row.replace(";", "\t").split("\t")]
            for row in text.splitlines()
            if row.strip() != ""
        ]
        if not rows_data:
            return
        rngs = self.selectedRanges()
        start_row = rngs[0].topRow() if rngs else 0
        start_col = rngs[0].leftColumn() if rngs else 0
        need_r = start_row + len(rows_data)
        need_c = start_col + max(len(r) for r in rows_data)
        if self.rowCount() < need_r:
            self.setRowCount(need_r)
        if self.columnCount() < need_c:
            self.setColumnCount(need_c)
        for i, row_vals in enumerate(rows_data):
            for j, val in enumerate(row_vals):
                r = start_row + i
                c = start_col + j
                it = self.item(r, c)
                if it is None:
                    it = QTableWidgetItem()
                    self.setItem(r, c, it)
                it.setText(val)


# ===================== Rooms dialog =====================
class RoomsDialog(QDialog):
    def __init__(self, parent=None, initial_rooms: List[Dict[str, str]] | None = None):
        super().__init__(parent)
        self.setWindowTitle("Настройки помещений")
        self.resize(700, 450)

        self.spin = QSpinBox()
        self.spin.setRange(0, 100)
        self.spin.setValue(len(initial_rooms or []))
        self.spin.valueChanged.connect(self._on_count)

        self.tbl = SpreadsheetTable()
        self.tbl.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.tbl.setSelectionBehavior(QAbstractItemView.SelectItems)
        self.tbl.setEditTriggers(QAbstractItemView.AllEditTriggers)

        headers = [
            "Номер помещения", "Наименование", "Класс чистоты",
            "Площадь, м²", "Объём, м³", "Δдавл., Pa",
            "Расход, м³/ч", "Кратность, не менее", "Темп., °C", "RH, %"
        ]
        self.tbl.setColumnCount(len(headers))
        self.tbl.setHorizontalHeaderLabels(headers)

        self._on_count(self.spin.value())
        if initial_rooms:
            for r, data in enumerate(initial_rooms):
                for c, key in enumerate(
                    ["num", "name", "klass", "area", "volume",
                     "dp", "airflow", "exchange", "temp", "rh"]
                ):
                    self.tbl.setItem(r, c, QTableWidgetItem(data.get(key, "")))

        btn_ok = QPushButton("ОК")
        btn_cancel = QPushButton("Отмена")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        hb = QHBoxLayout()
        hb.addStretch()
        hb.addWidget(btn_ok)
        hb.addWidget(btn_cancel)

        lay = QVBoxLayout(self)
        lay.addWidget(QLabel("Количество помещений:"))
        lay.addWidget(self.spin)
        lay.addWidget(self.tbl)
        lay.addLayout(hb)

    def _on_count(self, n: int) -> None:
        self.tbl.setRowCount(n)
        for r in range(n):
            for c in range(self.tbl.columnCount()):
                if not self.tbl.item(r, c):
                    self.tbl.setItem(r, c, QTableWidgetItem(""))

    def get_rooms(self) -> List[Dict[str, str]]:
        keys = ["num", "name", "klass", "area", "volume",
                "dp", "airflow", "exchange", "temp", "rh"]
        out: List[Dict[str, str]] = []
        for r in range(self.tbl.rowCount()):
            row: Dict[str, str] = {}
            for c, k in enumerate(keys):
                it = self.tbl.item(r, c)
                row[k] = (it.text().strip() if it else "")
            out.append(row)
        return out


# ===================== Equipment date helpers =====================
def _parse_any_date(v):
    """Пытается распарсить дату из datetime/date/строки/Excel-числа."""
    if v is None:
        return None

    if isinstance(v, datetime):
        return v.date()
    if isinstance(v, date):
        return v

    if isinstance(v, (int, float)):
        if v > 20000:
            base = datetime(1899, 12, 30)  # excel epoch
            return (base + timedelta(days=float(v))).date()
        return None

    s = str(v).strip()
    if not s:
        return None

    s2 = s.replace("T", " ").replace("Z", "").strip()

    try:
        dt = datetime.fromisoformat(s2)
        return dt.date()
    except Exception:
        pass

    m = re.match(r"^\s*(\d{4})[-/](\d{1,2})[-/](\d{1,2})", s2)
    if m:
        y, mo, d = map(int, m.groups())
        try:
            return date(y, mo, d)
        except Exception:
            return None

    m = re.match(r"^\s*(\d{1,2})[.\/](\d{1,2})[.\/](\d{4})", s2)
    if m:
        d, mo, y = map(int, m.groups())
        try:
            return date(y, mo, d)
        except Exception:
            return None

    return None


def _fmt_ddmmyyyy(v) -> str:
    d = _parse_any_date(v)
    return d.strftime("%d.%m.%Y") if d else ""


def _extract_two_dates_from_value(v) -> tuple[date | None, date | None]:
    """Если значение содержит 1 или 2 даты — вернём (d1, d2)."""
    if v is None:
        return (None, None)

    if isinstance(v, (datetime, date)):
        d = _parse_any_date(v)
        return (d, None)

    s = str(v).strip().replace("\u00a0", " ")
    if not s:
        return (None, None)

    s = re.sub(r"\s+", " ", s)
    cand = re.findall(
        r"\d{4}[-/]\d{1,2}[-/]\d{1,2}|\d{1,2}[./]\d{1,2}[./]\d{4}",
        s
    )
    dates_found: list[date] = []
    for c in cand:
        d = _parse_any_date(c)
        if d:
            dates_found.append(d)

    if not dates_found:
        d = _parse_any_date(s)
        return (d, None)

    if len(dates_found) == 1:
        return (dates_found[0], None)
    return (dates_found[0], dates_found[1])


# ===================== Equipment dialog =====================
class EquipmentDialog(QDialog):
    def __init__(self, parent=None, items_by_sheet: dict[str, list[dict]] | None = None):
        super().__init__(parent)
        self.setWindowTitle("Выбор оборудования")
        self.resize(900, 520)

        self.tabs = QTabWidget(self)
        self._lists: dict[str, QListWidget] = {}

        KEY_DATE_FROM = (
            "date_check", "check_date", "calibration_date", "date_from",
            "Дата поверки", "дата поверки", "дата_поверки",
        )
        KEY_DATE_TO = (
            "valid_to", "valid_until", "date_to", "expiry", "validity_to",
            "Действительно до", "действительно до",
            "Срок действия поверки", "срок действия поверки", "срок_действия_поверки",
        )
        KEY_DATE_RANGE = (
            "Дата поверки/ Действительно до:",
            "Дата поверки/Действительно до:",
            "Дата поверки / Действительно до:",
            "дата поверки/ действительно до",
            "дата поверки / действительно до",
            "дата поверки/действительно до",
        )

        def norm_key(s: str) -> str:
            return re.sub(r"\s+", " ", str(s or "")).strip().lower()

        def get_any(it: dict, keys: tuple[str, ...]):
            for k in keys:
                if k in it and it.get(k) not in (None, ""):
                    return it.get(k)

            nk = {norm_key(k): k for k in it.keys()}
            for k in keys:
                kk = nk.get(norm_key(k))
                if kk is not None and it.get(kk) not in (None, ""):
                    return it.get(kk)

            return None

        def find_range_value(it: dict):
            v = get_any(it, KEY_DATE_RANGE)
            if v not in (None, ""):
                return v

            for real_k, vv in it.items():
                if vv in (None, ""):
                    continue
                kk = norm_key(real_k)
                if ("поверк" in kk) and (("действ" in kk) or ("срок" in kk) or ("valid" in kk)):
                    return vv
            return None

        for sheet, items in (items_by_sheet or {}).items():
            lw = QListWidget()
            lw.setSelectionMode(QListWidget.MultiSelection)

            for it in items:
                name_sn = (it.get("name_sn") or it.get("name") or "").strip()

                raw_from = get_any(it, KEY_DATE_FROM)
                raw_to = get_any(it, KEY_DATE_TO)

                if (raw_from in (None, "")) and (raw_to in (None, "")):
                    rng = find_range_value(it)
                    d1, d2 = _extract_two_dates_from_value(rng)
                    if d1:
                        raw_from = d1
                    if d2:
                        raw_to = d2

                date_from = _fmt_ddmmyyyy(raw_from)
                date_to = _fmt_ddmmyyyy(raw_to)

                if date_to:
                    text = f"до {date_to} — {name_sn}"
                elif date_from:
                    text = f"поверка {date_from} — {name_sn}"
                else:
                    text = name_sn

                item = QListWidgetItem(text)
                item.setData(Qt.UserRole, it)
                item.setToolTip(
                    f"Поверка: {date_from or '—'}\n"
                    f"Действительно до: {date_to or '—'}"
                )
                lw.addItem(item)

            self.tabs.addTab(lw, f"{sheet} ({len(items)})")
            self._lists[sheet] = lw

        btn_ok = QPushButton("ОК")
        btn_cancel = QPushButton("Отмена")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        hb = QHBoxLayout()
        hb.addStretch()
        hb.addWidget(btn_ok)
        hb.addWidget(btn_cancel)

        lay = QVBoxLayout(self)
        lay.addWidget(self.tabs)
        lay.addLayout(hb)

    def get_selected(self) -> list[dict]:
        res: list[dict] = []
        for lw in self._lists.values():
            res.extend(i.data(Qt.UserRole) for i in lw.selectedItems())
        return res


# ===================== Tests dialog =====================
class TestsDialog(QDialog):
    """Выбор тестов + кнопка калькулятора рядом с тестом 11."""
    def __init__(self, parent=None, items: List[str] | None = None,
                 selected: List[str] | None = None,
                 calc_launcher=None):
        super().__init__(parent)
        self.setWindowTitle("Выберите тесты")
        self.resize(600, 400)

        self.selection_order: List[int] = []
        self.calc_launcher = calc_launcher

        lay = QVBoxLayout(self)
        self.list = QListWidget(self)
        self.list.setSelectionMode(QAbstractItemView.MultiSelection)
        lay.addWidget(self.list)

        self._base_text_by_row: list[str] = []
        self._label_by_row: list[QLabel] = []

        rx_calc = re.compile(r"(?i)тест\s*11.*расход.*приточ", re.IGNORECASE)

        for txt in (items or []):
            item = QListWidgetItem()
            roww = QWidget()
            hb = QHBoxLayout(roww)
            hb.setContentsMargins(6, 2, 6, 2)

            lbl = QLabel(txt)
            hb.addWidget(lbl)
            hb.addStretch()

            if rx_calc.search(txt) and callable(self.calc_launcher):
                btn = QPushButton("Кальк.")
                btn.setToolTip("Открыть калькулятор Тест 11")
                btn.clicked.connect(self.calc_launcher)
                hb.addWidget(btn)

            self.list.addItem(item)
            self.list.setItemWidget(item, roww)
            item.setSizeHint(roww.sizeHint())

            self._base_text_by_row.append(txt)
            self._label_by_row.append(lbl)

        if selected:
            by_text = {t: i for i, t in enumerate(self._base_text_by_row)}
            for txt in selected:
                r = by_text.get(txt)
                if r is not None:
                    self.list.item(r).setSelected(True)
                    self.selection_order.append(r)

        btn_ok = QPushButton("ОК")
        btn_cancel = QPushButton("Отмена")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        hb2 = QHBoxLayout()
        hb2.addStretch()
        hb2.addWidget(btn_ok)
        hb2.addWidget(btn_cancel)
        lay.addLayout(hb2)

        self.list.itemSelectionChanged.connect(self._on_selection_changed)
        self._renumber()

    def _on_selection_changed(self):
        current = {i.row() for i in self.list.selectedIndexes()}
        prev = set(self.selection_order)
        for r in list(self.selection_order):
            if r not in current:
                self.selection_order.remove(r)
        self.selection_order.extend(sorted(current - prev))
        self._renumber()

    def _renumber(self):
        for r, lbl in enumerate(self._label_by_row):
            lbl.setText(self._base_text_by_row[r])
        for n, r in enumerate(self.selection_order, start=1):
            lbl = self._label_by_row[r]
            base = self._base_text_by_row[r]
            lbl.setText(f"{n:02d}. {base}")

    def get_selected(self) -> List[str]:
        return [self._base_text_by_row[r] for r in self.selection_order]


# ===================== Appendix images (drag&drop + Ctrl+V) =====================
_ALLOWED_IMG_EXT = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif", ".pdf"}  # что можно выбрать/перетащить
DOCX_IMG_EXT = {".png", ".jpg", ".jpeg", ".gif", ".bmp"}  # что реально вставляет python-docx/docxtpl


class DropListWidget(QListWidget):
    def __init__(self, parent=None, *, paste_dir: Path | None = None, paste_prefix: str = "paste"):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setSelectionMode(QAbstractItemView.ExtendedSelection)

        base = Path(paste_dir) if paste_dir else Path(tempfile.gettempdir()) / "doc_generator_app" / "pasted_images"
        base.mkdir(parents=True, exist_ok=True)
        self._paste_dir = base
        self._paste_prefix = paste_prefix

    def dragEnterEvent(self, e: QDragEnterEvent):
        if e.mimeData().hasUrls():
            e.acceptProposedAction()
        else:
            super().dragEnterEvent(e)

    def dragMoveEvent(self, e):
        if e.mimeData().hasUrls():
            e.acceptProposedAction()
        else:
            super().dragMoveEvent(e)

    def dropEvent(self, e: QDropEvent):
        if e.mimeData().hasUrls():
            paths: list[str] = []
            for url in e.mimeData().urls():
                p = url.toLocalFile()
                if not p:
                    continue
                ext = Path(p).suffix.lower()
                if ext in _ALLOWED_IMG_EXT and Path(p).exists():
                    paths.append(str(Path(p).resolve()))
            self._add_paths(paths)
            e.acceptProposedAction()
            return
        super().dropEvent(e)

    def keyPressEvent(self, e):
        if e.matches(QKeySequence.Paste):
            if self.paste_from_clipboard():
                return
        super().keyPressEvent(e)

    def paste_from_clipboard(self) -> bool:
        cb = QGuiApplication.clipboard()
        md = cb.mimeData()

        if md.hasImage():
            img = cb.image()
            if img.isNull():
                return False

            fname = f"{self._paste_prefix}_{datetime.now():%Y%m%d_%H%M%S}_{uuid.uuid4().hex[:8]}.png"
            out = (self._paste_dir / fname).resolve()
            ok = img.save(str(out), "PNG")
            if not ok:
                return False

            self._add_paths([str(out)])
            return True

        if md.hasUrls():
            paths: list[str] = []
            for url in md.urls():
                p = url.toLocalFile()
                if not p:
                    continue
                pp = Path(p)
                if pp.exists() and pp.suffix.lower() in _ALLOWED_IMG_EXT:
                    paths.append(str(pp.resolve()))
            if paths:
                self._add_paths(paths)
                return True

        return False

    def _add_paths(self, paths: list[str]):
        existing = set(self.get_paths())
        for p in paths:
            if p in existing:
                continue
            self.addItem(QListWidgetItem(p))
            existing.add(p)

    def get_paths(self) -> list[str]:
        return [self.item(i).text() for i in range(self.count())]


class AppendixImagesDialog(QDialog):
    def __init__(self, parent=None, title: str = "Загрузка изображений/скриншотов",
                 initial_paths: list[str] | None = None,
                 paste_prefix: str = "app"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(760, 480)

        self.list = DropListWidget(self, paste_prefix=paste_prefix)
        if initial_paths:
            self.list._add_paths(initial_paths)

        self._sc_paste = QShortcut(QKeySequence.Paste, self)
        self._sc_paste.activated.connect(self._on_paste)

        hint = QLabel(
            "Перетащите файлы сюда или нажмите «Добавить…».\n"
            "Также можно вставить скриншот: Ctrl+V.\n"
            "Поддержка: PNG/JPG/BMP/TIFF/GIF/PDF.\n"
            "Важно: PDF/TIFF docxtpl вставлять не умеет — будут пропущены."
        )
        hint.setWordWrap(True)

        btn_add = QPushButton("Добавить…")
        btn_rm = QPushButton("Удалить выбранные")
        btn_clear = QPushButton("Очистить")

        btn_add.clicked.connect(self._add_files)
        btn_rm.clicked.connect(self._remove_selected)
        btn_clear.clicked.connect(self.list.clear)

        hb = QHBoxLayout()
        hb.addWidget(btn_add)
        hb.addWidget(btn_rm)
        hb.addWidget(btn_clear)
        hb.addStretch()

        btn_ok = QPushButton("ОК")
        btn_cancel = QPushButton("Отмена")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        hb2 = QHBoxLayout()
        hb2.addStretch()
        hb2.addWidget(btn_ok)
        hb2.addWidget(btn_cancel)

        lay = QVBoxLayout(self)
        lay.addWidget(hint)
        lay.addWidget(self.list)
        lay.addLayout(hb)
        lay.addLayout(hb2)

    def _on_paste(self):
        if not self.list.paste_from_clipboard():
            QApplication.beep()

    def _add_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Выберите файлы",
            os.getcwd(),
            "Images/PDF (*.png *.jpg *.jpeg *.bmp *.tif *.tiff *.gif *.pdf)"
        )
        if files:
            self.list._add_paths([str(Path(f).resolve()) for f in files])

    def _remove_selected(self):
        for it in list(self.list.selectedItems()):
            row = self.list.row(it)
            self.list.takeItem(row)

    def get_paths(self) -> list[str]:
        return self.list.get_paths()

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _cell_center_no_indent(cell):
    """
    Делает:
      - текст без табов/неразрывных пробелов
      - выравнивание по центру
      - убирает абзацные отступы
      - убирает внутренние поля ячейки (tcMar), которые обычно и дают "сдвиг"
      - фиксирует jc=center на уровне XML (чтобы стиль не перебил)
    """
    import re

    txt = (cell.text or "")
    txt = txt.replace("\t", "").replace("\u00A0", " ")
    txt = re.sub(r"\s+\n", "\n", txt)
    txt = re.sub(r"\n\s+", "\n", txt)
    txt = txt.strip()

    # ВАЖНО: text пересоздаёт абзацы -> форматируем уже после
    cell.text = txt

    # абзацы: центр + без отступов/таба
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pf = p.paragraph_format
        pf.left_indent = 0
        pf.right_indent = 0
        pf.first_line_indent = 0
        pf.space_before = 0
        pf.space_after = 0

        pPr = p._p.get_or_add_pPr()

        # убрать tabs
        tabs = pPr.find(qn("w:tabs"))
        if tabs is not None:
            pPr.remove(tabs)

        # убрать ind
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            pPr.remove(ind)

        # жёстко: center в XML
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "center")

    # убрать внутренние поля ячейки (это то, что чаще всего визуально мешает "центру")
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    def _set_mar(tag: str):
        el = tcMar.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            tcMar.append(el)
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")

    _set_mar("left")
    _set_mar("right")

    # опционально: вертикальный центр
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), "center")

# === Автозаполнение "Тест 11.2. Проверка кратности воздухообмена в ЧП" ===
def _fill_test_112(doc, rooms):
    import re

    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip().lower()

    # 1) находим таблицу теста 11.2
    t112 = None
    for t in doc.tables:
        if not t.rows:
            continue
        head = " ".join(c.text for c in t.rows[0].cells)
        h = norm(head)
        if ("тест" in h and ("11.2" in h or "11,2" in h)) and ("кратност" in h or "чп" in h):
            t112 = t
            break
    if t112 is None:
        return

    # 2) находим строку, где начинаются названия колонок (после "Результаты испытания")
    hdr_cols_row_idx = None
    for i, row in enumerate(t112.rows):
        line = norm(" ".join(c.text for c in row.cells))
        if "результаты испытани" in line:
            hdr_cols_row_idx = i + 1 if (i + 1) < len(t112.rows) else None
            break

    if hdr_cols_row_idx is None:
        for i, row in enumerate(t112.rows):
            line = norm(" ".join(c.text for c in row.cells))
            if ("номер" in line and "объем" in line) or ("общий расход" in line):
                hdr_cols_row_idx = i
                break

    if hdr_cols_row_idx is None:
        return

    # 3) ВАЖНО: из-за merge/gridSpan "Общий расход" может встречаться несколько раз.
    def find_cols(cells, *needles):
        out = []
        for ci, c in enumerate(cells):
            tt = norm(c.text)
            if all(n in tt for n in needles):
                out.append(ci)
        return out

    def find_col(cells, *needles):
        for ci, c in enumerate(cells):
            tt = norm(c.text)
            if all(n in tt for n in needles):
                return ci
        return None

    cols_row = t112.rows[hdr_cols_row_idx]
    col_totals = find_cols(cols_row.cells, "общий", "расход")     # <-- список
    col_fact = find_col(cols_row.cells, "фактическ")

    if not col_totals and col_fact is None:
        return

    # 4) Удаляем "пустую строку" от docxtpl (строка где был отдельный {% for %})
    def is_spacer(row):
        return all(norm(c.text) in ("", "¤") for c in row.cells)

    # Шапка занимает 2 строки: заголовки + подзаголовки
    data_start = hdr_cols_row_idx + 2

    while data_start < len(t112.rows) and is_spacer(t112.rows[data_start]):
        t112._tbl.remove(t112.rows[data_start]._tr)

    # 5) граница данных - до "КОММЕНТАРИИ"
    data_end = len(t112.rows)
    for i in range(data_start, len(t112.rows)):
        first_cell = norm(t112.rows[i].cells[0].text)
        if first_cell.startswith("комментар"):
            data_end = i
            break
    if data_end <= data_start:
        return

    # 6) если помещений больше, чем строк - дополняем копиями последней строки данных
    need = len(rooms)
    have = data_end - data_start
    if need > have:
        from copy import deepcopy
        sample_tr = deepcopy(t112.rows[data_end - 1]._tr)
        for _ in range(need - have):
            t112._tbl.append(deepcopy(sample_tr))
        data_end = data_start + need

    # 7) заполняем
    for idx, room in enumerate(rooms):
        r = data_start + idx
        if r >= len(t112.rows):
            break

        row = t112.rows[r]
        total_flow = (room.get("total_flow") or "").strip()
        exch_act = (room.get("exchange_actual") or "").strip()

        # общий расход: записываем (если есть) и ВСЕГДА форматируем все найденные колонки
        for ci in col_totals:
            if ci < len(row.cells):
                cell = row.cells[ci]
                if total_flow:
                    cell.text = total_flow
                _cell_center_no_indent(cell)

        # фактическая: заполняем как было
        if col_fact is not None and col_fact < len(row.cells) and exch_act:
            row.cells[col_fact].text = exch_act



from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re





# ===================== DOCX equipment date postprocess =====================
def _fix_weird_ddmmyyyy(s: str) -> str:
    m = re.match(r"^(\d{1,2})\.(\d{2})(\d{4})$", s)
    if m:
        return f"{m.group(1)}.{m.group(2)}.{m.group(3)}"
    return s


def _parse_date_any(s: str):
    s = (s or "").strip().replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).replace("г.", "").replace("г", "").strip(" .")
    if not s:
        return None

    s = s.split()[0]
    s = _fix_weird_ddmmyyyy(s)

    fmts = ["%Y-%m-%d", "%d.%m.%Y", "%d-%m-%Y", "%Y.%m.%d", "%d/%m/%Y", "%Y/%m/%d"]
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt).date()
        except Exception:
            pass

    m = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})$", s)
    if m:
        try:
            return date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except Exception:
            return None

    m = re.match(r"^(\d{1,2})\.(\d{1,2})\.(\d{4})$", s)
    if m:
        try:
            return date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
        except Exception:
            return None

    return None


def _fmt_date(d: date) -> str:
    return d.strftime("%d.%m.%Y")


def _format_date_range_cell(text: str) -> str:
    if not text or not text.strip():
        return text

    raw = re.sub(r"\s*\n\s*", " ", text).strip()

    if "/" in raw:
        parts = [p.strip() for p in raw.split("/") if p.strip()]
        if len(parts) >= 2:
            d1 = _parse_date_any(parts[0])
            d2 = _parse_date_any(parts[1])
            if d1 and d2:
                return f"{_fmt_date(d1)} / {_fmt_date(d2)}"
            return raw

    date_like = re.findall(
        r"\d{4}-\d{1,2}-\d{1,2}|\d{1,2}\.\d{1,2}\.\d{4}|\d{1,2}\.\d{2}\d{4}",
        raw
    )
    if len(date_like) >= 2:
        d1 = _parse_date_any(date_like[0])
        d2 = _parse_date_any(date_like[1])
        if d1 and d2:
            return f"{_fmt_date(d1)} / {_fmt_date(d2)}"

    d = _parse_date_any(raw)
    if d:
        return _fmt_date(d)

    return raw


def postprocess_equipment_dates(doc) -> None:
    """
    Ищет в документе столбец 'Дата поверки/ Действительно до:' и
    форматирует значения в 'ДД.ММ.ГГГГ / ДД.ММ.ГГГГ'.
    """
    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip().lower()

    hdr_needle = "дата поверки/ действительно до"
    for t in doc.tables:
        if not t.rows:
            continue

        hdr_row_idx = None
        for ri in range(min(3, len(t.rows))):
            row_text = " ".join(c.text for c in t.rows[ri].cells)
            if hdr_needle in norm(row_text):
                hdr_row_idx = ri
                break
        if hdr_row_idx is None:
            continue

        col_idx = None
        for ci, c in enumerate(t.rows[hdr_row_idx].cells):
            if hdr_needle in norm(c.text):
                col_idx = ci
                break
        if col_idx is None:
            continue

        for ri in range(hdr_row_idx + 1, len(t.rows)):
            cell = t.rows[ri].cells[col_idx]
            old = cell.text
            new = _format_date_range_cell(old)
            if new != old:
                cell.text = new


def update_fields_only_with_word(docx_path: str) -> None:
    """Обновляет поля Word (NUMPAGES и т.п.) через COM. Требует pywin32 и MS Word."""
    import win32com.client as win32

    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path))
        doc.Repaginate()
        doc.Fields.Update()
        for sec in doc.Sections:
            sec.Headers(1).Range.Fields.Update()
            sec.Footers(1).Range.Fields.Update()
        doc.Save()
        doc.Close()
    finally:
        word.Quit()


# ===================== Render Worker =====================



# ===================== Main Window =====================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Генератор OQ/PQ — PySide6 + DocxTemplate")

        base = Path.cwd()
        self.defaults = {
            "OQ": {
                "tpl": base / "Шаблон OQ.docx",
                "tests": base / "Тесты OQ.docx",
                "xls_tests": base / "tests OQ.xlsx",
                "xls_eq": base / "ПЕРЕЧЕНЬ ПРИБОРОВ OQ.xlsx",
                "scans_dir": base / "Сканы поверок"
            },
            "PQ": {
                "tpl": base / "Шаблон PQ.docx",
                "tests": base / "Тесты PQ.docx",
                "xls_tests": base / "tests PQ.xlsx",
                "xls_eq": base / "ПЕРЕЧЕНЬ ПРИБОРОВ OQ.xlsx",
                "scans_dir": base / "Сканы поверок"
            },
        }

        self.rooms: List[Dict[str, str]] = []
        self.equipment: List[Dict[str, str]] = []
        self.all_tests: List[str] = []
        self.selected_tests: List[str] = []
        self._batch_modes: List[str] = []

        self.app1_images: list[str] = []
        self.app4_images: list[str] = []

        self._build_ui()
        self.on_mode_changed("OQ")

    def _build_ui(self) -> None:
        cw = QWidget()
        g = QGridLayout()
        g.setSpacing(10)
        row = 0

        g.addWidget(QLabel("Режим:"), row, 0)
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["OQ", "PQ", "OQ и PQ"])
        self.mode_combo.currentTextChanged.connect(self.on_mode_changed)
        g.addWidget(self.mode_combo, row, 1, 1, 2)
        row += 1

        g.addWidget(QLabel("Шаблон (.docx):"), row, 0)
        self.tpl_path = QLineEdit()
        self.tpl_path.setReadOnly(True)
        g.addWidget(self.tpl_path, row, 1, 1, 1)
        btn_browse_tpl = QPushButton("…")
        btn_browse_tpl.clicked.connect(lambda: self._select(self.tpl_path, "Word Documents (*.docx)", save=False))
        g.addWidget(btn_browse_tpl, row, 2)
        row += 1

        g.addWidget(QLabel("Документ-тесты (.docx):"), row, 0)
        self.tests_path = QLineEdit()
        g.addWidget(self.tests_path, row, 1, 1, 1)
        btn_browse_tests = QPushButton("…")
        btn_browse_tests.clicked.connect(lambda: self._select(self.tests_path, "Word Documents (*.docx)", save=False))
        g.addWidget(btn_browse_tests, row, 2)
        row += 1

        btn_calc11 = QPushButton("Калькулятор Тест 11 (расход)…")
        btn_calc11.clicked.connect(self.launch_test11_calculator)
        g.addWidget(btn_calc11, row, 0, 1, 3)
        row += 1

        g.addWidget(QLabel("Excel с тестами:"), row, 0)
        self.xls_path = QLineEdit()
        self.xls_path.setReadOnly(True)
        g.addWidget(self.xls_path, row, 1, 1, 1)
        btn_browse_xls = QPushButton("…")
        btn_browse_xls.clicked.connect(lambda: self._select(self.xls_path, "Excel Files (*.xlsx)", save=False))
        g.addWidget(btn_browse_xls, row, 2)
        row += 1

        self.tests_btn = QPushButton("Выбрать тесты…")
        self.tests_btn.clicked.connect(self.choose_tests)
        g.addWidget(self.tests_btn, row, 0, 1, 3)
        row += 1

        g.addWidget(QLabel("Excel-оборудование:"), row, 0)
        self.equipment_xls_path = QLineEdit()
        self.equipment_xls_path.setReadOnly(True)
        g.addWidget(self.equipment_xls_path, row, 1, 1, 1)
        btn_eq = QPushButton("Загрузить оборудование")
        btn_eq.clicked.connect(lambda: self.load_equipment(silent=False))
        g.addWidget(btn_eq, row, 2)
        row += 1

        btn_rooms = QPushButton("Настройки помещений…")
        btn_rooms.clicked.connect(self.edit_rooms)
        g.addWidget(btn_rooms, row, 0, 1, 3)
        row += 1

        g.addWidget(QLabel("Папка со сканами (поверки):"), row, 0)
        self.scans_dir_input = QLineEdit()
        g.addWidget(self.scans_dir_input, row, 1, 1, 1)
        btn_dir = QPushButton("…")
        btn_dir.clicked.connect(lambda: self._select_dir(self.scans_dir_input))
        g.addWidget(btn_dir, row, 2)
        row += 1

        g.addWidget(QLabel("Приложение 1 (картинки/скриншоты):"), row, 0, Qt.AlignTop)
        self.app1_btn = QPushButton("Выбрать / перетащить файлы…")
        self.app1_btn.clicked.connect(self.choose_app1_images)
        g.addWidget(self.app1_btn, row, 1, 1, 2)
        row += 1

        self.app1_info = QLabel("Файлы не выбраны")
        self.app1_info.setWordWrap(True)
        g.addWidget(self.app1_info, row, 1, 1, 2)
        row += 1

        g.addWidget(QLabel("Приложение 4 (картинки/скриншоты):"), row, 0, Qt.AlignTop)
        self.app4_btn = QPushButton("Выбрать / перетащить файлы…")
        self.app4_btn.clicked.connect(self.choose_app4_images)
        g.addWidget(self.app4_btn, row, 1, 1, 2)
        row += 1

        self.app4_info = QLabel("Файлы не выбраны")
        self.app4_info.setWordWrap(True)
        g.addWidget(self.app4_info, row, 1, 1, 2)
        row += 1

        fields = [
            ("Объект:", "object_input"), ("ПРТ:", "prt_input"),
            ("Год:", "year_input"), ("Заказчик:", "customer_input"),
            ("Адрес объекта:", "address_input"),
            ("Разработал:", "developed_input"), ("Проверил:", "checked_input"),
        ]
        for label, attr in fields:
            g.addWidget(QLabel(label), row, 0)
            le = QLineEdit()
            setattr(self, attr, le)
            g.addWidget(le, row, 1, 1, 2)
            row += 1

        dates = [
            ("Дата разработки:", "date_dev"),
            ("Дата проверки:", "date_check"),
            ("Дата начала испытания:", "date_test"),
            ("Дата окончания:", "date_end"),
        ]
        for label, attr in dates:
            g.addWidget(QLabel(label), row, 0)
            de = QDateEdit(calendarPopup=True)
            de.setDate(date.today())
            setattr(self, attr, de)
            g.addWidget(de, row, 1, 1, 2)
            row += 1

        g.addWidget(QLabel("Сохранить как (.docx):"), row, 0)
        self.out_path = QLineEdit()
        g.addWidget(self.out_path, row, 1, 1, 1)
        btn_save = QPushButton("…")
        btn_save.clicked.connect(lambda: self._select(self.out_path, "Word Documents (*.docx)", save=True))
        g.addWidget(btn_save, row, 2)
        row += 1

        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 12)
        g.addWidget(self.progress_bar, row, 0, 1, 3)
        row += 1

        btn_generate = QPushButton("Сгенерировать")
        btn_generate.clicked.connect(self.start_render)
        g.addWidget(btn_generate, row, 0, 1, 3)

        cw.setLayout(g)
        self.setCentralWidget(cw)

    def _select(self, widget: QLineEdit, flt: str, save: bool = False) -> None:
        if save:
            path, _ = QFileDialog.getSaveFileName(self, "Сохранить файл", os.getcwd(), flt)
        else:
            path, _ = QFileDialog.getOpenFileName(self, "Выберите файл", os.getcwd(), flt)
        if path:
            widget.setText(path)

    def _select_dir(self, widget: QLineEdit) -> None:
        d = QFileDialog.getExistingDirectory(self, "Выберите папку со сканами", os.getcwd())
        if d:
            widget.setText(d)

    def choose_app1_images(self):
        dlg = AppendixImagesDialog(
            self,
            title="Приложение 1 — загрузка изображений/скриншотов",
            initial_paths=self.app1_images,
            paste_prefix="app1"
        )
        if dlg.exec() == QDialog.Accepted:
            self.app1_images = dlg.get_paths()
            self.app1_info.setText(f"Выбрано файлов: {len(self.app1_images)}")

    def choose_app4_images(self):
        dlg = AppendixImagesDialog(
            self,
            title="Приложение 4 — загрузка изображений/скриншотов",
            initial_paths=self.app4_images,
            paste_prefix="app4"
        )
        if dlg.exec() == QDialog.Accepted:
            self.app4_images = dlg.get_paths()
            self.app4_info.setText(f"Выбрано файлов: {len(self.app4_images)}")

    def on_mode_changed(self, mode: str) -> None:
        if mode in ("ОQ и PQ", "OQ и PQ"):
            self.tpl_path.setText("")
            self.tests_path.setText("")
            self.xls_path.setText("")
            self.equipment_xls_path.setText("")
            self.scans_dir_input.setText("")
            return

        d = self.defaults[mode]
        self.tpl_path.setText(str(d["tpl"]))
        self.tests_path.setText(str(d["tests"]))
        self.xls_path.setText(str(d["xls_tests"]))
        self.equipment_xls_path.setText(str(d["xls_eq"]))
        self.scans_dir_input.setText(str(d["scans_dir"]))
        self.load_tests(silent=True)
        self.load_equipment(silent=True)

    def load_tests(self, silent: bool = False) -> None:
        try:
            path = Path(self.xls_path.text())
            self.all_tests = io_manager.load_tests_list(path)
            self.selected_tests = []
            if not silent:
                QMessageBox.information(self, "Тесты загружены", f"Всего тестов: {len(self.all_tests)}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка загрузки тестов", str(e))

    def choose_tests(self) -> None:
        if not self.all_tests:
            QMessageBox.warning(self, "Нет тестов", "Тесты не загружены!")
            return

        dlg = TestsDialog(self, self.all_tests, self.selected_tests,
                          calc_launcher=self.launch_test11_calculator)
        if dlg.exec() == QDialog.Accepted:
            self.selected_tests = dlg.get_selected()

    def load_equipment(self, silent: bool = False) -> None:
        try:
            path = Path(self.equipment_xls_path.text())
            by_sheets = io_manager.load_equipment_by_sheets(path)
            if silent:
                self.equipment = [it for items in by_sheets.values() for it in items]
                return
            dlg = EquipmentDialog(self, by_sheets)
            if dlg.exec() == QDialog.Accepted:
                self.equipment = dlg.get_selected()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка загрузки оборудования", str(e))

    def edit_rooms(self) -> None:
        dlg = RoomsDialog(self, self.rooms)
        if dlg.exec() == QDialog.Accepted:
            self.rooms = dlg.get_rooms()

    def launch_test11_calculator(self, *_, auto_close: bool = True) -> int:
        import subprocess

        tpl_path_str = (self.tests_path.text().strip() or "")
        if not tpl_path_str:
            QMessageBox.critical(self, "Калькулятор Тест 11", "Не указан «Документ-тесты (.docx)».")
            return 1

        tpl_path = Path(tpl_path_str).expanduser().resolve()
        if not tpl_path.exists():
            QMessageBox.critical(self, "Калькулятор Тест 11", f"Файл не найден:\n{tpl_path}")
            return 1

        rooms: List[Dict[str, Any]] = list(self.rooms) if getattr(self, "rooms", None) else []
        n = max(1, len(rooms) or 1)

        def _get(room: Dict[str, Any], key: str, default: str) -> str:
            v = room.get(key)
            s = "" if v is None else str(v).strip()
            return s if s else default

        nums = [_get(r, "num", f"{i + 1}") for i, r in enumerate(rooms[:n])]
        names = [_get(r, "name", "") for r in rooms[:n]]
        areas = [_get(r, "area", "1") for r in rooms[:n]]
        filters = [_get(r, "filters", "2") for r in rooms[:n]]
        points = [_get(r, "points", "2") for r in rooms[:n]]
        airflows = [_get(r, "airflow", "") for r in rooms[:n]]

        def _pad(lst, filler):
            need = n - len(lst)
            return lst if need <= 0 else lst + [filler(len(lst) + k + 1) for k in range(need)]

        nums = _pad(nums, lambda k: str(k))
        names = _pad(names, lambda _: "")
        areas = _pad(areas, lambda _: "1")
        filters = _pad(filters, lambda _: "2")
        points = _pad(points, lambda _: "2")
        airflows = _pad(airflows, lambda _: "")

        def _join(xs: List[str]) -> str:
            return ";".join(x.replace(";", " ").strip() for x in xs)

        nums_s, names_s, areas_s = _join(nums), _join(names), _join(areas)
        filters_s, points_s = _join(filters), _join(points)
        airflows_s = _join(airflows)

        proj_root = Path(__file__).resolve().parents[1]
        candidates = [
            proj_root / "tools" / "test11_airflow_calc.py",
            Path.cwd() / "tools" / "test11_airflow_calc.py",
            proj_root / "tests" / "test11_airflow_calc.py",
            Path(__file__).resolve().parent / "test11_airflow_calc.py",
        ]
        script = next((p for p in candidates if p.exists()), None)
        if script is None:
            paths_list = "\n".join(str(p) for p in candidates)
            QMessageBox.critical(
                self, "Калькулятор Тест 11",
                "Не найден скрипт test11_airflow_calc.py по путям:\n\n" + paths_list
            )
            return 1

        tmp_dir = Path(tempfile.mkdtemp(prefix="test11_"))
        auto_save_path = tmp_dir / "test11_auto_save.docx"

        args: List[str] = [
            sys.executable, str(script),
            "--template", str(tpl_path),
            "--rooms", str(n),
            "--nums", nums_s,
            "--names", names_s,
            "--areas", areas_s,
            "--filters", filters_s,
            "--points", points_s,
            "--airflows", airflows_s,
            "--auto-save", str(auto_save_path),
        ]
        if auto_close:
            args.append("--auto-close")

        try:
            proc = subprocess.run(args, check=False)
            rc = int(proc.returncode or 0)
        except Exception as e:
            QMessageBox.critical(self, "Калькулятор Тест 11", f"Ошибка запуска:\n{e}")
            return 1

        if auto_save_path.exists():
            self.tests_path.setText(str(auto_save_path))

        setattr(self, "last_test11_autosave", str(auto_save_path))
        return rc

    def start_render(self) -> None:
        mode = self.mode_combo.currentText()
        if mode == "OQ и PQ":
            self._batch_modes = ["OQ", "PQ"]
            self._start_single_render(self._batch_modes.pop(0))
        else:
            self._batch_modes = []
            self._start_single_render(mode)

    def _start_single_render(self, mode: str) -> None:
        mode = mode.upper()
        use_ui_paths = (self.mode_combo.currentText().upper() == mode)

        if use_ui_paths:
            tpl_path = self.tpl_path.text().strip()
            tests_path = self.tests_path.text().strip()
            xls_tests = self.xls_path.text().strip()
            xls_eq = self.equipment_xls_path.text().strip()
            scans_dir = self.scans_dir_input.text().strip()
            sel_tests_raw = self.selected_tests[:]
            equipment_list = self.equipment[:]
        else:
            d = self.defaults[mode]
            tpl_path = str(d["tpl"])
            tests_path = str(d["tests"])
            xls_tests = str(d["xls_tests"])
            xls_eq = str(d["xls_eq"])
            scans_dir = str(d["scans_dir"])
            try:
                sel_tests_raw = io_manager.load_tests_list(Path(xls_tests))
            except Exception as e:
                QMessageBox.critical(self, f"Ошибка загрузки тестов ({mode})", str(e))
                return
            try:
                by_sheets = io_manager.load_equipment_by_sheets(Path(xls_eq))
                equipment_list = [it for items in by_sheets.values() for it in items]
            except Exception as e:
                QMessageBox.critical(self, f"Ошибка загрузки оборудования ({mode})", str(e))
                return

        app1_images = self.app1_images[:]
        app4_images = self.app4_images[:]

        out_base = Path(self.out_path.text().strip())
        if not out_base.name:
            QMessageBox.critical(self, "Ошибка", "Не указан путь сохранения")
            return
        out_path = out_base.with_name(out_base.stem + f"_{mode}.docx")

        missing_fields = []
        if not tpl_path: missing_fields.append("шаблон")
        if not tests_path: missing_fields.append("документ-тесты")
        if not xls_tests: missing_fields.append("Excel с тестами")
        if not xls_eq: missing_fields.append("Excel оборудование")
        if missing_fields:
            QMessageBox.critical(self, f"Ошибка ({mode})", "Не указаны: " + ", ".join(missing_fields))
            return

        sel_tests = [
            re.sub(r"(?i)^тест\s*\d+(?:\.\d+)?\.?\s*", "", s).strip()
            for s in sel_tests_raw if isinstance(s, str) and s.strip()
        ]

        raw_prt = self.prt_input.text().strip()
        prt_norm = raw_prt.upper()
        for pref in ("ПРТ-OQ-", "ПРТ-PQ-"):
            prt_norm = prt_norm.replace(pref, "")
        full_prt = f"ПРТ-{mode}-{prt_norm}" if prt_norm else ""
        prt = RichText(full_prt, bold=False, italic=False, underline=False)

        object_text = self.object_input.text().strip()
        object_rd = ""
        if object_text:
            morph = MorphAnalyzer()
            words = object_text.split()
            parsed = [morph.parse(w)[0] for w in words]
            adj_idx = next((i for i, p in enumerate(parsed[:-1]) if 'ADJF' in p.tag), None)
            noun_idx = adj_idx + 1 if adj_idx is not None and adj_idx + 1 < len(parsed) else None
            if adj_idx is not None and noun_idx is not None and 'NOUN' in parsed[noun_idx].tag:
                adj = parsed[adj_idx]
                noun = parsed[noun_idx]
                gender = noun.tag.gender
                number = noun.tag.number
                adj_gent = adj.inflect({'gent', gender, number})
                noun_gent = noun.inflect({'gent'})
                if adj_gent and noun_gent:
                    before = " ".join(words[:adj_idx])
                    after = " ".join(words[noun_idx + 1:])
                    object_rd = f"{before} {adj_gent.word} {noun_gent.word} {after}".strip()
                else:
                    object_rd = object_text
            else:
                first_noun = next((p for p in parsed if 'NOUN' in p.tag), None)
                if first_noun:
                    infl = first_noun.inflect({'gent'})
                    if infl:
                        idx = parsed.index(first_noun)
                        object_rd = " ".join(words[:idx] + [infl.word] + words[idx + 1:])
                    else:
                        object_rd = object_text
                else:
                    object_rd = object_text

        ctx_fields: Dict[str, Any] = {
            "объект": object_text,
            "объект1": object_rd,
            "prt": prt,
            "year": self.year_input.text().strip(),
            "customer": self.customer_input.text().strip(),
            "address": self.address_input.text().strip(),
            "Разработал": self.developed_input.text().strip(),
            "Проверил": self.checked_input.text().strip(),
            "Дата_Разработки": self.date_dev.date().toString("dd.MM.yyyy"),
            "Дата_Проверки": self.date_check.date().toString("dd.MM.yyyy"),
            "ДАТА_начала_испытания": self.date_test.date().toString("dd.MM.yyyy"),
            "ДАТА_окончания": self.date_end.date().toString("dd.MM.yyyy"),
            "Тесты_маркированным_списком": RichText("\n".join(f"• {t}" for t in sel_tests), bold=False),
        }

        self.setEnabled(False)
        self.progress_bar.setValue(0)

        self.worker = RenderWorker(
            tpl_path, tests_path, xls_tests, xls_eq, str(out_path),
            sel_tests, self.rooms, equipment_list, ctx_fields,
            scans_dir=scans_dir,
            app1_images=app1_images,
            app4_images=app4_images,
        )
        self.worker.progress.connect(self.on_progress)
        self.worker.finished.connect(self.on_finished)
        self.worker.start()

    def on_progress(self, step: int, message: str) -> None:
        self.progress_bar.setValue(step)
        self.statusBar().showMessage(message)

    def on_finished(self, success: bool, message: str, missing: list[str]) -> None:
        if self._batch_modes:
            next_mode = self._batch_modes.pop(0)
            if success:
                if missing:
                    QMessageBox.warning(self, "Внимание", "Не вставлены тесты:\n" + "\n".join(missing))
                logger.info(message)
            else:
                QMessageBox.critical(self, "Ошибка", message)
                self._batch_modes.clear()
                self.setEnabled(True)
                self.statusBar().clearMessage()
                return
            self._start_single_render(next_mode)
            return

        self.setEnabled(True)
        if success:
            if missing:
                QMessageBox.warning(self, "Внимание", "Не вставлены тесты:\n" + "\n".join(missing))
            QMessageBox.information(self, "Успех", message)
        else:
            QMessageBox.critical(self, "Ошибка", message)
        self.statusBar().clearMessage()


def main():
    app = QApplication(sys.argv)
    wnd = MainWindow()
    wnd.resize(900, 760)
    wnd.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()


from PySide6.QtWidgets import QMainWindow, QLabel
from PySide6.QtCore import Qt

class OldCleanroomWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Старый режим — Чистые помещения")
        label = QLabel("Режим чистых помещений пока реализован в старом виде.")
        label.setAlignment(Qt.AlignCenter)
        self.setCentralWidget(label)


def launch_old_app():
    return OldCleanroomWindow()


# ui/rooms_table.py  (новый файл)
from PySide6.QtWidgets import QTableWidget, QTableWidgetItem, QApplication
from PySide6.QtCore import Qt

class SpreadsheetTable(QTableWidget):
    """
    QTableWidget с excel-копипастой:
      • Ctrl+C — копирует текущее выделение как TSV
      • Ctrl+V — вставляет прямоугольный блок (TSV/CSV/строки)
      • если нет выделения — работаем от (0,0)
      • автоматически добавляет недостающие строки
    """
    def keyPressEvent(self, e):
        ctrl = e.modifiers() & Qt.ControlModifier
        if ctrl and e.key() == Qt.Key_C:
            self._copy_selection_to_clipboard()
            return
        if ctrl and e.key() == Qt.Key_V:
            self._paste_from_clipboard()
            return
        super().keyPressEvent(e)

    # --- helpers ---
    def _copy_selection_to_clipboard(self):
        rngs = self.selectedRanges()
        if rngs:
            r = rngs[0]
            rows = range(r.topRow(), r.bottomRow() + 1)
            cols = range(r.leftColumn(), r.rightColumn() + 1)
        else:
            # нет выделения — копируем всё содержимое
            rows = range(0, self.rowCount())
            cols = range(0, self.columnCount())

        lines = []
        for i in rows:
            vals = []
            for j in cols:
                it = self.item(i, j)
                vals.append("" if it is None else it.text())
            lines.append("\t".join(vals))
        QApplication.clipboard().setText("\n".join(lines))

    def _paste_from_clipboard(self):
        text = QApplication.clipboard().text()
        if not text:
            return

        # поддержка TSV/CSV
        rows_data = [
            [c.strip() for c in row.replace(";", "\t").split("\t")]
            for row in text.splitlines()
            if row.strip() != ""
        ]
        if not rows_data:
            return

        # точка вставки — левый верх выделения, иначе (0,0)
        rngs = self.selectedRanges()
        start_row = rngs[0].topRow() if rngs else 0
        start_col = rngs[0].leftColumn() if rngs else 0

        needed_rows = start_row + len(rows_data)
        if self.rowCount() < needed_rows:
            self.setRowCount(needed_rows)

        needed_cols = start_col + max(len(r) for r in rows_data)
        if self.columnCount() < needed_cols:
            self.setColumnCount(needed_cols)

        for i, row_vals in enumerate(rows_data):
            for j, val in enumerate(row_vals):
                r = start_row + i
                c = start_col + j
                it = self.item(r, c)
                if it is None:
                    it = QTableWidgetItem()
                    self.setItem(r, c, it)
                it.setText(val)


# ui/startup_dialog.py
from PySide6.QtWidgets import QDialog, QVBoxLayout, QPushButton

class StartupDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Выберите режим работы")
        self.resize(300, 150)
        self.choice = None

        layout = QVBoxLayout(self)

        btn_eq = QPushButton("Оборудование")
        btn_cr = QPushButton("Чистые помещения")

        btn_eq.clicked.connect(lambda: self._select("equipment"))
        btn_cr.clicked.connect(lambda: self._select("cleanrooms"))

        layout.addWidget(btn_eq)
        layout.addWidget(btn_cr)

    def _select(self, choice):
        self.choice = choice
        self.accept()

    def get_choice(self):
        return self.choice


from PySide6.QtWidgets import QApplication

def apply_compact_style(app: QApplication) -> None:
    app.setStyle("Fusion")
    app.setStyleSheet("""
        QWidget { font-size: 10pt; }
        QLineEdit, QComboBox, QDateEdit {
            padding: 5px 8px;
            border-radius: 6px;
        }
        QPushButton {
            padding: 6px 10px;
            border-radius: 6px;
        }
        QToolButton {
            padding: 4px 6px;
            border-radius: 6px;
        }
        QGroupBox {
            margin-top: 10px;
            font-weight: 600;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 4px;
        }
    """)



