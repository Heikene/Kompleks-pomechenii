import re
from copy import deepcopy
from typing import List, Dict
from docx.text.paragraph import Paragraph
from decimal import Decimal

def clear_paragraph(paragraph: Paragraph):
    for run in paragraph.runs:
        paragraph._p.remove(run._r)

Paragraph.clear = clear_paragraph
from docx import Document
from docx.table import Table
from docx.oxml import OxmlElement, CT_P, CT_Tbl
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from logger import logger


# -----------------------------------------------------------------------------
# 1) ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ДЛЯ ТЕСТОВЫХ ТАБЛИЦ
# -----------------------------------------------------------------------------
TEST_TITLE_RE = re.compile(r"^\s*Тест\s*11[\.\s]", re.IGNORECASE)

def make_test_titles_bold(doc: Document) -> None:
    """Делает жирным именно фразы вида 'Тест 11.x ...' в первой строке таблиц тестов."""
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        # работаем только с таблицами, у которых в первой строке есть 'Тест 11'
        if not _is_test11_table(tbl):
            # но всё равно проверим параграфы первой строки на всякий случай
            pass
        for cell in tbl.rows[0].cells:
            for p in cell.paragraphs:
                # ищем сам заголовок 'Тест 11.x ...' (а не служебные подписи типа 'Дата проведения')
                if TEST_TITLE_RE.search(p.text or ""):
                    if not p.runs:
                        p.add_run("")
                    for r in p.runs:
                        _force_face_only_tnr(r)   # гарнитура = TNR
                        r.font.bold = True        # только жирность
def find_table_obj_by_title(src_doc: Document, title: str) -> Table:
    sub = title.split('.', 1)[1] if '.' in title else title
    target = re.sub(r'[^\w\s]', '', sub.lower()).strip()
    for tbl in src_doc.tables:
        all_text = " ".join(cell.text for row in tbl.rows for cell in row.cells).lower()
        if target in re.sub(r'[^\w\s]', '', all_text).strip():
            return tbl
    body = src_doc.element.body
    for idx, child in enumerate(body):
        if isinstance(child, CT_P):
            txt = "".join(n.text for n in child.iter() if n.text).lower()
            norm = re.sub(r'[^\w\s]', '', txt).strip()
            if target in norm:
                for nxt in body[idx+1:]:
                    if isinstance(nxt, CT_Tbl):
                        for t in src_doc.tables:
                            if t._tbl is nxt:
                                return t
    raise ValueError(f"Test-table «{title}» not found.")


def create_page_break() -> OxmlElement:
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def fix_table_xml(tbl_xml: OxmlElement) -> OxmlElement:
    tblPr = tbl_xml.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_xml.insert(0, tblPr)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    return tbl_xml


# ---------------------------------------------------------------------------
# ХЕЛПЕРЫ: шрифт Times New Roman и распознавание тест-таблиц
# ---------------------------------------------------------------------------


def _force_face_only_tnr(run) -> None:
    """Ставит TNR для всех семейств (ascii/hAnsi/cs/eastAsia), НЕ меняя размер и жирность."""
    run.font.name = "Times New Roman"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), "Times New Roman")


def _apply_face_only_to_paragraphs(paragraphs) -> None:
    for p in paragraphs:
        if not p.runs:
            p.add_run("")
        for r in p.runs:
            _force_face_only_tnr(r)

def enforce_tnr_face_only_everywhere(doc: Document) -> None:
    """Ставит TNR для всего документа, не меняя размер/жирность/курсив и т.п."""
    # Тело документа: абзацы вне таблиц
    _apply_face_only_to_paragraphs(doc.paragraphs)

    # Таблицы в теле документа
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                _apply_face_only_to_paragraphs(cell.paragraphs)

    # Колонтитулы всех секций
    for sec in doc.sections:
        _apply_face_only_to_paragraphs(sec.header.paragraphs)
        for t in sec.header.tables:
            for row in t.rows:
                for cell in row.cells:
                    _apply_face_only_to_paragraphs(cell.paragraphs)

        _apply_face_only_to_paragraphs(sec.footer.paragraphs)
        for t in sec.footer.tables:
            for row in t.rows:
                for cell in row.cells:
                    _apply_face_only_to_paragraphs(cell.paragraphs)



def _is_test11_table(tbl) -> bool:
    if not tbl.rows or not tbl.rows[0].cells:
        return False
    head = " ".join(c.text for c in tbl.rows[0].cells)
    return re.search(r"^\s*Тест\s*11[\.\s]", head, flags=re.IGNORECASE) is not None

def _force_run_font_tnr(run, size_pt: int, bold: bool | None = None):
    """
    Жёстко ставим Times New Roman указанного размера для всех семейств (ascii, hAnsi, cs, eastAsia),
    чтобы кириллица не «прыгала» в Calibri и т.п.
    """
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), "Times New Roman")


def _looks_like_test_table(tbl) -> bool:
    """
    Распознаём ВСЕ тестовые таблицы:
      • «результатные»: содержит «номер помещения» и «площад» и («точк» или «№»),
        ИЛИ содержит «результат» и «испытан»;
      • «описательные»: содержит хотя бы одно из: «дата проведения»,
        «контролируемый параметр», «критерий приемлемости», «квалификационное испытание».
    """
    if not tbl.rows:
        return False
    hdr = " ".join(c.text for c in tbl.rows[0].cells).lower()

    # результатные варианты
    points_like = ("номер помещения" in hdr and "площад" in hdr and ("точк" in hdr or "№" in hdr))
    results_like = ("результат" in hdr and "испытан" in hdr)

    # описательные варианты
    desc_keys = ("дата провед", "контролируемый параметр", "критерий приемлемости", "квалификационное испытание")
    desc_like = any(k in hdr for k in desc_keys)

    return points_like or results_like or desc_like




# def _apply_tnr_to_test_table(tbl):
#     """Шапка: TNR 11pt жирный, тело: TNR 10pt."""
#     for ri, row in enumerate(tbl.rows):
#         is_header = (ri == 0)
#         for cell in row.cells:
#             for p in cell.paragraphs:
#                 for run in p.runs:
#                     _force_run_font_tnr(run, 11 if is_header else 10, bold=is_header)
#
#     try:
#         for tbl in doc.tables:
#             if _looks_like_test_table(tbl):
#                 _apply_tnr_to_test_table(tbl)
#                 _zero_cell_margins(tbl)
#     except NameError:
#         pass


def _zero_cell_margins(tbl) -> None:
    """
    Обнуляет внутренние поля ячеек таблицы (top/left/bottom/right = 0).
    Делается через поиск узлов w:tblCellMar и w:top/left/bottom/right.
    """
    t = tbl._tbl
    tblPr = t.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        t.insert(0, tblPr)

    # w:tblCellMar
    mar = tblPr.find(qn('w:tblCellMar'))
    if mar is None:
        mar = OxmlElement('w:tblCellMar')
        tblPr.append(mar)

    # w:top/left/bottom/right
    for side in ('top', 'left', 'bottom', 'right'):
        el = mar.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            mar.append(el)
        el.set(qn('w:w'), '0')     # 0 twips
        el.set(qn('w:type'), 'dxa')


def _format_table_tnr_no_margins(tbl) -> None:
    # 0) отключить авто-подгон
    tbl.allow_autofit = False

    # 1) обнулить внутренние поля ячеек (снимает «поля»/паддинги)
    _zero_cell_margins(tbl)

    # 2) абзацные параметры: без отступов/интервалов, межстрочный ~1.1
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.first_line_indent = Cm(0)
                pf.left_indent = Cm(0)
                pf.right_indent = Cm(0)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.1
                # на всякий случай выравнивание влево, чтобы не выглядело как «отступ»
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 3) только гарнитура TNR (без изменения размера/жирности)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if not p.runs:
                    p.add_run("")
                for r in p.runs:
                    _force_face_only_tnr(r)


# -----------------------------------------------------------------------------
# 2) ОБРАБОТКА ТАБЛИЦ ПОМЕЩЕНИЙ
# -----------------------------------------------------------------------------

def process_rooms_table(doc: Document, rooms: List[Dict[str, str]]) -> None:
    logger.debug("Начинаем process_rooms_table")
    tbl = next(
        (t for t in doc.tables
         if t.rows and "Номер помещения" in t.rows[0].cells[0].text),
        None
    )
    if not tbl:
        raise ValueError("Таблица помещений не найдена (header 'Номер помещения').")

    header_map = {
        "#": "Номер помещения", "##": "Наименование помещений",
        "###": "Класс чистоты", "####": "Площадь, м²",
        "#$": "Объём, м³", "#$$": "Перепад давления, Па (±5 Па)",
        "#$$$": "Расход приточного воздуха, м³/ч", "#%": "Кратность воздухообмена, ч⁻¹, не менее",
        "#%%": "Температура, °C", "#%%%": "Относительная влажность, %"
    }
    hdr = tbl.rows[0]
    hdr.height = Cm(3.4)
    hdr.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for cell in hdr.cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Центрирование
            pf = para.paragraph_format
            pf.left_indent = Cm(0)
            pf.right_indent = Cm(0)
            pf.first_line_indent = Cm(0)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.1
            for run in para.runs:
                key = run.text.strip()
                if key in header_map:
                    run.text = header_map[key]
                    run.font.name = "Times New Roman"
                    _force_face_only_tnr(run)

    ph2key = {
        "#": "num", "##": "name", "###": "klass", "####": "area",
        "#$": "volume", "#$$": "dp", "#$$$": "airflow", "#%": "exchange",
        "#%%": "temp", "#%%%": "rh"
    }

    # Шаблонная строка
    sample_tr = None
    for row in tbl.rows[1:]:
        if any(cell.text.strip() in ph2key for cell in row.cells):
            sample_tr = deepcopy(row._tr)
            break

    if sample_tr is None:
        raise ValueError("В таблице помещений не найден placeholder row.")

    # Удаляем все строки после заголовка
    while len(tbl.rows) > 1:
        tbl._tbl.remove(tbl.rows[1]._tr)

    for room in rooms:
        tbl._tbl.append(deepcopy(sample_tr))
        new_row = tbl.rows[-1]
        # Устанавливаем высоту строки и вертикальное выравнивание
        new_row.height = Cm(1.46)
        new_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for cell in new_row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for idx, cell in enumerate(new_row.cells):
            placeholder = cell.text.strip()
            cell.text = ""
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf = para.paragraph_format
                pf.first_line_indent = Cm(0)
                pf.left_indent = Cm(0)
                pf.right_indent = Cm(0)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.1
            if placeholder in ph2key:
                run = cell.paragraphs[0].add_run(room.get(ph2key[placeholder], ""))
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    logger.debug("process_rooms_table выполнен")



# -----------------------------------------------------------------------------
# 3) ОБРАБОТКА ТАБЛИЦ ОБОРУДОВАНИЯ
# -----------------------------------------------------------------------------

def process_equipment_table(doc: Document, equipment: List[Dict[str, str]]) -> None:
    """
    Заполняет таблицу средств измерений.
    Ищем таблицу, где в ЛЮБОЙ строке встречаются плейсхолдеры @1..@4.
    Строка, содержащая любой из @1..@4, используется как шаблон.
    Все строки после неё удаляются и заменяются клонами с данными из `equipment`.
    """
    logger.debug("Начинаем process_equipment_table")

    PH = ("@1", "@2", "@3", "@4")

    # 1) Найти таблицу по наличию плейсхолдеров в ЛЮБОЙ строке
    tbl = None
    for t in doc.tables:
        found = False
        for r in t.rows:
            if any(any(ph in c.text for ph in PH) for c in r.cells):
                found = True
                break
        if found:
            tbl = t
            break
    if tbl is None:
        raise ValueError("Таблица оборудования не найдена (в таблице нет @1..@4).")

    # 2) Найти шаблонную строку — ту, где реально лежат @1..@4
    sample_idx = None
    for i, row in enumerate(tbl.rows):
        if any(any(ph in c.text for ph in PH) for c in row.cells):
            sample_idx = i
            break
    if sample_idx is None:
        raise ValueError("Не нашли шаблонную строку в таблице оборудования.")

    # Сохраняем все строки ДО шаблонной (это заголовочный блок)
    header_rows_to_keep = sample_idx
    sample_tr = deepcopy(tbl.rows[sample_idx]._tr)

    # 3) Удалить все строки, начиная с sample_idx
    for i in range(len(tbl.rows) - 1, header_rows_to_keep - 1, -1):
        tbl._tbl.remove(tbl.rows[i]._tr)

    # 4) Заполнять
    for eq in equipment:
        tbl._tbl.append(deepcopy(sample_tr))
        new_row = tbl.rows[-1]

        # Сформировать значения
        date_str = (eq.get("date") or "").strip()
        until_str = (eq.get("until") or "").strip()
        date_full = f"{date_str} / {until_str}".strip(" /")

        values = [
            eq.get("name_sn", "") or "",   # @1
            eq.get("params", "") or "",    # @2
            eq.get("cert", "") or "",      # @3
            date_full,                     # @4
        ]

        # 5) Очистка содержимого ячеек строки-шаблона и запись текста
        for ci, cell in enumerate(new_row.cells):
            # чистим все ранны в абзацах
            for p in cell.paragraphs:
                # формат абзаца
                pf = p.paragraph_format
                pf.first_line_indent = Cm(0)
                pf.left_indent = Cm(0)
                pf.right_indent = Cm(0)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.1
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                for r in list(p.runs):
                    p._p.remove(r._r)

            text = values[ci] if ci < len(values) else ""
            run = cell.paragraphs[0].add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    logger.debug("process_equipment_table выполнен")





# -----------------------------------------------------------------------------
# 4) ВСТАВКА ТЕСТОВЫХ ТАБЛИЦ
# -----------------------------------------------------------------------------
def extract_total_flows_from_test11(doc: Document) -> list[Decimal]:
    """
    Возвращает список «фактический суммарный» по помещениям в порядке следования,
    сканируя таблицу «Тест 11. Проверка расхода приточного воздуха».
    Берём столбец с заголовком, содержащим 'фактический суммарный',
    и собираем непустые верхушки vMerge-блоков (по одной на помещение).
    """
    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip().lower()

    for tbl in doc.tables:
        if not tbl.rows:
            continue
        hdr = " ".join(norm(c.text) for c in tbl.rows[0].cells)
        if "расход приточного воздуха" in hdr and "фактический суммарный" in hdr:
            # найти индекс колонки «фактический суммарный»
            fact_sum_col = None
            for ci, c in enumerate(tbl.rows[0].cells):
                if "фактический суммарный" in norm(c.text):
                    fact_sum_col = ci
                    break
            if fact_sum_col is None:
                continue

            vals: list[Decimal] = []
            last_seen = None
            for r in tbl.rows[1:]:
                if fact_sum_col >= len(r.cells):
                    continue
                t = norm(r.cells[fact_sum_col].text)
                # в vMerge верхняя ячейка содержит число, нижние — пустые
                if t and t != last_seen:
                    # вытащим число (запятая/точка)
                    num = re.sub(r"[^\d,.\-]", "", t).replace(",", ".")
                    try:
                        vals.append(Decimal(num))
                    except Exception:
                        pass
                    last_seen = t
            return vals
    return []



def insert_test_tables(
    doc: Document,
    tests_docx_path: str,
    selected_tests: List[str]
) -> List[str]:
    """
    Вставляет выбранные тест-таблицы из tests_docx_path в место плейсхолдера {{TABLE}}
    (поддерживаются варианты {{ TABLE }}, __TABLE__, ___TABLE_PLACEHOLDER___),
    перенумеровывает заголовки "Тест 11.x …" и приводит таблицы к единому формату:
    Times New Roman (шапка 11pt жирный, тело 10pt), без абзацных отступов и без внутренних полей ячеек.
    """
    src = Document(tests_docx_path)
    elems: List[OxmlElement] = []
    missing: List[str] = []

    # 1) подобрать и подготовить таблицы из источника
    for idx, title in enumerate(selected_tests, start=1):
        try:
            tbl = find_table_obj_by_title(src, title)

            # перенумеровка заголовка "Тест 11.{idx} ..."
            header_cell = tbl.rows[0].cells[0]
            header_text = header_cell.text.strip()
            m = re.match(r'(Тест)\s*\d+(?:\.\d+)?\.?\s*(.+)', header_text)
            if m:
                base, rest = m.groups()
                new_title = f"{base} 11.{idx} {rest}"
                p0 = header_cell.paragraphs[0]
                for run in list(p0.runs):
                    p0._p.remove(run._r)
                r = p0.add_run(new_title)
                _force_face_only_tnr(r)  # только гарнитура
                r.font.bold = True

            elems.append(fix_table_xml(deepcopy(tbl._tbl)))
        except Exception as ex:
            missing.append(f"{title}: {ex}")

    # 2) вставка вместо плейсхолдера
    PLACEHOLDERS = ("{{TABLE}}", "{{ TABLE }}", "__TABLE__", "___TABLE_PLACEHOLDER___")
    anchor_found = False
    for para in doc.paragraphs:
        full = "".join(run.text for run in para.runs) or para.text or ""
        if any(ph in full for ph in PLACEHOLDERS):
            parent = para._p.getparent()
            pos = list(parent).index(para._p)
            parent.remove(para._p)
            for i, xml in enumerate(elems):
                parent.insert(pos, xml); pos += 1
                if i < len(elems) - 1:
                    parent.insert(pos, create_page_break()); pos += 1
            anchor_found = True
            break

    if not anchor_found and elems:
        logger.warning("Плейсхолдер таблиц не найден. Убедитесь, что в шаблоне есть {{ TABLE }} в отдельном абзаце.")

    # 3) единая типографика/шрифт/отступы для ВСТАВЛЕННЫХ тестовых таблиц
    for tbl in doc.tables:
        tbl.allow_autofit = False
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.left_indent = Cm(0)
                    pf.right_indent = Cm(0)
                    pf.first_line_indent = Cm(0)
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = 1.1
                    pf.contextual_spacing = False

        # Шапка: Times New Roman 11, жирный
        hdr = tbl.rows[0]
        for cell in hdr.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    _force_face_only_tnr(r)

# -----------------------------------------------------------------------------
# 5) ЗАПОЛНЕНИЕ РЕЗУЛЬТАТОВ ТЕСТОВЫХ ТАБЛИЦ ДАННЫМИ ПОМЕЩЕНИЙ
# -----------------------------------------------------------------------------

def process_test_results_tables(doc: Document, rooms: List[Dict[str, str]]) -> None:
    """
    Ищет в результатных тестовых таблицах блоки точек измерений и:
      - объединяет колонки 0 и 1 ТОЛЬКО по строкам с номером точки (цифра в колонке "№ точки"),
      - НЕ включает строку "Среднее" в объединение (это ключ к исправлению),
      - центрирует ячейки блока (кроме первых двух),
      - в конце приводит тест-таблицы к единому формату (если есть _looks_like_test_table/_format_table_tnr_no_margins).
    """
    logger.debug("Начинаем process_test_results_tables")

    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip()

    def is_results_header(row) -> bool:
        j = " ".join(norm(c.text).lower() for c in row.cells)
        return ("номер" in j and "помещени" in j) and ("площад" in j) and ("точк" in j or "№" in j)

    def is_comments_row(row) -> bool:
        return "КОММЕНТАРИИ" in norm(row.cells[0].text).upper() if row.cells else False

    def find_point_col(header_row) -> int | None:
        for ci, c in enumerate(header_row.cells):
            t = norm(c.text).lower()
            # "№ точки" / "точк" / просто "№"
            if "точк" in t or t == "№" or "№ точки" in t:
                return ci
        return None

    # Достаём номер точки из ячейки. Для "1¤", "1α", "1." и т.п. — берём первую группу цифр.
    def get_point_number(cell_text: str) -> int | None:
        t = norm(cell_text)
        m = re.search(r"\d+", t)
        if not m:
            return None
        try:
            return int(m.group(0))
        except Exception:
            return None

    # Полная перезапись текста ячейки без изменения tcPr (vMerge и т.п. остаётся)
    def rewrite_cell_text_center(cell, text: str, size_pt: int = 10):
        # удалить все параграфы
        for p in list(cell.paragraphs):
            p._element.getparent().remove(p._element)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0)
        pf.right_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.1

        r = p.add_run(text or "")
        r.font.name = "Times New Roman"
        r.font.size = Pt(size_pt)

    for tbl in doc.tables:
        if not tbl.rows:
            continue

        # 1) Найти строку шапки результата
        hdr_idx = None
        for i, row in enumerate(tbl.rows):
            if is_results_header(row):
                hdr_idx = i
        if hdr_idx is None:
            continue

        hdr_row = tbl.rows[hdr_idx]

        # 2) Найти колонку "№ точки"
        point_col = find_point_col(hdr_row)
        if point_col is None:
            # запасной вариант (часто это 2)
            point_col = 2
        if point_col < 0:
            continue

        # 3) Диапазон данных: от строки после шапки до "КОММЕНТАРИИ" или конца таблицы
        end_idx = len(tbl.rows)
        for i in range(hdr_idx + 1, len(tbl.rows)):
            if is_comments_row(tbl.rows[i]):
                end_idx = i
                break

        data_start = hdr_idx + 1
        if data_start >= end_idx:
            continue

        # 4) Собрать блоки ТОЛЬКО из непрерывных строк с номером точки
        #    + дополнительно: если встретили "1" и блок уже начат — начинаем новый (новый фильтр)
        blocks: list[list[int]] = []
        cur: list[int] = []

        for ri in range(data_start, end_idx):
            row = tbl.rows[ri]
            if point_col >= len(row.cells):
                # если вдруг “кривой” ряд — закрываем текущий блок
                if cur:
                    blocks.append(cur)
                    cur = []
                continue

            num = get_point_number(row.cells[point_col].text)

            if num is None:
                # это "Среднее" / пусто / текст — закрываем блок
                if cur:
                    blocks.append(cur)
                    cur = []
                continue

            # если новая серия начинается с 1 — закрываем предыдущую серию
            if num == 1 and cur:
                blocks.append(cur)
                cur = []

            cur.append(ri)

        if cur:
            blocks.append(cur)

        # 5) Применить объединения (колонки 0 и 1) только по строкам точек
        for block in blocks:
            if len(block) <= 1:
                continue

            top = block[0]

            # Колонка 0
            try:
                c0 = tbl.rows[top].cells[0]
                v0 = norm(c0.text)
                for ri in block[1:]:
                    c0 = c0.merge(tbl.rows[ri].cells[0])
                rewrite_cell_text_center(c0, v0, size_pt=10)
            except Exception as e:
                logger.warning(f"Не удалось объединить колонку 0 в блоке {block}: {e}")

            # Колонка 1
            try:
                c1 = tbl.rows[top].cells[1]
                v1 = norm(c1.text)
                for ri in block[1:]:
                    c1 = c1.merge(tbl.rows[ri].cells[1])
                rewrite_cell_text_center(c1, v1, size_pt=10)
            except Exception as e:
                logger.warning(f"Не удалось объединить колонку 1 в блоке {block}: {e}")

            # Центрирование остальных колонок внутри блока (только строк точек)
            for ri in block:
                row = tbl.rows[ri]
                for ci, cell in enumerate(row.cells):
                    if ci >= 2 and cell.paragraphs:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ФИНАЛ: привести тест-таблицы к общему виду (если ваши хелперы определены)
    try:
        for t in doc.tables:
            if _looks_like_test_table(t):
                _format_table_tnr_no_margins(t)
    except Exception:
        pass

    logger.debug("process_test_results_tables завершён")
