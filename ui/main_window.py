# ui/main_window.py
from __future__ import annotations

import sys
import os
import re
import tempfile
import uuid
from copy import deepcopy
from collections import OrderedDict
from datetime import datetime, date
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from ui import word_table2_otch_splitter

from PySide6.QtCore import Qt, QThread, Signal
from PySide6.QtGui import QDragEnterEvent, QDropEvent, QGuiApplication, QKeySequence, QShortcut
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QListWidget,
    QListWidgetItem, QPushButton, QMessageBox, QFileDialog,
    QGridLayout, QDateEdit, QDialog, QVBoxLayout, QHBoxLayout,
    QSpinBox, QTableWidget, QTableWidgetItem, QProgressBar, QComboBox,
    QTabWidget, QAbstractItemView
)

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxtpl import DocxTemplate, RichText, InlineImage
from docx.shared import Mm, Pt
from docx.image.exceptions import UnrecognizedImageError
from docx.image.image import Image

from pymorphy3 import MorphAnalyzer

from file_utils import temp_docx
import io_manager
import template_renderer
import table_processor
from logger import logger

from risk_table5 import get_risk_rows, insert_table5_into_doc
import word_table5_splitter
import word_update_all




# =============================================================================
# 1) Защита от падения tpl.render() из-за битых/неподдерживаемых картинок
# =============================================================================
# docxtpl вставляет картинки лениво: ошибка может прилететь при tpl.render(),
# когда InlineImage превращается в строку. Этот патч гарантирует "не упасть".
from docxtpl.inline_image import InlineImage as _TplInlineImage  # noqa

_old_inlineimage_str = _TplInlineImage.__str__


def _safe_inlineimage_str(self):
    try:
        return _old_inlineimage_str(self)
    except UnrecognizedImageError:
        desc = getattr(self, "image_descriptor", None)
        logger.warning(f"Пропущена битая/неподдерживаемая картинка (UnrecognizedImageError): {desc}")
        return ""
    except Exception as e:
        desc = getattr(self, "image_descriptor", None)
        logger.warning(f"Пропущена картинка (ошибка вставки: {e}): {desc}")
        return ""


_TplInlineImage.__str__ = _safe_inlineimage_str


# =============================================================================
# 2) Безопасное создание InlineImage + фильтрация расширений
# =============================================================================
_IMG_EXT = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif"}  # только картинки (НЕ pdf)


def safe_inline_image(tpl: DocxTemplate, path: str, *, width_mm: int = 160, label: str = ""):
    """
    1) Отсекаем несуществующие / не файлы
    2) Отсекаем расширения (только картинки)
    3) Проверяем, что это реально картинка: Image.from_file()
    4) Создаём InlineImage
    """
    if not path:
        return None

    pp = Path(path)
    if not pp.exists() or not pp.is_file():
        logger.warning(f"Пропущен файл для {label} (не найден): {path}")
        return None

    ext = pp.suffix.lower()
    if ext not in _IMG_EXT:
        logger.warning(f"Пропущен файл для {label} (не картинка / pdf не поддерживается): {path}")
        return None

    # ранняя проверка (иначе может упасть позже на render)
    try:
        Image.from_file(str(pp))
    except UnrecognizedImageError:
        logger.warning(f"Пропущен файл для {label} (битая/неподдерживаемая картинка): {path}")
        return None
    except Exception as e:
        logger.warning(f"Пропущен файл для {label} (ошибка проверки картинки: {e}): {path}")
        return None

    try:
        return InlineImage(tpl, str(pp), width=Mm(width_mm))
    except Exception as e:
        logger.warning(f"Пропущен файл для {label} (ошибка создания InlineImage: {e}): {path}")
        return None


def make_inline_images(tpl: DocxTemplate, paths: list[str], *, label: str, width_mm: int = 160) -> list[InlineImage]:
    out: list[InlineImage] = []
    for p in (paths or []):
        ii = safe_inline_image(tpl, p, width_mm=width_mm, label=label)
        if ii is not None:
            out.append(ii)
    return out


# =============================================================================

# =============================================================================
# 4) Spreadsheet-like table
# =============================================================================
class SpreadsheetTable(QTableWidget):
    """QTableWidget с excel-копипастой."""
    def keyPressEvent(self, e):
        ctrl = bool(e.modifiers() & Qt.ControlModifier)
        if ctrl and e.key() == Qt.Key_C:
            self._copy_selection_to_clipboard()
            return
        if ctrl and e.key() == Qt.Key_V:
            self._paste_from_clipboard()
            return
        super().keyPressEvent(e)

    def _copy_selection_to_clipboard(self):
        rngs = self.selectedRanges()
        if rngs:
            r = rngs[0]
            rows = range(r.topRow(), r.bottomRow() + 1)
            cols = range(r.leftColumn(), r.rightColumn() + 1)
        else:
            rows = range(0, self.rowCount())
            cols = range(0, self.columnCount())

        lines = []
        for i in rows:
            vals = []
            for j in cols:
                it = self.item(i, j)
                vals.append("" if it is None else it.text())
            lines.append("\t".join(vals))
        QApplication.clipboard().setText("\n".join(lines))

    def _paste_from_clipboard(self):
        text = QApplication.clipboard().text()
        if not text:
            return

        rows_data = [
            [c.strip() for c in row.replace(";", "\t").split("\t")]
            for row in text.splitlines()
            if row.strip() != ""
        ]
        if not rows_data:
            return

        rngs = self.selectedRanges()
        start_row = rngs[0].topRow() if rngs else max(0, self.currentRow())
        start_col = rngs[0].leftColumn() if rngs else max(0, self.currentColumn())

        need_r = start_row + len(rows_data)
        need_c = start_col + max(len(r) for r in rows_data)

        if self.rowCount() < need_r:
            self.setRowCount(need_r)
        if self.columnCount() < need_c:
            self.setColumnCount(need_c)

        for i, row_vals in enumerate(rows_data):
            for j, val in enumerate(row_vals):
                r = start_row + i
                c = start_col + j
                it = self.item(r, c)
                if it is None:
                    it = QTableWidgetItem()
                    self.setItem(r, c, it)
                it.setText(val)


# =============================================================================
# 5) Dialogs
# =============================================================================
class RoomsDialog(QDialog):
    def __init__(self, parent=None, initial_rooms: List[Dict[str, str]] | None = None):
        super().__init__(parent)
        self.setWindowTitle("Настройки помещений")
        self.resize(700, 450)

        self.spin = QSpinBox()
        self.spin.setRange(0, 100)
        self.spin.setValue(len(initial_rooms or []))
        self.spin.valueChanged.connect(self._on_count)

        self.tbl = SpreadsheetTable()
        self.tbl.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.tbl.setSelectionBehavior(QAbstractItemView.SelectItems)
        self.tbl.setEditTriggers(QAbstractItemView.AllEditTriggers)

        headers = [
            "Номер помещения", "Наименование", "Класс чистоты",
            "Площадь, м²", "Объём, м³", "Δдавл., Pa",
            "Расход, м³/ч", "Кратность, не менее", "Темп., °C", "RH, %"
        ]
        self.tbl.setColumnCount(len(headers))
        self.tbl.setHorizontalHeaderLabels(headers)

        self._on_count(self.spin.value())
        if initial_rooms:
            for r, data in enumerate(initial_rooms):
                for c, key in enumerate(
                    ["num", "name", "klass", "area", "volume",
                     "dp", "airflow", "exchange", "temp", "rh"]
                ):
                    self.tbl.setItem(r, c, QTableWidgetItem(data.get(key, "")))

        btn_ok = QPushButton("ОК")
        btn_cancel = QPushButton("Отмена")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        hb = QHBoxLayout()
        hb.addStretch()
        hb.addWidget(btn_ok)
        hb.addWidget(btn_cancel)

        lay = QVBoxLayout(self)
        lay.addWidget(QLabel("Количество помещений:"))
        lay.addWidget(self.spin)
        lay.addWidget(self.tbl)
        lay.addLayout(hb)

    def _on_count(self, n: int) -> None:
        self.tbl.setRowCount(n)
        for r in range(n):
            for c in range(self.tbl.columnCount()):
                if not self.tbl.item(r, c):
                    self.tbl.setItem(r, c, QTableWidgetItem(""))

    def get_rooms(self) -> List[Dict[str, str]]:
        keys = ["num", "name", "klass", "area", "volume",
                "dp", "airflow", "exchange", "temp", "rh"]
        out: List[Dict[str, str]] = []
        for r in range(self.tbl.rowCount()):
            row: Dict[str, str] = {}
            for c, k in enumerate(keys):
                it = self.tbl.item(r, c)
                row[k] = (it.text().strip() if it else "")
            out.append(row)
        return out


class EquipmentDialog(QDialog):
    def __init__(self, parent=None, items_by_sheet: dict[str, list[dict]] | None = None):
        super().__init__(parent)
        self.setWindowTitle("Выбор оборудования")
        self.resize(600, 420)

        self.tabs = QTabWidget(self)
        self._lists: dict[str, QListWidget] = {}

        for sheet, items in (items_by_sheet or {}).items():
            lw = QListWidget()
            lw.setSelectionMode(QListWidget.MultiSelection)
            for it in items:
                item = QListWidgetItem(it.get("name_sn", ""))
                item.setData(Qt.UserRole, it)
                lw.addItem(item)
            self.tabs.addTab(lw, f"{sheet} ({len(items)})")
            self._lists[sheet] = lw

        btn_ok = QPushButton("ОК")
        btn_cancel = QPushButton("Отмена")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        hb = QHBoxLayout()
        hb.addStretch()
        hb.addWidget(btn_ok)
        hb.addWidget(btn_cancel)

        lay = QVBoxLayout(self)
        lay.addWidget(self.tabs)
        lay.addLayout(hb)

    def get_selected(self) -> list[dict]:
        res: list[dict] = []
        for lw in self._lists.values():
            res.extend(i.data(Qt.UserRole) for i in lw.selectedItems())
        return res


class TestsDialog(QDialog):
    """Выбор тестов + кнопка калькулятора рядом с тестом 11."""
    def __init__(self, parent=None, items: List[str] | None = None,
                 selected: List[str] | None = None,
                 calc_launcher=None):
        super().__init__(parent)
        self.setWindowTitle("Выберите тесты")
        self.resize(600, 400)

        self.selection_order: List[int] = []
        self.calc_launcher = calc_launcher

        lay = QVBoxLayout(self)
        self.list = QListWidget(self)
        self.list.setSelectionMode(QAbstractItemView.MultiSelection)
        lay.addWidget(self.list)

        self._base_text_by_row: list[str] = []
        self._label_by_row: list[QLabel] = []

        rx_calc = re.compile(r"(?i)тест\s*11.*расход.*приточ", re.IGNORECASE)

        for txt in (items or []):
            item = QListWidgetItem()
            roww = QWidget()
            hb = QHBoxLayout(roww)
            hb.setContentsMargins(6, 2, 6, 2)

            lbl = QLabel(txt)
            hb.addWidget(lbl)
            hb.addStretch()

            if rx_calc.search(txt) and callable(self.calc_launcher):
                btn = QPushButton("Кальк.")
                btn.setToolTip("Открыть калькулятор Тест 11")
                btn.clicked.connect(self.calc_launcher)
                hb.addWidget(btn)

            self.list.addItem(item)
            self.list.setItemWidget(item, roww)
            item.setSizeHint(roww.sizeHint())

            self._base_text_by_row.append(txt)
            self._label_by_row.append(lbl)

        if selected:
            by_text = {t: i for i, t in enumerate(self._base_text_by_row)}
            for txt in selected:
                r = by_text.get(txt)
                if r is not None:
                    self.list.item(r).setSelected(True)
                    self.selection_order.append(r)

        btn_ok = QPushButton("ОК")
        btn_cancel = QPushButton("Отмена")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        hb2 = QHBoxLayout()
        hb2.addStretch()
        hb2.addWidget(btn_ok)
        hb2.addWidget(btn_cancel)
        lay.addLayout(hb2)

        self.list.itemSelectionChanged.connect(self._on_selection_changed)
        self._renumber()

    def _on_selection_changed(self):
        current = {i.row() for i in self.list.selectedIndexes()}
        prev = set(self.selection_order)
        for r in list(self.selection_order):
            if r not in current:
                self.selection_order.remove(r)
        self.selection_order.extend(sorted(current - prev))
        self._renumber()

    def _renumber(self):
        for r, lbl in enumerate(self._label_by_row):
            lbl.setText(self._base_text_by_row[r])
        for n, r in enumerate(self.selection_order, start=1):
            lbl = self._label_by_row[r]
            base = self._base_text_by_row[r]
            lbl.setText(f"{n:02d}. {base}")

    def get_selected(self) -> List[str]:
        return [self._base_text_by_row[r] for r in self.selection_order]


# =============================================================================
# 6) Приложения 1/4: drag&drop + Ctrl+V
# =============================================================================
_ALLOWED_FILE_EXT = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif", ".pdf"}


class DropListWidget(QListWidget):
    def __init__(self, parent=None, *, paste_dir: Path | None = None, paste_prefix: str = "paste"):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setSelectionMode(QAbstractItemView.ExtendedSelection)

        base = Path(paste_dir) if paste_dir else Path(tempfile.gettempdir()) / "doc_generator_app" / "pasted_images"
        base.mkdir(parents=True, exist_ok=True)
        self._paste_dir = base
        self._paste_prefix = paste_prefix

    def dragEnterEvent(self, e: QDragEnterEvent):
        if e.mimeData().hasUrls():
            e.acceptProposedAction()
        else:
            super().dragEnterEvent(e)

    def dragMoveEvent(self, e):
        if e.mimeData().hasUrls():
            e.acceptProposedAction()
        else:
            super().dragMoveEvent(e)

    def dropEvent(self, e: QDropEvent):
        if e.mimeData().hasUrls():
            paths: list[str] = []
            for url in e.mimeData().urls():
                p = url.toLocalFile()
                if not p:
                    continue
                ext = Path(p).suffix.lower()
                if ext in _ALLOWED_FILE_EXT and Path(p).exists():
                    paths.append(str(Path(p).resolve()))
            self._add_paths(paths)
            e.acceptProposedAction()
            return
        super().dropEvent(e)

    def keyPressEvent(self, e):
        if e.matches(QKeySequence.Paste):
            if self.paste_from_clipboard():
                return
        super().keyPressEvent(e)

    def paste_from_clipboard(self) -> bool:
        cb = QGuiApplication.clipboard()
        md = cb.mimeData()

        # 1) картинка (скриншот)
        if md.hasImage():
            img = cb.image()
            if img.isNull():
                return False

            fname = f"{self._paste_prefix}_{datetime.now():%Y%m%d_%H%M%S}_{uuid.uuid4().hex[:8]}.png"
            out = (self._paste_dir / fname).resolve()
            ok = img.save(str(out), "PNG")
            if not ok:
                return False

            self._add_paths([str(out)])
            return True

        # 2) файлы из проводника
        if md.hasUrls():
            paths: list[str] = []
            for url in md.urls():
                p = url.toLocalFile()
                if not p:
                    continue
                pp = Path(p)
                if pp.exists() and pp.suffix.lower() in _ALLOWED_FILE_EXT:
                    paths.append(str(pp.resolve()))
            if paths:
                self._add_paths(paths)
                return True

        return False

    def _add_paths(self, paths: list[str]):
        existing = set(self.get_paths())
        for p in paths:
            if p in existing:
                continue
            self.addItem(QListWidgetItem(p))
            existing.add(p)

    def get_paths(self) -> list[str]:
        return [self.item(i).text() for i in range(self.count())]


class AppendixImagesDialog(QDialog):
    def __init__(self, parent=None, title: str = "Загрузка изображений/скриншотов",
                 initial_paths: list[str] | None = None,
                 paste_prefix: str = "app"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(760, 480)

        self.list = DropListWidget(self, paste_prefix=paste_prefix)
        if initial_paths:
            self.list._add_paths(initial_paths)

        # Ctrl+V на всё окно
        self._sc_paste = QShortcut(QKeySequence.Paste, self)
        self._sc_paste.activated.connect(self._on_paste)

        hint = QLabel(
            "Перетащите файлы сюда или нажмите «Добавить…».\n"
            "Также можно вставить скриншот: Ctrl+V.\n"
            "Поддержка: PNG/JPG/BMP/TIFF/GIF/PDF.\n"
            "Важно: PDF как картинку docxtpl вставлять не умеет — PDF будут пропущены."
        )
        hint.setWordWrap(True)

        btn_add = QPushButton("Добавить…")
        btn_rm = QPushButton("Удалить выбранные")
        btn_clear = QPushButton("Очистить")

        btn_add.clicked.connect(self._add_files)
        btn_rm.clicked.connect(self._remove_selected)
        btn_clear.clicked.connect(self.list.clear)

        hb = QHBoxLayout()
        hb.addWidget(btn_add)
        hb.addWidget(btn_rm)
        hb.addWidget(btn_clear)
        hb.addStretch()

        btn_ok = QPushButton("ОК")
        btn_cancel = QPushButton("Отмена")
        btn_ok.clicked.connect(self.accept)
        btn_cancel.clicked.connect(self.reject)

        hb2 = QHBoxLayout()
        hb2.addStretch()
        hb2.addWidget(btn_ok)
        hb2.addWidget(btn_cancel)

        lay = QVBoxLayout(self)
        lay.addWidget(hint)
        lay.addWidget(self.list)
        lay.addLayout(hb)
        lay.addLayout(hb2)

    def _on_paste(self):
        if not self.list.paste_from_clipboard():
            QApplication.beep()

    def _add_files(self):
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "Выберите файлы",
            os.getcwd(),
            "Images/PDF (*.png *.jpg *.jpeg *.bmp *.tif *.tiff *.gif *.pdf)"
        )
        if files:
            self.list._add_paths([str(Path(f).resolve()) for f in files])

    def _remove_selected(self):
        for it in list(self.list.selectedItems()):
            row = self.list.row(it)
            self.list.takeItem(row)

    def get_paths(self) -> list[str]:
        return self.list.get_paths()


# =============================================================================
# 7) Автозаполнение "Тест 11.2. Проверка кратности воздухообмена в ЧП"
# =============================================================================
def _fill_test_112(doc: DocxDocument, rooms: List[Dict[str, Any]]):
    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip().lower()

    t112 = None
    for t in doc.tables:
        if not t.rows:
            continue
        head = " ".join(c.text for c in t.rows[0].cells)
        h = norm(head)
        if ("тест" in h and ("11.2" in h or "11,2" in h)) and ("кратност" in h or "чп" in h):
            t112 = t
            break
    if t112 is None:
        return

    hdr_cols_row_idx = None
    for i, row in enumerate(t112.rows):
        line = norm(" ".join(c.text for c in row.cells))
        if "результаты испытани" in line:
            hdr_cols_row_idx = i + 1 if (i + 1) < len(t112.rows) else None
            break
    if hdr_cols_row_idx is None:
        for i, row in enumerate(t112.rows):
            line = norm(" ".join(c.text for c in row.cells))
            if ("номер" in line and "объем" in line) or ("общий расход" in line):
                hdr_cols_row_idx = i
                break
    if hdr_cols_row_idx is None:
        return

    def find_col(cells, *needles):
        for ci, c in enumerate(cells):
            tt = norm(c.text)
            if all(n in tt for n in needles):
                return ci
        return None

    cols_row = t112.rows[hdr_cols_row_idx]
    col_total = find_col(cols_row.cells, "общий", "расход")
    col_fact = find_col(cols_row.cells, "фактическ")
    if col_total is None and col_fact is None:
        return

    def is_spacer(row):
        return all(norm(c.text) in ("", "¤") for c in row.cells)

    data_start = hdr_cols_row_idx + 2  # пропускаем вторую строку заголовка
    while data_start < len(t112.rows) and is_spacer(t112.rows[data_start]):
        data_start += 1

    data_end = len(t112.rows)
    for i in range(data_start, len(t112.rows)):
        first_cell = norm(t112.rows[i].cells[0].text)
        if first_cell.startswith("комментар"):
            data_end = i
            break
    if data_end <= data_start:
        return

    need = len(rooms)
    have = data_end - data_start
    if need > have:
        sample_tr = deepcopy(t112.rows[data_end - 1]._tr)
        for _ in range(need - have):
            t112._tbl.append(deepcopy(sample_tr))
        data_end = data_start + need

    for idx, room in enumerate(rooms):
        r = data_start + idx
        if r >= len(t112.rows):
            break
        row = t112.rows[r]
        total_flow = (room.get("total_flow") or "").strip()
        exch_act = (room.get("exchange_actual") or "").strip()

        if col_total is not None and col_total < len(row.cells) and total_flow:
            row.cells[col_total].text = total_flow
        if col_fact is not None and col_fact < len(row.cells) and exch_act:
            row.cells[col_fact].text = exch_act


# =============================================================================
# 8) Форматирование дат поверки в таблице оборудования
# =============================================================================
def _fix_weird_ddmmyyyy(s: str) -> str:
    # "26.062026" -> "26.06.2026"
    m = re.match(r"^(\d{1,2})\.(\d{2})(\d{4})$", s)
    if m:
        return f"{m.group(1)}.{m.group(2)}.{m.group(3)}"
    return s


def _parse_date_any(s: str):
    s = (s or "").strip().replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s).replace("г.", "").replace("г", "").strip(" .")
    if not s:
        return None

    s = s.split()[0]  # убрать время
    s = _fix_weird_ddmmyyyy(s)

    fmts = ["%Y-%m-%d", "%d.%m.%Y", "%d-%m-%Y", "%Y.%m.%d", "%d/%m/%Y", "%Y/%m/%d"]
    for fmt in fmts:
        try:
            return datetime.strptime(s, fmt).date()
        except Exception:
            pass

    m = re.match(r"^(\d{4})-(\d{1,2})-(\d{1,2})$", s)
    if m:
        try:
            return date(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except Exception:
            return None

    m = re.match(r"^(\d{1,2})\.(\d{1,2})\.(\d{4})$", s)
    if m:
        try:
            return date(int(m.group(3)), int(m.group(2)), int(m.group(1)))
        except Exception:
            return None

    return None


def _fmt_date(d: date) -> str:
    return d.strftime("%d.%m.%Y")


def _format_date_range_cell(text: str) -> str:
    if not text or not text.strip():
        return text

    raw = re.sub(r"\s*\n\s*", " ", text).strip()

    if "/" in raw:
        parts = [p.strip() for p in raw.split("/") if p.strip()]
        if len(parts) >= 2:
            d1 = _parse_date_any(parts[0])
            d2 = _parse_date_any(parts[1])
            if d1 and d2:
                return f"{_fmt_date(d1)} / {_fmt_date(d2)}"
            return raw

    date_like = re.findall(r"\d{4}-\d{1,2}-\d{1,2}|\d{1,2}\.\d{1,2}\.\d{4}|\d{1,2}\.\d{2}\d{4}", raw)
    if len(date_like) >= 2:
        d1 = _parse_date_any(date_like[0])
        d2 = _parse_date_any(date_like[1])
        if d1 and d2:
            return f"{_fmt_date(d1)} / {_fmt_date(d2)}"

    d = _parse_date_any(raw)
    if d:
        return _fmt_date(d)

    return raw


def postprocess_equipment_dates(doc: DocxDocument) -> None:
    """
    Ищет в документе столбец 'Дата поверки/ Действительно до:' и
    форматирует значения в 'ДД.ММ.ГГГГ / ДД.ММ.ГГГГ'.
    """
    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip().lower()

    hdr_needle = "дата поверки/ действительно до"
    for t in doc.tables:
        if not t.rows:
            continue

        hdr_row_idx = None
        for ri in range(min(3, len(t.rows))):
            row_text = " ".join(c.text for c in t.rows[ri].cells)
            if hdr_needle in norm(row_text):
                hdr_row_idx = ri
                break
        if hdr_row_idx is None:
            continue

        col_idx = None
        for ci, c in enumerate(t.rows[hdr_row_idx].cells):
            if hdr_needle in norm(c.text):
                col_idx = ci
                break
        if col_idx is None:
            continue

        for ri in range(hdr_row_idx + 1, len(t.rows)):
            cell = t.rows[ri].cells[col_idx]
            old = cell.text
            new = _format_date_range_cell(old)
            if new != old:
                cell.text = new


import time
import re
import win32com.client as win32
from logger import logger



import re
from typing import Optional

import win32com.client as win32
from logger import logger


import re
from copy import deepcopy
from typing import Optional

from docx.document import Document as DocxDocument
from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _norm_basic(s: str) -> str:
    s = "" if s is None else str(s)
    s = (
        s.replace("\u00A0", " ")
         .replace("\u202F", " ")
         .replace("\u200B", "")
         .replace("\u00AD", "")
    )
    s = re.sub(r"\s+", " ", s).strip().lower()
    s = s.replace("ё", "е")
    return s


def _match_key(s: str) -> str:
    return re.sub(r"[^0-9a-zа-я]+", "", _norm_basic(s))


def _looks_like_table5(tbl: Table) -> bool:
    """
    Узнаём Таблицу 5 по шапке (стабильно для ваших шаблонов):
      - в первых 2 строках должны встречаться "Риск", "Возможная причина"
      - и "Аттестационное испытание"
    """
    if not tbl.rows:
        return False

    head0 = " ".join(_norm_basic(c.text) for c in tbl.rows[0].cells)
    head1 = " ".join(_norm_basic(c.text) for c in tbl.rows[1].cells) if len(tbl.rows) > 1 else ""

    return (
        ("риск" in head0 or "риск" in head1)
        and ("возможная причина" in head0 or "возможная причина" in head1)
        and ("аттестационное испытание" in head0 or "аттестационное испытание" in head1)
    )


def _insert_page_break_paragraph_after_table(tbl: Table) -> OxmlElement:
    """
    Вставляет ПУСТОЙ абзац с разрывом страницы СРАЗУ после таблицы.
    """
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    tbl._tbl.addnext(p)
    return p


def _insert_bold_tnr_paragraph_after(el: OxmlElement, text: str) -> OxmlElement:
    """
    Вставляет абзац сразу после el:
      - Times New Roman
      - 12pt
      - bold
      - текст: text
    """
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    for a in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{a}"), "Times New Roman")
    rPr.append(rFonts)

    b = OxmlElement("w:b")
    rPr.append(b)

    # 12pt = 24 half-points
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)

    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(szCs)

    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)

    p.append(r)
    el.addnext(p)
    return p


def split_table5_with_continuation_open_doc(
    doc: DocxDocument,
    *,
    split_risk_phrase: str = "Класс чистоты ЧП не соответствует установленным требованиям",
    continuation_title: str = "Продолжение таблицы 5",
    header_rows_count: int = 2,
) -> bool:
    """
    Делает разделение Таблицы 5 на 2 таблицы, как в "правильном шаблоне":

    - Находит Таблицу 5 по шапке (Риск / Возможная причина / Аттестационное испытание).
    - Находит строку, где в 1-й колонке встречается split_risk_phrase.
    - Всё начиная с этой строки переносит в НОВУЮ таблицу (копией XML),
      при этом шапка (header_rows_count) остаётся и в первой, и во второй таблице.
    - Между таблицами вставляет:
        1) разрыв страницы
        2) абзац "Продолжение таблицы 5" (TNR 12 bold)
    - Ничего не "перерисовывает" — копирует XML => формат ЕДИНЫЙ 1:1.

    Возвращает True если разделение сделано, иначе False.
    """

    # 1) Найти Таблицу 5
    t5: Optional[Table] = None
    for t in doc.tables:
        if _looks_like_table5(t):
            t5 = t
            break
    if t5 is None:
        return False

    # 2) Если уже есть "Продолжение таблицы 5" где-то в документе — не плодим дубль
    if any(continuation_title in (p.text or "") for p in doc.paragraphs):
        return False

    if len(t5.rows) <= header_rows_count:
        return False

    # 3) Найти row-индекс (внутри t5), с которого нужно переносить во 2-ю таблицу
    target_key = _match_key(split_risk_phrase)
    split_row_idx = None
    for ri in range(header_rows_count, len(t5.rows)):
        c0 = t5.rows[ri].cells[0].text
        if target_key and target_key in _match_key(c0):
            split_row_idx = ri
            break

    # Если фразы нет — НЕ делим (как ты и просишь: делим именно там)
    if split_row_idx is None:
        return False

    # Если попали прямо в начало данных — тоже смысла делить нет
    if split_row_idx <= header_rows_count:
        return False

    # 4) Достаём XML-строки
    tbl_el = t5._tbl
    tr_elems = tbl_el.xpath("./w:tr")
    if len(tr_elems) != len(t5.rows):
        # крайне редкая ситуация, но если вдруг — не рискуем
        return False

    header_trs = tr_elems[:header_rows_count]
    data_trs = tr_elems[header_rows_count:]

    split_data_pos = split_row_idx - header_rows_count
    if not (0 < split_data_pos < len(data_trs)):
        return False

    # 5) Создаём новую таблицу как полную копию, затем чистим строки
    new_tbl_el = deepcopy(tbl_el)
    new_tr_elems = new_tbl_el.xpath("./w:tr")

    # В новой таблице оставляем: header + data[split:]
    keep_idxs_new = set(range(0, header_rows_count)) | set(
        range(header_rows_count + split_data_pos, header_rows_count + len(data_trs))
    )
    for i in range(len(new_tr_elems) - 1, -1, -1):
        if i not in keep_idxs_new:
            new_tbl_el.remove(new_tr_elems[i])

    # В старой (первой) таблице оставляем: header + data[:split]
    for i in range(len(tr_elems) - 1, -1, -1):
        if i >= header_rows_count + split_data_pos:
            tbl_el.remove(tr_elems[i])

    # 6) Вставляем: разрыв страницы + "Продолжение..." + новую таблицу
    p_break = _insert_page_break_paragraph_after_table(t5)
    p_title = _insert_bold_tnr_paragraph_after(p_break, continuation_title)
    p_title.addnext(new_tbl_el)

    return True




def update_fields_with_word(docx_path: str) -> None:
    """
    Открывает DOCX в Word, режет Таблицу 5, обновляет поля/колонтитулы, сохраняет.
    """
    word = win32.gencache.EnsureDispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(docx_path))

        doc.Repaginate()
        split_table5_with_continuation_open_doc(doc)
        doc.Repaginate()

        doc.Fields.Update()
        for sec in doc.Sections:
            sec.Headers(1).Range.Fields.Update()
            sec.Footers(1).Range.Fields.Update()

        doc.Save()
        doc.Close()
    finally:
        word.Quit()

def _set_cell_text_keep_style(cell, text: str) -> None:
    """
    Ставит текст, стараясь не ломать форматирование ячейки:
    - если есть runs, пишем в первый run и чистим остальные
    - иначе fallback на cell.text
    """
    text = "" if text is None else str(text)

    if cell.paragraphs:
        p0 = cell.paragraphs[0]
        if p0.runs:
            p0.runs[0].text = text
            for r in p0.runs[1:]:
                r.text = ""
            # чистим остальные параграфы в ячейке
            for p in cell.paragraphs[1:]:
                for r in p.runs:
                    r.text = ""
            return

    cell.text = text


from copy import deepcopy
import re
from pathlib import Path
import openpyxl
from docx.table import Table
from docx.document import Document as DocxDocument


def _norm_key(s: str) -> str:
    s = "" if s is None else str(s)
    s = re.sub(r"(?i)^\s*тест\s*\d+(?:[.,]\d+)?\.?\s*", "", s).strip()
    s = (
        s.replace("\u00A0", " ")
         .replace("\u202F", " ")
         .replace("\u200B", "")
         .replace("\u00AD", "")
    )
    s = re.sub(r"\s+", " ", s).strip().lower()
    s = s.replace("ё", "е")
    # ключ без знаков
    return re.sub(r"[^0-9a-zа-я]+", "", s)


def _set_cell_text_keep_style(cell, text: str) -> None:
    text = "" if text is None else str(text)

    if cell.paragraphs:
        p0 = cell.paragraphs[0]
        if p0.runs:
            p0.runs[0].text = text
            for r in p0.runs[1:]:
                r.text = ""
            for p in cell.paragraphs[1:]:
                for r in p.runs:
                    r.text = ""
            return

    cell.text = text


def _load_table2_rows_from_excel(xlsx_path: str, sheet_name: str | None = None) -> dict[str, dict]:
    """
    Возвращает словарь:
      norm(test_name) -> {"test":..., "crit":..., "fact":..., "eval":...}
    """
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb[wb.sheetnames[0]]

    # ожидаем шапку в первой строке
    # A: тест, B: критерий, C: факт, D: оценка
    out: dict[str, dict] = {}
    for r in range(2, ws.max_row + 1):
        test = ws.cell(r, 1).value
        crit = ws.cell(r, 2).value
        fact = ws.cell(r, 3).value
        evl  = ws.cell(r, 4).value

        if not test:
            continue

        row = {
            "test": str(test).strip(),
            "crit": "" if crit is None else str(crit).strip(),
            "fact": "" if fact is None else str(fact).strip(),
            "eval": "" if evl  is None else str(evl).strip(),
        }
        out[_norm_key(row["test"])] = row

    return out


def fill_report_table2_from_excel(
    doc: DocxDocument,
    selected_tests: list[str],
    xlsx_path: str,
    *,
    sheet_name: str | None = None,
    default_eval: str = "Соответствует",
) -> tuple[bool, list[str]]:
    """
    Заполняет Таблицу 2 в ОТЧ-OQ данными из Excel по выбранным тестам.
    Возвращает: (ok, missing_tests)
    """

    # 1) грузим Excel
    rows_map = _load_table2_rows_from_excel(xlsx_path, sheet_name=sheet_name)

    # 2) находим таблицу 2 (по заголовку столбцов)
    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip().lower()

    target_table: Table | None = None
    for t in doc.tables:
        if not t.rows:
            continue
        head = " ".join(c.text for c in t.rows[0].cells)
        h = norm(head)
        if ("тест" in h) and ("критер" in h) and ("фактичес" in h) and ("оценк" in h):
            target_table = t
            break

    if target_table is None:
        return False, selected_tests[:]  # не нашли таблицу

    # 3) ищем шаблонную строку с маркерами (лучший вариант)
    token_row_idx = None
    tokens = {"#T2_TEST", "#T2_CRIT", "#T2_FACT", "#T2_EVAL"}
    for ri, row in enumerate(target_table.rows):
        row_tokens = { (c.text or "").strip() for c in row.cells }
        if tokens.issubset(row_tokens):
            token_row_idx = ri
            break

    # если шаблонной строки нет:
    # - если есть хотя бы 2 строки, возьмем 2ю как базовую
    # - иначе создадим строку через add_row (стиль может быть не идеальный)
    if token_row_idx is None:
        if len(target_table.rows) >= 2:
            token_row_idx = 1
        else:
            target_table.add_row()
            token_row_idx = 1

    base_tr = deepcopy(target_table.rows[token_row_idx]._tr)

    # 4) готовим данные в порядке выбора пользователя
    missing: list[str] = []
    chosen_rows: list[dict] = []
    for tname in selected_tests:
        key = _norm_key(tname)
        row = rows_map.get(key)
        if not row:
            missing.append(tname)
            # можно пропускать, либо вставлять пустую строку — выберу пропуск
            continue
        chosen_rows.append(row)

    # если ничего не нашли — хотя бы очистим шаблонную строку
    need = max(1, len(chosen_rows))

    # сколько строк данных сейчас начиная с token_row_idx
    have = len(target_table.rows) - token_row_idx
    if need > have:
        for _ in range(need - have):
            target_table._tbl.append(deepcopy(base_tr))

    # 5) заполняем строки
    for i in range(need):
        row = target_table.rows[token_row_idx + i]
        data = chosen_rows[i] if i < len(chosen_rows) else {"test": "", "crit": "", "fact": "", "eval": ""}

        for cell in row.cells:
            token = (cell.text or "").strip()

            if token == "#T2_TEST":
                _set_cell_text_keep_style(cell, data.get("test", ""))
            elif token == "#T2_CRIT":
                _set_cell_text_keep_style(cell, data.get("crit", ""))
            elif token == "#T2_FACT":
                _set_cell_text_keep_style(cell, data.get("fact", ""))
            elif token == "#T2_EVAL":
                ev = data.get("eval", "").strip() or default_eval
                _set_cell_text_keep_style(cell, ev)
            else:
                # если маркеров нет (fallback-режим), попробуем по колонкам
                # (тут можно оставить как есть)
                pass

    return True, missing


def fill_report_table1_rooms_by_hashes(doc: DocxDocument, rooms: list[dict]) -> bool:
    """
    Ищет таблицу, где строка данных содержит маркеры:
      #, ##, ###, ####, #$, #$$, #$$$, #%, #%%, #%%%
    и заполняет её из rooms (num, name, klass, area, volume, dp, airflow, exchange, temp, rh).
    """
    token_to_key = {
        "#": "num",
        "##": "name",
        "###": "klass",
        "####": "area",
        "#$": "volume",
        "#$$": "dp",
        "#$$$": "airflow",
        "#%": "exchange",
        "#%%": "temp",
        "#%%%": "rh",
    }

    target_table = None
    tmpl_row_idx = None

    # 1) найти таблицу и "шаблонную" строку с #..#%%%
    for t in doc.tables:
        for ri, row in enumerate(t.rows):
            row_tokens = [c.text.strip() for c in row.cells]
            if ("#" in row_tokens) and ("##" in row_tokens) and ("#%%%" in row_tokens):
                target_table = t
                tmpl_row_idx = ri
                break
        if target_table:
            break

    if not target_table or tmpl_row_idx is None:
        return False

    # 2) копируем шаблонную строку (чтобы формат 1:1 сохранялся)
    base_tr = deepcopy(target_table.rows[tmpl_row_idx]._tr)

    need = max(1, len(rooms))  # если комнат 0 — оставим одну строку пустой
    have = len(target_table.rows) - tmpl_row_idx

    if need > have:
        for _ in range(need - have):
            target_table._tbl.append(deepcopy(base_tr))

    # 3) заполняем
    for i in range(need):
        row = target_table.rows[tmpl_row_idx + i]
        room = rooms[i] if i < len(rooms) else {}

        for cell in row.cells:
            token = (cell.text or "").strip()
            key = token_to_key.get(token)
            if not key:
                continue
            _set_cell_text_keep_style(cell, (room.get(key) or "").strip())

    return True

def fix_table2_caption_glue(doc: DocxDocument, caption: str = "Таблица 2") -> bool:
    """
    1) Удаляет пустые абзацы между подписью и таблицей
    2) Делает таблицу inline (убирает tblpPr/tblOverlap)
    3) КЛЮЧЕВОЕ: ставит keepNext на абзац подписи, чтобы Word не оставлял подпись сиротой
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def norm(s: str) -> str:
        s = (s or "").replace("\u00A0", " ")
        s = re.sub(r"\s+", " ", s).strip().lower()
        return s

    cap_p = None
    for p in doc.paragraphs:
        if norm(p.text) == norm(caption):
            cap_p = p
            break
    if cap_p is None:
        return False

    # ✅ Главное: подпись не должна отрываться от таблицы
    pf = cap_p.paragraph_format
    pf.keep_with_next = True
    pf.keep_together = True
    pf.widow_control = True

    def norm(s: str) -> str:
        s = (s or "").replace("\u00A0", " ")
        s = re.sub(r"\s+", " ", s).strip().lower()
        return s

    # 0) найти абзац подписи
    cap_p = None
    for p in doc.paragraphs:
        if norm(p.text) == norm(caption):
            cap_p = p
            break
    if cap_p is None:
        return False

    # 1) keepNext на подпись (чтобы не отрывалась от таблицы)
    pPr = cap_p._p.get_or_add_pPr()
    if pPr.find(qn("w:keepNext")) is None:
        pPr.append(OxmlElement("w:keepNext"))

    # (опционально) чуть безопаснее: не разрывать саму подпись
    if pPr.find(qn("w:keepLines")) is None:
        pPr.append(OxmlElement("w:keepLines"))

    # 2) найти следующий элемент и удалить пустые абзацы между подписью и таблицей
    p_el = cap_p._p
    sib = p_el.getnext()

    while sib is not None and sib.tag != qn("w:tbl"):
        if sib.tag == qn("w:p"):
            # удаляем реально пустые (в т.ч. с одними пробелами)
            txt = "".join(t.text for t in sib.iter() if t.tag == qn("w:t"))
            if norm(txt) == "":
                nxt = sib.getnext()
                sib.getparent().remove(sib)
                sib = nxt
                continue
        sib = sib.getnext()

    if sib is None or sib.tag != qn("w:tbl"):
        return False

    # 3) убрать "плавающее" позиционирование таблицы
    tbl_el = sib
    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is not None:
        for child_tag in (qn("w:tblpPr"), qn("w:tblOverlap")):
            ch = tblPr.find(child_tag)
            if ch is not None:
                tblPr.remove(ch)

    return True




# =============================================================================
# 9) Render Worker
# =============================================================================

class RenderWorker(QThread):
    progress = Signal(int, str)
    finished = Signal(bool, str, list)

    def __init__(
            self,
            tpl_path: str,
            tests_doc_path: str,
            xls_tests_path: str,
            xls_eq_path: str,
            out_path: str,
            selected_tests: list[str],
            rooms: list[dict],
            equipment: list[dict],
            ctx_fields: dict,
            *,
            risk_path: str,
            scans_dir: str,
            app1_images: list[str],
            app4_images: list[str],
            app5_images: list[str],

            # NEW:
            tpl_report_path: str | None = None,
            out_report_path: str | None = None,
            ctx_fields_report: dict | None = None,
            xls_report_path: str | None = None,
    ):
        super().__init__()

        self.tpl_path = tpl_path
        self.xls_report_path = xls_report_path
        self.tests_doc_path = tests_doc_path
        self.xls_tests_path = xls_tests_path
        self.xls_eq_path = xls_eq_path
        self.out_path = out_path
        self.selected_tests = selected_tests
        self.rooms = rooms
        self.equipment = equipment
        self.ctx_fields = ctx_fields
        self.risk_path = risk_path
        self.scans_dir = scans_dir
        self.app1_images = app1_images or []
        self.app4_images = app4_images or []
        self.app5_images = app5_images or []

        # NEW:
        self.tpl_report_path = tpl_report_path
        self.out_report_path = out_report_path
        self.ctx_fields_report = ctx_fields_report

    def run(self):
        missing: list[str] = []
        try:
            step = 0
            self.progress.emit(step, "Валидация файлов…")
            io_manager.validate_file(Path(self.tpl_path))
            io_manager.validate_file(Path(self.tests_doc_path))
            io_manager.validate_file(Path(self.xls_tests_path))
            io_manager.validate_file(Path(self.xls_eq_path))
            io_manager.validate_file(Path(self.risk_path))
            Path(self.out_path).parent.mkdir(parents=True, exist_ok=True)

            # второй документ (ОТЧ-OQ) — опционально
            do_report = bool(self.tpl_report_path and self.out_report_path and self.ctx_fields_report is not None)
            if do_report:
                io_manager.validate_file(Path(self.tpl_report_path))
                Path(self.out_report_path).parent.mkdir(parents=True, exist_ok=True)

                if self.xls_report_path:
                    io_manager.validate_file(Path(self.xls_report_path))
                else:
                    logger.warning("ОТЧ-OQ: не указан Excel-файл с данными для Таблицы 2.")

            # ---------- 1. Базовый контекст ----------
            step += 1
            self.progress.emit(step, "Сбор контекста…")
            context = template_renderer.build_context(self.ctx_fields, self.rooms)

            # ---------- 2. Скан-файлы (поверки оборудования) ----------
            step += 1
            self.progress.emit(step, "Сбор сканов оборудования…")
            scan_paths: list[str] = []

            def _serial_key(s: str) -> str:
                return re.sub(r"[^0-9a-zA-Z]+", "", (s or "")).lower()

            if self.equipment and self.scans_dir:
                for eq in self.equipment:
                    name_sn = eq.get("name_sn", "")
                    try:
                        _, serial = [p.strip() for p in name_sn.rsplit(",", 1)]
                    except ValueError:
                        serial = name_sn.strip()

                    key = _serial_key(serial)
                    found = None
                    for root, _, files in os.walk(self.scans_dir):
                        for fn in files:
                            # pdf пропускаем как картинку
                            if Path(fn).suffix.lower() not in (".jpg", ".jpeg", ".png", ".pdf"):
                                continue
                            if _serial_key(Path(fn).stem) == key:
                                found = os.path.join(root, fn)
                                break
                        if found:
                            scan_paths.append(found)
                            break

            if scan_paths:
                context["Scan_paths"] = scan_paths

            # ---------- 3. Рендер тестового DOCX (1-й проход) ----------
            step += 1
            self.progress.emit(step, "Рендер тестового документа…")
            tpl_tests = DocxTemplate(self.tests_doc_path)
            tpl_tests.render(context)

            with temp_docx() as tmp_tests_path:
                tpl_tests.save(tmp_tests_path)

                # ---------- 4. Извлечь расход и посчитать кратность ----------
                tests_doc_parsed = Document(tmp_tests_path)

                def _robust_extract_total_flows_from_test11(docx_doc: DocxDocument):
                    flows: list[Decimal | None] = []

                    def norm(s: str) -> str:
                        return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip().lower()

                    def to_dec(s: str) -> Decimal | None:
                        s = (s or "").strip().replace("\u00a0", "").replace(" ", "").replace(",", ".")
                        try:
                            return Decimal(s)
                        except InvalidOperation:
                            return None

                    for t in docx_doc.tables:
                        head = " ".join(c.text for c in t.rows[0].cells) if t.rows else ""
                        h = norm(head)
                        if "проверка" in h and "расхода" in h and "приточ" in h:
                            hdr_idx = None
                            for i, row in enumerate(t.rows):
                                if "результаты испытани" in norm(" ".join(c.text for c in row.cells)):
                                    hdr_idx = i + 1 if (i + 1) < len(t.rows) else None
                                    break
                            if hdr_idx is None:
                                continue

                            rows_for_header = [t.rows[hdr_idx]]
                            if hdr_idx + 1 < len(t.rows):
                                rows_for_header.append(t.rows[hdr_idx + 1])

                            def col_idx(*needles):
                                for row in rows_for_header:
                                    for ci, c in enumerate(row.cells):
                                        tt = norm(c.text)
                                        if all(n in tt for n in needles):
                                            return ci
                                return None

                            col_fact = col_idx("фактическ")
                            col_sum = col_idx("фактическ", "суммарн")

                            cur_sum: Decimal | None = None
                            room_started = False

                            data_start = hdr_idx + (2 if len(rows_for_header) == 2 else 1)

                            for r in range(data_start, len(t.rows)):
                                row = t.rows[r]
                                row_text = norm(" ".join(c.text for c in row.cells))

                                if row_text.startswith("комментар"):
                                    if room_started:
                                        flows.append(cur_sum)
                                    break

                                if row_text.startswith("помещение"):
                                    if room_started:
                                        flows.append(cur_sum)
                                    room_started = True
                                    cur_sum = None
                                    continue

                                if not room_started:
                                    continue

                                if col_sum is not None and col_sum < len(row.cells):
                                    v = to_dec(row.cells[col_sum].text)
                                    if v is not None:
                                        cur_sum = v
                                        continue

                                if col_fact is not None and col_fact < len(row.cells):
                                    v = to_dec(row.cells[col_fact].text)
                                    if v is not None:
                                        cur_sum = (cur_sum or Decimal("0")) + v

                            if room_started:
                                flows.append(cur_sum)

                    return flows

                _total1 = table_processor.extract_total_flows_from_test11(tests_doc_parsed)
                total_flows = (
                    _total1 if (_total1 and any(v is not None for v in _total1))
                    else _robust_extract_total_flows_from_test11(tests_doc_parsed)
                )

                for idx, room in enumerate(self.rooms):
                    if idx < len(total_flows) and total_flows[idx] is not None:
                        val = total_flows[idx].quantize(Decimal("1.00"))
                        room["total_flow"] = f"{val}".replace(".", ",")
                    else:
                        pr = (room.get("airflow") or "").strip().replace(" ", "").replace("\u00a0", "")
                        pr = pr.replace(".", ",")
                        try:
                            q = Decimal(pr.replace(",", ".")).quantize(Decimal("1.00"))
                            room["total_flow"] = f"{q}".replace(".", ",")
                        except InvalidOperation:
                            room["total_flow"] = pr or ""

                    vol_raw = (room.get("volume") or "").replace(",", ".")
                    try:
                        vol = Decimal(vol_raw)
                        tf_raw = (room.get("total_flow") or "").replace(",", ".")
                        tf = Decimal(tf_raw) if tf_raw else None
                        if vol and vol > 0 and tf is not None:
                            exch = tf / vol
                            room["exchange_actual"] = str(exch.quantize(Decimal("1.00"))).replace(".", ",")
                        else:
                            room["exchange_actual"] = ""
                    except (InvalidOperation, ZeroDivisionError):
                        room["exchange_actual"] = ""

                # ---------- 5. Пересобрать контекст ----------
                context = template_renderer.build_context(self.ctx_fields, self.rooms)
                if scan_paths:
                    context["Scan_paths"] = scan_paths

                # ---------- 6. Перерендер тестового DOCX (2-й проход) ----------
                step += 1
                self.progress.emit(step, "Перерендер тестовых таблиц с расчётами…")
                tpl_tests2 = DocxTemplate(self.tests_doc_path)
                tpl_tests2.render(context)
                tpl_tests2.save(tmp_tests_path)

                # ---------- 7. Рендер основного шаблона ----------
                step += 1
                self.progress.emit(step, "Рендер основного шаблона…")
                tpl_main = DocxTemplate(self.tpl_path)

                context["App1Scans"] = make_inline_images(tpl_main, self.app1_images, label="Приложение 1")
                context["App4Scans"] = make_inline_images(tpl_main, self.app4_images, label="Приложение 4")
                context["App5Scans"] = make_inline_images(tpl_main, self.app5_images, label="Приложение 5")
                context["Scans"] = make_inline_images(tpl_main, scan_paths, label="Приложение 2")

                tpl_main.render(context)

                with temp_docx() as tmp_main_path:
                    tpl_main.save(tmp_main_path)
                    doc = Document(tmp_main_path)

                    # ---------- 8. Таблицы помещений/оборудования ----------
                    step += 1
                    self.progress.emit(step, "Обработка таблицы помещений…")
                    table_processor.process_rooms_table(doc, self.rooms)

                    step += 1
                    self.progress.emit(step, "Обработка таблицы оборудования…")
                    table_processor.process_equipment_table(doc, self.equipment)
                    postprocess_equipment_dates(doc)

                    # ---------- 9. Вставка выбранных тестов ----------
                    step += 1
                    self.progress.emit(step, "Вставка тестовых таблиц…")
                    missing = table_processor.insert_test_tables(doc, tmp_tests_path, self.selected_tests)

                    # ---------- 10. Таблица 5 ----------
                    step += 1
                    self.progress.emit(step, "Вставка Таблицы 5 (анализ рисков)…")
                    try:
                        risk_rows = get_risk_rows(self.risk_path, self.selected_tests)
                        insert_table5_into_doc(doc, risk_rows)
                    except Exception as e:
                        logger.warning(f"Таблица 5 не вставлена: {e}")

                    # ---------- 11. Постобработка ----------
                    step += 1
                    self.progress.emit(step, "Заполнение результатов тестов…")
                    table_processor.process_test_results_tables(doc, self.rooms)
                    _fill_test_112(doc, self.rooms)

                    step += 1
                    self.progress.emit(step, "Унификация шрифта…")
                    table_processor.enforce_tnr_face_only_everywhere(doc)

                    try:
                        for table in doc.tables:
                            table.style = "Table Grid"
                    except Exception as e:
                        logger.warning(f"Не удалось применить стиль 'Table Grid': {e}")

                    # ---------- 12. Сохранение ----------
                    step += 1
                    self.progress.emit(step, "Сохранение…")
                    doc.save(self.out_path)

                try:
                    word_table5_splitter.update_fields_with_word(self.out_path)
                except Exception as e:
                    logger.warning(f"Не удалось обновить поля/разрезать таблицу 5 через Word: {e}")

                # ---------- 13. Рендер ОТЧ-OQ ----------
                if do_report:
                    step += 1
                    self.progress.emit(step, "Рендер ОТЧ-OQ…")

                    context_r = template_renderer.build_context(self.ctx_fields_report, self.rooms)

                    if scan_paths:
                        context_r["Scan_paths"] = scan_paths

                    tpl_r = DocxTemplate(self.tpl_report_path)

                    context_r["App1Scans"] = make_inline_images(tpl_r, self.app1_images, label="Приложение 1")
                    context_r["App4Scans"] = make_inline_images(tpl_r, self.app4_images, label="Приложение 4")
                    context_r["App5Scans"] = make_inline_images(tpl_r, self.app5_images, label="Приложение 5")
                    context_r["Scans"] = make_inline_images(tpl_r, scan_paths, label="Приложение 2")

                    tpl_r.render(context_r)

                    with temp_docx() as tmp_r_path:
                        tpl_r.save(tmp_r_path)
                        doc_r = Document(tmp_r_path)

                        # 1) Таблица 1 (помещения)
                        ok1 = fill_report_table1_rooms_by_hashes(doc_r, self.rooms)
                        if not ok1:
                            logger.warning(
                                "ОТЧ-OQ: Таблица 1 с маркерами #/##/... не найдена — помещения не заполнены.")

                        # 2) Таблица 2 (из Excel по выбранным тестам)
                        if self.xls_report_path:
                            ok2, missing2 = fill_report_table2_from_excel(
                                doc_r,
                                self.selected_tests,  # порядок как выбран пользователем
                                self.xls_report_path,
                                default_eval="Соответствует",
                            )
                            if not ok2:
                                logger.warning("ОТЧ-OQ: Таблица 2 не найдена (по заголовку/маркерам).")
                            if missing2:
                                logger.warning("ОТЧ-OQ: в Excel нет строк для тестов:\n" + "\n".join(missing2))
                        else:
                            logger.warning("ОТЧ-OQ: Excel отчёта ОТЧ-OQ не задан — Таблица 2 не заполнена.")

                        # 3) Исправление разрыва: "Таблица 2" приклеить к следующей таблице
                        try:
                            fix_table2_caption_glue(doc_r)  # <-- ВАЖНО: функция должна быть определена
                        except Exception as e:
                            logger.warning(f"ОТЧ-OQ: не удалось применить fix_table2_caption_glue: {e}")

                        # 4) Сохранение один раз
                        doc_r.save(self.out_report_path)

                    # если надо обновлять поля/колонтитулы через Word — можно оставить
                    # Таблица 2: деление по странице + "Продолжение таблицы 2" + обновление полей
                    try:
                        word_table2_otch_splitter.update_fields_and_split_table2(self.out_report_path)
                    except Exception as e:
                        logger.warning(f"ОТЧ-OQ: не удалось разрезать Таблицу 2 / обновить поля через Word: {e}")

                    # ✅ финальный проход: обновить ВСЁ (PAGE/NUMPAGES + TOC + поля)
                    try:
                        import word_update_all
                        word_update_all.update_all(self.out_report_path)
                    except Exception as e:
                        logger.warning(f"ОТЧ-OQ: не удалось обновить все поля/содержание через Word: {e}")

            msg = f"Документ сохранён:\n{self.out_path}"
            if do_report and self.out_report_path:
                msg += f"\n\nОТЧ-OQ сохранён:\n{self.out_report_path}"
            self.finished.emit(True, msg, missing)

        except Exception as e:
            logger.exception("Ошибка в RenderWorker:")
            self.finished.emit(False, str(e), [])


# =============================================================================
# 10) Main Window
# =============================================================================
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Генератор OQ/PQ — PySide6 + DocxTemplate")

        base = Path(__file__).resolve().parents[1]
        self.defaults = {
            "OQ": {
                "tpl": base / "Шаблон OQ.docx",
                "tests": base / "Тесты OQ.docx",
                "xls_tests": base / "tests OQ.xlsx",
                "xls_eq": base / "ПЕРЕЧЕНЬ ПРИБОРОВ OQ.xlsx",
                "scans_dir": base / "Сканы поверок",
                "risk_doc": base / "risk_analysis_from_docx.xlsx",
            },
            "PQ": {
                "tpl": base / "Шаблон PQ.docx",
                "tests": base / "Тесты PQ.docx",
                "xls_tests": base / "tests PQ.xlsx",
                "xls_eq": base / "ПЕРЕЧЕНЬ ПРИБОРОВ OQ.xlsx",
                "scans_dir": base / "Сканы поверок",
                "risk_doc": base / "risk_analysis_from_docx.xlsx",
            },
        }

        self.rooms: List[Dict[str, str]] = []
        self.equipment: List[Dict[str, str]] = []
        self.all_tests: List[str] = []
        self.selected_tests: List[str] = []
        self._batch_modes: List[str] = []

        # Приложения 1, 4, 5
        self.app1_images: list[str] = []
        self.app4_images: list[str] = []
        self.app5_images: list[str] = []

        self._build_ui()
        self.on_mode_changed("OQ")

    def _build_ui(self) -> None:
        cw = QWidget()
        g = QGridLayout()
        g.setSpacing(10)
        row = 0

        g.addWidget(QLabel("Режим:"), row, 0)
        self.mode_combo = QComboBox()
        self.mode_combo.addItems(["OQ", "PQ", "OQ и PQ"])
        self.mode_combo.currentTextChanged.connect(self.on_mode_changed)
        g.addWidget(self.mode_combo, row, 1, 1, 2)
        row += 1

        g.addWidget(QLabel("Шаблон (.docx):"), row, 0)
        self.tpl_path = QLineEdit()
        self.tpl_path.setReadOnly(True)
        g.addWidget(self.tpl_path, row, 1, 1, 1)
        btn_browse_tpl = QPushButton("…")
        btn_browse_tpl.clicked.connect(lambda: self._select(self.tpl_path, "Word Documents (*.docx)", save=False))
        g.addWidget(btn_browse_tpl, row, 2)
        row += 1

        g.addWidget(QLabel("Документ-тесты (.docx):"), row, 0)
        self.tests_path = QLineEdit()
        g.addWidget(self.tests_path, row, 1, 1, 1)
        btn_browse_tests = QPushButton("…")
        btn_browse_tests.clicked.connect(lambda: self._select(self.tests_path, "Word Documents (*.docx)", save=False))
        g.addWidget(btn_browse_tests, row, 2)
        row += 1

        btn_calc11 = QPushButton("Калькулятор Тест 11 (расход)…")
        btn_calc11.clicked.connect(self.launch_test11_calculator)
        g.addWidget(btn_calc11, row, 0, 1, 3)
        row += 1

        g.addWidget(QLabel("Excel с тестами:"), row, 0)
        self.xls_path = QLineEdit()
        self.xls_path.setReadOnly(True)
        g.addWidget(self.xls_path, row, 1, 1, 1)
        btn_browse_xls = QPushButton("…")
        btn_browse_xls.clicked.connect(lambda: self._select(self.xls_path, "Excel Files (*.xlsx)", save=False))
        g.addWidget(btn_browse_xls, row, 2)
        row += 1

        self.tests_btn = QPushButton("Выбрать тесты…")
        self.tests_btn.clicked.connect(self.choose_tests)
        g.addWidget(self.tests_btn, row, 0, 1, 3)
        row += 1

        g.addWidget(QLabel("Excel-оборудование:"), row, 0)
        self.equipment_xls_path = QLineEdit()
        self.equipment_xls_path.setReadOnly(True)
        g.addWidget(self.equipment_xls_path, row, 1, 1, 1)
        btn_eq = QPushButton("Загрузить оборудование")
        btn_eq.clicked.connect(lambda: self.load_equipment(silent=False))
        g.addWidget(btn_eq, row, 2)
        row += 1

        btn_rooms = QPushButton("Настройки помещений…")
        btn_rooms.clicked.connect(self.edit_rooms)
        g.addWidget(btn_rooms, row, 0, 1, 3)
        row += 1

        g.addWidget(QLabel("Папка со сканами (поверки):"), row, 0)
        self.scans_dir_input = QLineEdit()
        g.addWidget(self.scans_dir_input, row, 1, 1, 1)
        btn_dir = QPushButton("…")
        btn_dir.clicked.connect(lambda: self._select_dir(self.scans_dir_input))
        g.addWidget(btn_dir, row, 2)
        row += 1

        # Приложение 1
        g.addWidget(QLabel("Приложение 1:"), row, 0, Qt.AlignTop)
        self.app1_btn = QPushButton("Выбрать / перетащить файлы…")
        self.app1_btn.clicked.connect(self.choose_app1_images)
        g.addWidget(self.app1_btn, row, 1, 1, 2)
        row += 1

        self.app1_info = QLabel("Файлы не выбраны")
        self.app1_info.setWordWrap(True)
        g.addWidget(self.app1_info, row, 1, 1, 2)
        row += 1

        # Приложение 4
        g.addWidget(QLabel("Приложение 4:"), row, 0, Qt.AlignTop)
        self.app4_btn = QPushButton("Выбрать / перетащить файлы…")
        self.app4_btn.clicked.connect(self.choose_app4_images)
        g.addWidget(self.app4_btn, row, 1, 1, 2)
        row += 1

        self.app4_info = QLabel("Файлы не выбраны")
        self.app4_info.setWordWrap(True)
        g.addWidget(self.app4_info, row, 1, 1, 2)
        row += 1

        # Приложение 5
        g.addWidget(QLabel("Приложение 5:"), row, 0, Qt.AlignTop)
        self.app5_btn = QPushButton("Выбрать / перетащить файлы…")
        self.app5_btn.clicked.connect(self.choose_app5_images)
        g.addWidget(self.app5_btn, row, 1, 1, 2)
        row += 1

        self.app5_info = QLabel("Файлы не выбраны")
        self.app5_info.setWordWrap(True)
        g.addWidget(self.app5_info, row, 1, 1, 2)
        row += 1

        fields = [
            ("Объект:", "object_input"), ("ПРТ:", "prt_input"),
            ("Год:", "year_input"), ("Заказчик:", "customer_input"),
            ("Адрес объекта:", "address_input"),
            ("Разработал:", "developed_input"), ("Проверил:", "checked_input"),
        ]
        for label, attr in fields:
            g.addWidget(QLabel(label), row, 0)
            le = QLineEdit()
            setattr(self, attr, le)
            g.addWidget(le, row, 1, 1, 2)
            row += 1

        dates = [
            ("Дата разработки:", "date_dev"),
            ("Дата проверки:", "date_check"),
            ("Дата начала испытания:", "date_test"),
            ("Дата окончания:", "date_end"),
        ]
        for label, attr in dates:
            g.addWidget(QLabel(label), row, 0)
            de = QDateEdit(calendarPopup=True)
            de.setDate(date.today())
            setattr(self, attr, de)
            g.addWidget(de, row, 1, 1, 2)
            row += 1

        g.addWidget(QLabel("Сохранить как (.docx):"), row, 0)
        self.out_path = QLineEdit()
        g.addWidget(self.out_path, row, 1, 1, 1)
        btn_save = QPushButton("…")
        btn_save.clicked.connect(lambda: self._select(self.out_path, "Word Documents (*.docx)", save=True))
        g.addWidget(btn_save, row, 2)
        row += 1

        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 12)
        g.addWidget(self.progress_bar, row, 0, 1, 3)
        row += 1

        btn_generate = QPushButton("Сгенерировать")
        btn_generate.clicked.connect(self.start_render)
        g.addWidget(btn_generate, row, 0, 1, 3)

        cw.setLayout(g)
        self.setCentralWidget(cw)

    def _select(self, widget: QLineEdit, flt: str, save: bool = False) -> None:
        if save:
            path, _ = QFileDialog.getSaveFileName(self, "Сохранить файл", os.getcwd(), flt)
        else:
            path, _ = QFileDialog.getOpenFileName(self, "Выберите файл", os.getcwd(), flt)
        if path:
            widget.setText(path)

    def _select_dir(self, widget: QLineEdit) -> None:
        d = QFileDialog.getExistingDirectory(self, "Выберите папку со сканами", os.getcwd())
        if d:
            widget.setText(d)

    def choose_app1_images(self):
        dlg = AppendixImagesDialog(
            self,
            title="Приложение 1 — загрузка изображений/скриншотов",
            initial_paths=self.app1_images,
            paste_prefix="app1"
        )
        if dlg.exec() == QDialog.Accepted:
            self.app1_images = dlg.get_paths()
            self.app1_info.setText(f"Выбрано файлов: {len(self.app1_images)}")

    def choose_app4_images(self):
        dlg = AppendixImagesDialog(
            self,
            title="Приложение 4 — загрузка изображений/скриншотов",
            initial_paths=self.app4_images,
            paste_prefix="app4"
        )
        if dlg.exec() == QDialog.Accepted:
            self.app4_images = dlg.get_paths()
            self.app4_info.setText(f"Выбрано файлов: {len(self.app4_images)}")

    def choose_app5_images(self):
        dlg = AppendixImagesDialog(
            self,
            title="Приложение 5 — загрузка изображений/скриншотов",
            initial_paths=self.app5_images,
            paste_prefix="app5"
        )
        if dlg.exec() == QDialog.Accepted:
            self.app5_images = dlg.get_paths()
            self.app5_info.setText(f"Выбрано файлов: {len(self.app5_images)}")

    def on_mode_changed(self, mode: str) -> None:
        if mode in ("OQ и PQ", "ОQ и PQ"):
            self.tpl_path.setText("")
            self.tests_path.setText("")
            self.xls_path.setText("")
            self.equipment_xls_path.setText("")
            self.scans_dir_input.setText("")
            return

        d = self.defaults[mode]
        self.tpl_path.setText(str(d["tpl"]))
        self.tests_path.setText(str(d["tests"]))
        self.xls_path.setText(str(d["xls_tests"]))
        self.equipment_xls_path.setText(str(d["xls_eq"]))
        self.scans_dir_input.setText(str(d["scans_dir"]))
        self.load_tests(silent=True)
        self.load_equipment(silent=True)

    def load_tests(self, silent: bool = False) -> None:
        try:
            path = Path(self.xls_path.text())
            self.all_tests = io_manager.load_tests_list(path)
            self.selected_tests = []
            if not silent:
                QMessageBox.information(self, "Тесты загружены", f"Всего тестов: {len(self.all_tests)}")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка загрузки тестов", str(e))

    def choose_tests(self) -> None:
        if not self.all_tests:
            QMessageBox.warning(self, "Нет тестов", "Тесты не загружены!")
            return

        dlg = TestsDialog(self, self.all_tests, self.selected_tests,
                          calc_launcher=self.launch_test11_calculator)
        if dlg.exec() == QDialog.Accepted:
            self.selected_tests = dlg.get_selected()

    def load_equipment(self, silent: bool = False) -> None:
        try:
            path = Path(self.equipment_xls_path.text())
            by_sheets = io_manager.load_equipment_by_sheets(path)
            if silent:
                self.equipment = [it for items in by_sheets.values() for it in items]
                return
            dlg = EquipmentDialog(self, by_sheets)
            if dlg.exec() == QDialog.Accepted:
                self.equipment = dlg.get_selected()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка загрузки оборудования", str(e))

    def edit_rooms(self) -> None:
        dlg = RoomsDialog(self, self.rooms)
        if dlg.exec() == QDialog.Accepted:
            self.rooms = dlg.get_rooms()

    def launch_test11_calculator(self, *_, auto_close: bool = True) -> int:
        import subprocess

        tpl_path_str = (self.tests_path.text().strip() or "")
        if not tpl_path_str:
            QMessageBox.critical(self, "Калькулятор Тест 11", "Не указан «Документ-тесты (.docx)».")
            return 1

        tpl_path = Path(tpl_path_str).expanduser().resolve()
        if not tpl_path.exists():
            QMessageBox.critical(self, "Калькулятор Тест 11", f"Файл не найден:\n{tpl_path}")
            return 1

        rooms: List[Dict[str, Any]] = list(self.rooms) if getattr(self, "rooms", None) else []
        n = max(1, len(rooms) or 1)

        def _get(room: Dict[str, Any], key: str, default: str) -> str:
            v = room.get(key)
            s = "" if v is None else str(v).strip()
            return s if s else default

        # забираем из "Настройки помещений": num, name, klass, area, airflow
        nums = [_get(r, "num", f"{i + 1}") for i, r in enumerate(rooms[:n])]
        names = [_get(r, "name", "") for r in rooms[:n]]
        klasses = [_get(r, "klass", "") for r in rooms[:n]]
        areas = [_get(r, "area", "1") for r in rooms[:n]]

        # этих полей в RoomsDialog нет -> остаются дефолты (если не добавлял их отдельно)
        filters = [_get(r, "filters", "2") for r in rooms[:n]]
        points = [_get(r, "points", "2") for r in rooms[:n]]

        airflows = [_get(r, "airflow", "") for r in rooms[:n]]

        def _pad(lst, filler):
            need = n - len(lst)
            return lst if need <= 0 else lst + [filler(len(lst) + k + 1) for k in range(need)]

        nums = _pad(nums, lambda k: str(k))
        names = _pad(names, lambda _: "")
        klasses = _pad(klasses, lambda _: "")
        areas = _pad(areas, lambda _: "1")
        filters = _pad(filters, lambda _: "2")
        points = _pad(points, lambda _: "2")
        airflows = _pad(airflows, lambda _: "")

        def _join(xs: List[str]) -> str:
            # безопасно для разделителя ;
            return ";".join(x.replace(";", " ").strip() for x in xs)

        nums_s = _join(nums)
        names_s = _join(names)
        klasses_s = _join(klasses)
        areas_s = _join(areas)
        filters_s = _join(filters)
        points_s = _join(points)
        airflows_s = _join(airflows)

        proj_root = Path(__file__).resolve().parents[1]
        candidates = [
            proj_root / "tools" / "test11_airflow_calc.py",
            Path.cwd() / "tools" / "test11_airflow_calc.py",
            proj_root / "tests" / "test11_airflow_calc.py",
            Path(__file__).resolve().parent / "test11_airflow_calc.py",
        ]
        script = next((p for p in candidates if p.exists()), None)
        if script is None:
            paths_list = "\n".join(str(p) for p in candidates)
            QMessageBox.critical(
                self, "Калькулятор Тест 11",
                "Не найден скрипт test11_airflow_calc.py по путям:\n\n" + paths_list
            )
            return 1

        tmp_dir = Path(tempfile.mkdtemp(prefix="test11_"))
        auto_save_path = tmp_dir / "test11_auto_save.docx"

        args: List[str] = [
            sys.executable, str(script),
            "--template", str(tpl_path),
            "--rooms", str(n),
            "--nums", nums_s,
            "--names", names_s,
            "--klasses", klasses_s,  # <-- ВАЖНО
            "--areas", areas_s,
            "--filters", filters_s,
            "--points", points_s,
            "--airflows", airflows_s,
            "--auto-save", str(auto_save_path),
        ]
        if auto_close:
            args.append("--auto-close")

        try:
            proc = subprocess.run(args, check=False)
            rc = int(proc.returncode or 0)
        except Exception as e:
            QMessageBox.critical(self, "Калькулятор Тест 11", f"Ошибка запуска:\n{e}")
            return 1

        if auto_save_path.exists():
            self.tests_path.setText(str(auto_save_path))

        setattr(self, "last_test11_autosave", str(auto_save_path))
        return rc

    def start_render(self) -> None:
        mode = self.mode_combo.currentText()
        if mode == "OQ и PQ":
            self._batch_modes = ["OQ", "PQ"]
            self._start_single_render(self._batch_modes.pop(0))
        else:
            self._batch_modes = []
            self._start_single_render(mode)

    def _start_single_render(self, mode: str) -> None:
        mode = mode.upper()
        use_ui_paths = (self.mode_combo.currentText().upper() == mode)

        if use_ui_paths:
            tpl_path = self.tpl_path.text().strip()
            tests_path = self.tests_path.text().strip()
            xls_tests = self.xls_path.text().strip()
            xls_eq = self.equipment_xls_path.text().strip()
            scans_dir = self.scans_dir_input.text().strip()
            sel_tests_raw = self.selected_tests[:]
            equipment_list = self.equipment[:]
            # risk берем из defaults (UI поля под него нет)
            risk_path = str(self.defaults[mode]["risk_doc"])
        else:
            d = self.defaults[mode]
            tpl_path = str(d["tpl"])
            tests_path = str(d["tests"])
            xls_tests = str(d["xls_tests"])
            xls_eq = str(d["xls_eq"])
            scans_dir = str(d["scans_dir"])
            risk_path = str(d["risk_doc"])
            try:
                sel_tests_raw = io_manager.load_tests_list(Path(xls_tests))
            except Exception as e:
                QMessageBox.critical(self, f"Ошибка загрузки тестов ({mode})", str(e))
                return
            try:
                by_sheets = io_manager.load_equipment_by_sheets(Path(xls_eq))
                equipment_list = [it for items in by_sheets.values() for it in items]
            except Exception as e:
                QMessageBox.critical(self, f"Ошибка загрузки оборудования ({mode})", str(e))
                return

        app1_images = self.app1_images[:]
        app4_images = self.app4_images[:]
        app5_images = self.app5_images[:]

        out_base = Path(self.out_path.text().strip())
        if not out_base.name:
            QMessageBox.critical(self, "Ошибка", "Не указан путь сохранения")
            return

        out_path = out_base.with_name(out_base.stem + f"_{mode}.docx")

        # ---------- ОТЧ-OQ генерируем ТОЛЬКО вместе с OQ ----------
        base = Path(__file__).resolve().parents[1]
        tpl_report_path = None
        out_report_path = None
        ctx_fields_report = None
        xls_report_path = None

        if mode == "OQ":
            # шаблон отчёта лежит в корне проекта
            candidate_tpl = base / "Шаблон ОТЧ-OQ.docx"
            if candidate_tpl.exists():
                tpl_report_path = str(candidate_tpl)
                out_report_path = str(out_base.with_name(out_base.stem + f"_{mode}_ОТЧ-OQ.docx"))
            else:
                logger.warning(f"Шаблон отчёта ОТЧ-OQ не найден: {candidate_tpl}")

            candidate_xls = base / "ОТЧ-OQ.xlsx"  # имя файла поменяй под свой
            if candidate_xls.exists():
                xls_report_path = str(candidate_xls)
            else:
                logger.warning(f"Excel отчёта ОТЧ-OQ не найден: {candidate_xls}")

        missing_fields = []
        if not tpl_path:
            missing_fields.append("шаблон")
        if not tests_path:
            missing_fields.append("документ-тесты")
        if not xls_tests:
            missing_fields.append("Excel с тестами")
        if not xls_eq:
            missing_fields.append("Excel оборудование")
        if missing_fields:
            QMessageBox.critical(self, f"Ошибка ({mode})", "Не указаны: " + ", ".join(missing_fields))
            return

        sel_tests = [
            re.sub(r"(?i)^тест\s*\d+(?:\.\d+)?\.?\s*", "", s).strip()
            for s in sel_tests_raw if isinstance(s, str) and s.strip()
        ]

        def _norm_doc_id(raw: str) -> str:
            s = (raw or "").strip().upper()
            # если пользователь вставил уже с префиксом — срежем
            prefixes = (
                "ПРТ-OQ-", "ПРТ-PQ-", "ПРТ-IQ-",
                "ОТЧ-OQ-", "ОТЧ-PQ-", "ОТЧ-IQ-",
                "ПРТ-", "ОТЧ-",
            )
            for p in prefixes:
                if s.startswith(p):
                    s = s[len(p):]
                    break
            return s

        doc_id = _norm_doc_id(self.prt_input.text())

        # для основного документа (OQ/PQ)
        prt = RichText(f"ПРТ-{doc_id}" if doc_id else "", bold=False, italic=False, underline=False)

        # для отчёта ОТЧ-OQ
        prt_report = RichText(f"ОТЧ-OQ-{doc_id}" if doc_id else "", bold=False, italic=False, underline=False)

        # ✅ ВОТ ЭТО ДОБАВЬ: для {{prt1}} в Шаблон ОТЧ-OQ.docx
        prt1_report = RichText(f"ПРТ-OQ-{doc_id}" if doc_id else "", bold=False, italic=False, underline=False)

        # для основного документа (OQ/PQ): {{prt}} = ПРТ-<ввод>
        prt = RichText(f"ПРТ-{doc_id}" if doc_id else "", bold=False, italic=False, underline=False)

        # для отчёта ОТЧ-OQ: {{prt}} = ОТЧ-OQ-<ввод>
        prt_report = RichText(f"ОТЧ-OQ-{doc_id}" if doc_id else "", bold=False, italic=False, underline=False)

        object_text = self.object_input.text().strip()
        object_rd = ""
        if object_text:
            morph = MorphAnalyzer()
            words = object_text.split()
            parsed = [morph.parse(w)[0] for w in words]
            adj_idx = next((i for i, p in enumerate(parsed[:-1]) if "ADJF" in p.tag), None)
            noun_idx = adj_idx + 1 if adj_idx is not None and adj_idx + 1 < len(parsed) else None
            if adj_idx is not None and noun_idx is not None and "NOUN" in parsed[noun_idx].tag:
                adj = parsed[adj_idx]
                noun = parsed[noun_idx]
                gender = noun.tag.gender
                number = noun.tag.number
                adj_gent = adj.inflect({"gent", gender, number})
                noun_gent = noun.inflect({"gent"})
                if adj_gent and noun_gent:
                    before = " ".join(words[:adj_idx])
                    after = " ".join(words[noun_idx + 1:])
                    object_rd = f"{before} {adj_gent.word} {noun_gent.word} {after}".strip()
                else:
                    object_rd = object_text
            else:
                first_noun = next((p for p in parsed if "NOUN" in p.tag), None)
                if first_noun:
                    infl = first_noun.inflect({"gent"})
                    if infl:
                        idx = parsed.index(first_noun)
                        object_rd = " ".join(words[:idx] + [infl.word] + words[idx + 1:])
                    else:
                        object_rd = object_text
                else:
                    object_rd = object_text

        ctx_fields: Dict[str, Any] = {
            "объект": object_text,
            "объект1": object_rd,
            "prt": prt,
            "year": self.year_input.text().strip(),
            "customer": self.customer_input.text().strip(),
            "address": self.address_input.text().strip(),
            "Разработал": self.developed_input.text().strip(),
            "Проверил": self.checked_input.text().strip(),
            "Дата_Разработки": self.date_dev.date().toString("dd.MM.yyyy"),
            "Дата_Проверки": self.date_check.date().toString("dd.MM.yyyy"),
            "ДАТА_начала_испытания": self.date_test.date().toString("dd.MM.yyyy"),
            "ДАТА_окончания": self.date_end.date().toString("dd.MM.yyyy"),
            "Тесты_маркированным_списком": RichText("\n".join(f"• {t}" for t in sel_tests), bold=False),
        }

        if mode == "OQ" and tpl_report_path and out_report_path:
            ctx_fields_report = dict(ctx_fields)
            ctx_fields_report["prt"] = prt_report
            ctx_fields_report["prt1"] = prt1_report

        self.setEnabled(False)
        self.progress_bar.setValue(0)

        self.worker = RenderWorker(
            tpl_path, tests_path, xls_tests, xls_eq, str(out_path),
            sel_tests, self.rooms, equipment_list, ctx_fields,
            risk_path=risk_path,
            scans_dir=scans_dir,
            app1_images=app1_images,
            app4_images=app4_images,
            app5_images=app5_images,

            # ОТЧ-OQ (только для OQ):
            tpl_report_path=tpl_report_path,
            out_report_path=out_report_path,
            ctx_fields_report=ctx_fields_report,
            xls_report_path=xls_report_path,
        )

        self.worker.progress.connect(self.on_progress)
        self.worker.finished.connect(self.on_finished)
        self.worker.start()

    def on_progress(self, step: int, message: str) -> None:
        self.progress_bar.setValue(step)
        self.statusBar().showMessage(message)

    def on_finished(self, success: bool, message: str, missing: list[str]) -> None:
        if self._batch_modes:
            next_mode = self._batch_modes.pop(0)
            if success:
                if missing:
                    QMessageBox.warning(self, "Внимание", "Не вставлены тесты:\n" + "\n".join(missing))
                logger.info(message)
            else:
                QMessageBox.critical(self, "Ошибка", message)
                self._batch_modes.clear()
                self.setEnabled(True)
                self.statusBar().clearMessage()
                return
            self._start_single_render(next_mode)
            return

        self.setEnabled(True)
        if success:
            if missing:
                QMessageBox.warning(self, "Внимание", "Не вставлены тесты:\n" + "\n".join(missing))
            QMessageBox.information(self, "Успех", message)
        else:
            QMessageBox.critical(self, "Ошибка", message)
        self.statusBar().clearMessage()


def main():
    app = QApplication(sys.argv)
    wnd = MainWindow()
    wnd.resize(900, 760)
    wnd.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
