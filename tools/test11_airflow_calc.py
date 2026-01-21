from __future__ import annotations

import argparse
import re
import sys
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP, getcontext
from pathlib import Path
from typing import List, Optional

from PySide6.QtCore import Qt, QTimer, QObject, QEvent
from PySide6.QtGui import QKeySequence
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QLabel, QLineEdit, QPushButton,
    QFileDialog, QMessageBox, QSpinBox, QTableWidget, QTableWidgetItem,
    QGridLayout, QVBoxLayout, QHBoxLayout, QScrollArea, QFrame
)

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- Настройки точности ----------
getcontext().prec = 28

ROUNDING_MODE = ROUND_HALF_UP
SPEED_PLACES = 2
FLOW_PLACES = 2
AREA_PLACES = 2
SEC_PER_HOUR = Decimal(3600)

# ---------- CLI ----------
CLI = argparse.Namespace(
    template="",
    rooms=0,
    nums=[],
    names=[],
    klasses=[],
    areas=[],
    filters=[],
    points=[],
    airflows=[],  # проектные расходы м3/ч по помещениям
    auto_save="",
    auto_close=False,
)

# ---------- Decimal utils ----------
def to_decimal(s: str) -> Decimal:
    s = (str(s).strip()
         .replace("\u00A0", "")
         .replace("\u202F", "")
         .replace(" ", "")
         .replace(",", "."))
    try:
        return Decimal(s)
    except InvalidOperation:
        raise ValueError(f"Некорректное число: «{s}»")

def fmt(n: Decimal, places=2) -> str:
    q = Decimal(10) ** -places
    return str(n.quantize(q, rounding=ROUNDING_MODE)).replace(".", ",")

def fmt_speed(x: Decimal) -> str: return fmt(x, SPEED_PLACES)
def fmt_flow(x: Decimal) -> str: return fmt(x, FLOW_PLACES)
def fmt_area(x: Decimal) -> str: return fmt(x, AREA_PLACES)

def _safe_parse_decimal(s: str) -> Optional[Decimal]:
    s = (s or "").strip()
    if not s:
        return None
    try:
        return to_decimal(s)
    except Exception:
        return None

def _to_int(s: str, default: int = 0) -> int:
    try:
        return int(str(s).strip())
    except Exception:
        return default

# ---------- DOCX helpers ----------
def write_cell_text(cell, text: str):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

def cell_text_all_runs(cell) -> str:
    if cell is None:
        return ""
    return "".join(run.text for p in cell.paragraphs for run in p.runs) \
        .replace("\u00A0", " ").replace("\u202F", " ")

def iter_cells_safe(row):
    try:
        return list(row.cells)
    except Exception:
        return []

def get_cell_safe(row, col_idx):
    try:
        cells = row.cells
        if 0 <= col_idx < len(cells):
            return cells[col_idx]
    except Exception:
        return None
    return None

def write_if_cell(row, col_idx, text: str):
    c = get_cell_safe(row, col_idx)
    if c is not None:
        write_cell_text(c, text)

def _norm_head(s: str) -> str:
    s = s.replace("\u00A0", " ").replace("\u202F", " ").lower()
    s = s.replace("ё", "е")
    return re.sub(r"[^a-zа-я]+", "", s)

def insert_copies_after_tr(last_tr, template_tr, times=1):
    cur = last_tr
    for _ in range(times):
        cp = deepcopy(template_tr)
        cur.addnext(cp)
        cur = cp
    return cur

# --- vMerge ---
def clear_vmerge(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith("}vMerge"):
            tcPr.remove(child)

def set_vmerge(cell, restart=False):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith("}vMerge"):
            tcPr.remove(child)
    vMerge = OxmlElement("w:vMerge")
    vMerge.set(qn("w:val"), "restart" if restart else "continue")
    tcPr.append(vMerge)

def vmerge_cells(cells, top_value):
    if not cells:
        return
    for c in cells:
        clear_vmerge(c)
    write_cell_text(cells[0], top_value)
    set_vmerge(cells[0], restart=True)
    for c in cells[1:]:
        write_cell_text(c, "")
        set_vmerge(c, restart=False)

def collect_column_cells(table, start_row, n_rows, col_idx):
    if col_idx is None:
        return []
    if start_row < 0 or start_row + n_rows > len(table.rows):
        return []
    out = []
    for r in range(start_row, start_row + n_rows):
        row = table.rows[r]
        c = get_cell_safe(row, col_idx)
        if c is None:
            return []
        out.append(c)
    return out

# --- поиск таблицы Test11 ---
def find_table_and_template_row(doc):
    for table in doc.tables:
        for r, row in enumerate(table.rows):
            cells = iter_cells_safe(row)
            if not any("{num}" in cell_text_all_runs(c) for c in cells):
                continue

            col_filter = col_S = col_num = col_speed = None
            fact_cols = []
            for c, cell in enumerate(cells):
                txt = cell_text_all_runs(cell)
                if "{filter_num}" in txt and col_filter is None: col_filter = c
                if "{S}" in txt and col_S is None: col_S = c
                if "{num}" in txt and col_num is None: col_num = c
                if "{avg_speed}" in txt and col_speed is None: col_speed = c
                if "{fact}" in txt: fact_cols.append(c)

            if None not in (col_filter, col_S, col_num, col_speed):
                col_fact_left = fact_cols[0] if len(fact_cols) >= 1 else None
                col_fact_right = fact_cols[1] if len(fact_cols) >= 2 else None
                return table, r, col_filter, col_S, col_num, col_speed, col_fact_left, col_fact_right

    return None, None, None, None, None, None, None, None

def find_header_row(table, below_row_idx):
    for r in range(below_row_idx - 1, -1, -1):
        row_text = " | ".join(cell_text_all_runs(c).lower() for c in iter_cells_safe(table.rows[r]))
        if ("помещен" in row_text) or ("{$#}" in row_text) or ("{room}" in row_text) or ("###" in row_text):
            return r
    return None

def find_yes_no_columns(table, before_row_idx):
    col_yes = col_no = None
    for r in range(before_row_idx - 1, -1, -1):
        for c, cell in enumerate(iter_cells_safe(table.rows[r])):
            t = _norm_head(cell_text_all_runs(cell))
            if col_yes is None and (("да" in t) or ("yes" in t)):
                col_yes = c
            if col_no is None and (("нет" in t) or ("no" in t)):
                col_no = c
        if col_yes is not None and col_no is not None:
            break
    return col_yes, col_no

def find_mean_row_near(table, start_row, span):
    keys = ("средне", "average")
    def norm(s: str) -> str:
        return " ".join(s.replace("\u00A0", " ").replace("\u202F", " ").lower().split())
    end = min(start_row + span + 2, len(table.rows))
    for r in range(start_row, end):
        row_text = norm(" | ".join(cell_text_all_runs(c) for c in iter_cells_safe(table.rows[r])))
        if any(k in row_text for k in keys):
            return r
    return None

def make_room_value(num: str, name: str) -> str:
    num = (num or "").strip()
    name = (name or "").strip()
    if num and name:
        return f"{num}:{name}"
    return num or name

def replace_room_in_row(row, room_value: str, klass: str = ""):
    klass = (klass or "").strip()
    done = False

    for cell in iter_cells_safe(row):
        txt = cell_text_all_runs(cell)
        new_txt = txt

        if "{$#}" in new_txt:
            new_txt = new_txt.replace("{$#}", room_value)
        if "{room}" in new_txt or "{ROOM}" in new_txt:
            new_txt = new_txt.replace("{room}", room_value).replace("{ROOM}", room_value)

        if klass:
            new_txt = new_txt.replace("###", klass).replace("{klass}", klass).replace("{KLASS}", klass)

        if new_txt != txt:
            write_cell_text(cell, new_txt)
            done = True

    if done:
        return

    pattern = re.compile(r"(помещени[ея]\s*[:=—\-]?\s*)(.*)$", flags=re.IGNORECASE)
    for cell in iter_cells_safe(row):
        txt = cell_text_all_runs(cell)
        if "помещен" in txt.lower():
            if pattern.search(txt):
                new_txt = pattern.sub(lambda m: m.group(1) + room_value, txt)
            else:
                new_txt = txt.strip() + " " + room_value
            if klass:
                new_txt = new_txt.replace("###", klass).replace("{klass}", klass).replace("{KLASS}", klass)
            write_cell_text(cell, new_txt)
            return

    c0 = get_cell_safe(row, 0)
    if c0 is not None:
        base = f"Помещение {room_value}"
        if klass:
            base = base.replace("###", klass)
        write_cell_text(c0, base)

# ---------- Excel-like QTableWidget ----------
class SpreadsheetTable(QTableWidget):
    """QTableWidget с Excel-вставкой/копированием."""
    def keyPressEvent(self, e):
        ctrl = bool(e.modifiers() & Qt.ControlModifier)
        if ctrl and e.key() == Qt.Key_C:
            self._copy_selection()
            return
        if ctrl and e.key() == Qt.Key_V:
            if self._paste_clipboard():
                return
        super().keyPressEvent(e)

    def _copy_selection(self):
        rngs = self.selectedRanges()
        if rngs:
            r = rngs[0]
            rows = range(r.topRow(), r.bottomRow() + 1)
            cols = range(r.leftColumn(), r.rightColumn() + 1)
        else:
            rows = range(0, self.rowCount())
            cols = range(0, self.columnCount())

        lines = []
        for i in rows:
            vals = []
            for j in cols:
                it = self.item(i, j)
                vals.append("" if it is None else it.text())
            lines.append("\t".join(vals))
        QApplication.clipboard().setText("\n".join(lines))

    def _paste_clipboard(self) -> bool:
        text = QApplication.clipboard().text()
        if not text:
            return False

        rows_data = [
            [c.strip() for c in row.replace(";", "\t").split("\t")]
            for row in text.splitlines()
            if row.strip() != ""
        ]
        if not rows_data:
            return False

        rngs = self.selectedRanges()
        start_row = rngs[0].topRow() if rngs else 0
        start_col = rngs[0].leftColumn() if rngs else 0

        for i, row_vals in enumerate(rows_data):
            for j, val in enumerate(row_vals):
                r = start_row + i
                c = start_col + j
                if r >= self.rowCount() or c >= self.columnCount():
                    continue
                it = self.item(r, c)
                if it is None:
                    it = QTableWidgetItem()
                    self.setItem(r, c, it)
                it.setText(val)
        return True

# ---------- Пастинг матрицы в QLineEdit скоростей ----------
def _split_paste_matrix(text: str) -> List[str]:
    text = text.strip()
    if not text:
        return []
    lines = re.split(r"\r?\n", text)
    vals: List[str] = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if "\t" in line:
            vals.extend([c.strip() for c in line.split("\t") if c.strip() != ""])
        elif ";" in line:
            vals.extend([c.strip() for c in line.split(";") if c.strip() != ""])
        else:
            vals.extend([c.strip() for c in re.split(r"\s+", line) if c.strip() != ""])
    return vals

class SpeedPasteFilter(QObject):
    """
    Ctrl+V в любом speed QLineEdit:
    если в буфере матрица, вставляем последовательно в speed_fields_flat начиная с текущего.
    """
    def __init__(self, get_flat_list, schedule_live_update):
        super().__init__()
        self._get_flat_list = get_flat_list
        self._schedule_live_update = schedule_live_update

    def eventFilter(self, obj, event):
        if event.type() == QEvent.KeyPress:
            if event.matches(QKeySequence.Paste):
                data = QApplication.clipboard().text()
                if any(sym in data for sym in ("\n", "\r", "\t", ";")):
                    vals = _split_paste_matrix(data)
                    if not vals:
                        return True
                    flat = self._get_flat_list()
                    if not hasattr(obj, "_flat_idx"):
                        return False
                    start = int(getattr(obj, "_flat_idx"))
                    for w, v in zip(flat[start:], vals):
                        w.setText(v)
                    self._schedule_live_update()
                    return True
        return False

# ---------- Data ----------
@dataclass
class RoomRow:
    num: str
    name: str
    s_m2: str
    filters: str
    points: str

# ---------- Main Window ----------
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Калькулятор расхода воздуха (мультипомещения)")

        self._live_job: Optional[QTimer] = None

        self.rooms_table = SpreadsheetTable()
        self.rooms_table.setColumnCount(5)
        self.rooms_table.setHorizontalHeaderLabels(["Помещение", "Название", "S, м²", "Фильтров", "Точек"])
        self.rooms_table.verticalHeader().setVisible(True)

        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.speeds_root = QWidget()
        self.speeds_layout = QVBoxLayout(self.speeds_root)
        self.speeds_layout.setContentsMargins(6, 6, 6, 6)
        self.speeds_layout.setSpacing(10)
        self.scroll.setWidget(self.speeds_root)

        self.speed_fields: List[List[List[QLineEdit]]] = []  # [room][filter][point]
        self.speed_fields_flat: List[QLineEdit] = []

        self.room_crit_labels: List[QLabel] = []
        self.room_total_labels: List[QLabel] = []
        self.room_pass_labels: List[QLabel] = []
        self.filter_s_labels: List[List[QLabel]] = []
        self.filter_avg_labels: List[List[QLabel]] = []
        self.filter_flow_labels: List[List[QLabel]] = []

        self._build_ui()

        self._speed_paste_filter = SpeedPasteFilter(
            get_flat_list=lambda: self.speed_fields_flat,
            schedule_live_update=self._schedule_live_update
        )

        if CLI.template:
            self.template_path.setText(CLI.template)
        if CLI.rooms and CLI.rooms > 0:
            self.rooms_count.setValue(CLI.rooms)
        else:
            self.rooms_count.setValue(max(1, int(CLI.rooms or 1)))

        self.create_rooms()

    # ---------- UI ----------
    def _build_ui(self):
        cw = QWidget()
        main = QVBoxLayout(cw)

        r1 = QHBoxLayout()
        r1.addWidget(QLabel("Шаблон DOCX (опц.):"))
        self.template_path = QLineEdit()
        r1.addWidget(self.template_path, 1)
        btn_browse = QPushButton("Обзор…")
        btn_browse.clicked.connect(self.choose_template)
        r1.addWidget(btn_browse)
        main.addLayout(r1)

        r2 = QHBoxLayout()
        r2.addWidget(QLabel("Кол-во помещений:"))
        self.rooms_count = QSpinBox()
        self.rooms_count.setRange(1, 300)
        r2.addWidget(self.rooms_count)

        btn_rooms = QPushButton("Создать помещения")
        btn_rooms.clicked.connect(self.create_rooms)
        r2.addWidget(btn_rooms)

        btn_speeds = QPushButton("Создать поля для скоростей")
        btn_speeds.clicked.connect(self.create_speed_fields)
        r2.addWidget(btn_speeds)

        btn_calc = QPushButton("Рассчитать")
        btn_calc.clicked.connect(self.calculate_only)
        r2.addWidget(btn_calc)

        btn_save = QPushButton("Сохранить DOCX")
        btn_save.clicked.connect(self.save_docx)
        r2.addWidget(btn_save)

        r2.addStretch(1)
        main.addLayout(r2)

        main.addWidget(self.rooms_table, 0)

        line = QFrame()
        line.setFrameShape(QFrame.HLine)
        line.setFrameShadow(QFrame.Sunken)
        main.addWidget(line)

        main.addWidget(self.scroll, 1)

        r4 = QHBoxLayout()
        r4.addWidget(QLabel("Средняя скорость (справочно):"))
        self.avg_ref = QLineEdit()
        self.avg_ref.setReadOnly(True)
        self.avg_ref.setFixedWidth(120)
        r4.addWidget(self.avg_ref)
        r4.addStretch(1)
        main.addLayout(r4)

        self.setCentralWidget(cw)
        self.resize(1000, 750)

        self.rooms_table.itemChanged.connect(lambda *_: self._schedule_live_update())

    def choose_template(self):
        p, _ = QFileDialog.getOpenFileName(self, "Выберите шаблон Word", "", "Word Document (*.docx)")
        if p:
            self.template_path.setText(p)

    # ---------- rooms ----------
    def create_rooms(self):
        n = int(self.rooms_count.value())
        self.rooms_table.blockSignals(True)
        self.rooms_table.setRowCount(n)

        for r in range(n):
            num = (CLI.nums[r] if r < len(CLI.nums) and CLI.nums[r] else str(r + 1))
            name = (CLI.names[r] if r < len(CLI.names) else "")
            s = (CLI.areas[r] if r < len(CLI.areas) else "1")
            f = (CLI.filters[r] if r < len(CLI.filters) else "1")
            p = (CLI.points[r] if r < len(CLI.points) else "2")

            for c, val in enumerate([num, name, s, f, p]):
                it = self.rooms_table.item(r, c)
                if it is None:
                    it = QTableWidgetItem()
                    self.rooms_table.setItem(r, c, it)
                it.setText(str(val))

        self.rooms_table.blockSignals(False)
        self._schedule_live_update()

    def _get_rooms(self) -> List[RoomRow]:
        out: List[RoomRow] = []
        for r in range(self.rooms_table.rowCount()):
            def cell(c: int) -> str:
                it = self.rooms_table.item(r, c)
                return (it.text().strip() if it else "")
            out.append(RoomRow(
                num=cell(0),
                name=cell(1),
                s_m2=cell(2),
                filters=cell(3),
                points=cell(4),
            ))
        return out

    # ---------- speeds ----------
    def _clear_speeds_ui(self):
        while self.speeds_layout.count():
            item = self.speeds_layout.takeAt(0)
            w = item.widget()
            if w is not None:
                w.deleteLater()

        self.speed_fields = []
        self.speed_fields_flat = []

        self.room_crit_labels = []
        self.room_total_labels = []
        self.room_pass_labels = []
        self.filter_s_labels = []
        self.filter_avg_labels = []
        self.filter_flow_labels = []

    def create_speed_fields(self):
        rooms = self._get_rooms()
        if not rooms:
            QMessageBox.warning(self, "Ошибка", "Сначала создайте помещения.")
            return

        for i, rr in enumerate(rooms, start=1):
            if not rr.name.strip():
                QMessageBox.warning(self, "Ошибка", f"Не заполнено название (строка {i}).")
                return
            S = _safe_parse_decimal(rr.s_m2)
            if S is None or S <= 0:
                QMessageBox.warning(self, "Ошибка", f"Площадь S должна быть > 0 (строка {i}).")
                return
            nf = _to_int(rr.filters, 0)
            np = _to_int(rr.points, 0)
            if nf <= 0 or np <= 0:
                QMessageBox.warning(self, "Ошибка", f"Фильтров и Точек должны быть > 0 (строка {i}).")
                return

        self._clear_speeds_ui()

        for r_idx, rr in enumerate(rooms):
            nf = _to_int(rr.filters, 1)
            np = _to_int(rr.points, 2)

            room_box = QFrame()
            room_box.setFrameShape(QFrame.StyledPanel)
            vb = QVBoxLayout(room_box)
            vb.setContentsMargins(8, 8, 8, 8)
            vb.setSpacing(6)

            title = QLabel(f"Помещение {rr.num}")
            title.setStyleSheet("font-weight: 700;")
            vb.addWidget(title)

            summ = QHBoxLayout()
            summ.addWidget(QLabel("Критерий (м³/ч):"))
            lbl_crit = QLabel("—")
            lbl_crit.setFixedWidth(120)
            summ.addWidget(lbl_crit)

            summ.addSpacing(16)
            summ.addWidget(QLabel("Итого факт. (м³/ч):"))
            lbl_total = QLabel("—")
            lbl_total.setFixedWidth(120)
            summ.addWidget(lbl_total)

            summ.addSpacing(16)
            summ.addWidget(QLabel("Результат:"))
            lbl_pass = QLabel("—")
            lbl_pass.setStyleSheet("font-weight: 700;")
            lbl_pass.setFixedWidth(60)
            summ.addWidget(lbl_pass)

            summ.addStretch(1)
            vb.addLayout(summ)

            self.room_crit_labels.append(lbl_crit)
            self.room_total_labels.append(lbl_total)
            self.room_pass_labels.append(lbl_pass)

            room_speed_filters: List[List[QLineEdit]] = []
            room_s_lbls: List[QLabel] = []
            room_avg_lbls: List[QLabel] = []
            room_flow_lbls: List[QLabel] = []

            for f_idx in range(nf):
                fl_title = QLabel(f"  Фильтр {f_idx+1}")
                fl_title.setStyleSheet("font-style: italic;")
                vb.addWidget(fl_title)

                grid = QGridLayout()
                grid.setHorizontalSpacing(10)
                grid.setVerticalSpacing(4)

                filt_points: List[QLineEdit] = []
                for p_idx in range(np):
                    lab = QLabel(f"    № точки {p_idx+1} (м/с):")
                    e = QLineEdit()
                    e.setFixedWidth(120)
                    e._flat_idx = len(self.speed_fields_flat)  # type: ignore[attr-defined]
                    e.installEventFilter(self._speed_paste_filter)
                    e.textChanged.connect(lambda *_: self._schedule_live_update())
                    grid.addWidget(lab, p_idx, 0, Qt.AlignRight)
                    grid.addWidget(e, p_idx, 1, Qt.AlignLeft)

                    filt_points.append(e)
                    self.speed_fields_flat.append(e)

                vb.addLayout(grid)

                row_sum = QHBoxLayout()
                row_sum.addWidget(QLabel("    S, м²:"))
                lbl_s = QLabel("—")
                lbl_s.setFixedWidth(100)
                row_sum.addWidget(lbl_s)

                row_sum.addSpacing(16)
                row_sum.addWidget(QLabel("Средняя (м/с):"))
                lbl_avg = QLabel("—")
                lbl_avg.setFixedWidth(100)
                row_sum.addWidget(lbl_avg)

                row_sum.addSpacing(16)
                row_sum.addWidget(QLabel("Итого факт (м³/ч):"))
                lbl_flow = QLabel("—")
                lbl_flow.setFixedWidth(120)
                row_sum.addWidget(lbl_flow)

                row_sum.addStretch(1)
                vb.addLayout(row_sum)

                room_s_lbls.append(lbl_s)
                room_avg_lbls.append(lbl_avg)
                room_flow_lbls.append(lbl_flow)

                room_speed_filters.append(filt_points)

            self.speed_fields.append(room_speed_filters)
            self.filter_s_labels.append(room_s_lbls)
            self.filter_avg_labels.append(room_avg_lbls)
            self.filter_flow_labels.append(room_flow_lbls)

            self.speeds_layout.addWidget(room_box)

        self.speeds_layout.addStretch(1)
        self._schedule_live_update()

    # ---------- live update ----------
    def _schedule_live_update(self):
        if self._live_job is None:
            self._live_job = QTimer(self)
            self._live_job.setSingleShot(True)
            self._live_job.timeout.connect(self._live_update)
        self._live_job.start(120)

    def _live_update(self):
        rooms = self._get_rooms()
        n = len(rooms)

        crit: List[Optional[Decimal]] = []
        for r in range(n):
            txt = (CLI.airflows[r] if r < len(CLI.airflows) else "") or ""
            crit.append(_safe_parse_decimal(txt))

        for r_idx, rr in enumerate(rooms):
            S = _safe_parse_decimal(rr.s_m2)
            if S is not None and S <= 0:
                S = None

            if r_idx < len(self.room_crit_labels):
                self.room_crit_labels[r_idx].setText(fmt_flow(crit[r_idx]) if crit[r_idx] is not None else "—")

            if r_idx >= len(self.speed_fields):
                if r_idx < len(self.room_total_labels):
                    self.room_total_labels[r_idx].setText("—")
                if r_idx < len(self.room_pass_labels):
                    self.room_pass_labels[r_idx].setText("—")
                continue

            total_room = Decimal(0)
            have_all = True

            nf = len(self.speed_fields[r_idx])
            for f_idx in range(nf):
                if r_idx < len(self.filter_s_labels) and f_idx < len(self.filter_s_labels[r_idx]):
                    self.filter_s_labels[r_idx][f_idx].setText(fmt_area(S) if S is not None else "—")

                pts = self.speed_fields[r_idx][f_idx]
                vals: List[Decimal] = []
                ok = True
                for e in pts:
                    v = _safe_parse_decimal(e.text())
                    if v is None or v < 0:
                        ok = False
                        break
                    vals.append(v)

                if (not ok) or (S is None) or (len(vals) != len(pts)):
                    if r_idx < len(self.filter_avg_labels) and f_idx < len(self.filter_avg_labels[r_idx]):
                        self.filter_avg_labels[r_idx][f_idx].setText("—")
                    if r_idx < len(self.filter_flow_labels) and f_idx < len(self.filter_flow_labels[r_idx]):
                        self.filter_flow_labels[r_idx][f_idx].setText("—")
                    have_all = False
                    continue

                avg = sum(vals, Decimal(0)) / Decimal(len(vals))
                flow = S * avg * SEC_PER_HOUR

                if r_idx < len(self.filter_avg_labels) and f_idx < len(self.filter_avg_labels[r_idx]):
                    self.filter_avg_labels[r_idx][f_idx].setText(fmt_speed(avg))
                if r_idx < len(self.filter_flow_labels) and f_idx < len(self.filter_flow_labels[r_idx]):
                    self.filter_flow_labels[r_idx][f_idx].setText(fmt_flow(flow))

                total_room += flow

            if r_idx < len(self.room_total_labels):
                self.room_total_labels[r_idx].setText(fmt_flow(total_room) if have_all and nf > 0 else "—")
            if r_idx < len(self.room_pass_labels):
                if have_all and nf > 0 and (crit[r_idx] is not None):
                    self.room_pass_labels[r_idx].setText("ДА" if total_room >= crit[r_idx] else "НЕТ")
                else:
                    self.room_pass_labels[r_idx].setText("—")

        try:
            if self.speed_fields and self.speed_fields[0]:
                vals = []
                for e in self.speed_fields[0][0]:
                    v = _safe_parse_decimal(e.text())
                    if v is None:
                        vals = []
                        break
                    vals.append(v)
                if vals:
                    avg0 = sum(vals, Decimal(0)) / Decimal(len(vals))
                    self.avg_ref.setText(fmt_speed(avg0))
                else:
                    self.avg_ref.setText("—")
        except Exception:
            self.avg_ref.setText("—")

    # ---------- calc / save ----------
    def _collect_inputs_strict(self):
        rooms = self._get_rooms()
        if not rooms:
            raise ValueError("Нет помещений.")

        n_rooms = len(rooms)
        room_names: List[str] = []
        room_nums: List[str] = []
        room_S: List[Decimal] = []
        room_filters: List[int] = []
        room_points: List[int] = []

        for r, rr in enumerate(rooms):
            if not rr.name.strip():
                raise ValueError(f"Не заполнено название для помещения {r+1}.")
            room_names.append(rr.name.strip())

            room_nums.append(rr.num.strip() or str(r + 1))

            S = to_decimal(rr.s_m2)
            if S <= 0:
                raise ValueError(f"Площадь S должна быть > 0 (помещение {r+1}).")
            room_S.append(S)

            nf = _to_int(rr.filters, 0)
            np = _to_int(rr.points, 0)
            if nf <= 0:
                raise ValueError(f"Фильтров на помещение должно быть > 0 (помещение {r+1}).")
            if np <= 0:
                raise ValueError(f"Количество точек должно быть > 0 (помещение {r+1}).")
            room_filters.append(nf)
            room_points.append(np)

        if not self.speed_fields or len(self.speed_fields) != n_rooms:
            raise ValueError("Нажмите «Создать поля для скоростей» и заполните скорости.")

        speeds: List[List[List[Decimal]]] = []
        avgs: List[List[Decimal]] = []

        for r in range(n_rooms):
            if len(self.speed_fields[r]) != room_filters[r]:
                raise ValueError(f"Несовпадение кол-ва фильтров в скоростях (помещение {r+1}).")

            sp_room: List[List[Decimal]] = []
            avg_room: List[Decimal] = []

            for f in range(room_filters[r]):
                pts = self.speed_fields[r][f]
                if len(pts) != room_points[r]:
                    raise ValueError(f"Несовпадение кол-ва точек (помещение {r+1}, фильтр {f+1}).")

                row_vals: List[Decimal] = []
                for i in range(room_points[r]):
                    raw = pts[i].text().strip()
                    if raw == "":
                        raise ValueError(f"Пустая скорость: пом. {r+1}, фильтр {f+1}, точка {i+1}.")
                    v = to_decimal(raw)
                    if v < 0:
                        raise ValueError(f"Отрицательная скорость: пом. {r+1}, фильтр {f+1}, точка {i+1}.")
                    row_vals.append(v)

                sp_room.append(row_vals)
                avg_room.append(sum(row_vals, Decimal(0)) / Decimal(room_points[r]))

            speeds.append(sp_room)
            avgs.append(avg_room)

        room_sum_flow: List[Decimal] = []
        for r in range(n_rooms):
            total = sum((room_S[r] * avgs[r][f] * SEC_PER_HOUR for f in range(room_filters[r])), Decimal(0))
            room_sum_flow.append(total)

        criterion_texts: List[Optional[str]] = []
        for r in range(n_rooms):
            txt = ((CLI.airflows[r] if r < len(CLI.airflows) else "") or "").strip()
            if txt:
                try:
                    criterion_texts.append(fmt_flow(to_decimal(txt)))
                except Exception:
                    criterion_texts.append(txt)
            else:
                criterion_texts.append(None)

        return (n_rooms, room_nums, room_names, room_S, room_filters, room_points,
                speeds, avgs, room_sum_flow, criterion_texts)

    def calculate_only(self):
        try:
            self._live_update()
            (n_rooms, *_rest) = self._collect_inputs_strict()
            QMessageBox.information(self, "Ок", f"Данные корректны. Помещений: {n_rooms}.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def save_docx(self):
        try:
            (n_rooms, room_nums, room_names, room_S, room_filters, room_points,
             speeds, avgs, room_sum_flow, criterion_texts) = self._collect_inputs_strict()

            # ВАЖНО: room_klasses вычисляем ТОЛЬКО тут (когда уже есть n_rooms)
            room_klasses = list(getattr(CLI, "klasses", []) or [])
            if len(room_klasses) < n_rooms:
                room_klasses += [""] * (n_rooms - len(room_klasses))

            template_path = self.template_path.text().strip()
            if not template_path:
                p, _ = QFileDialog.getOpenFileName(self, "Выберите DOCX шаблон Test 11", "", "Word Document (*.docx)")
                if not p:
                    return
                template_path = p
                self.template_path.setText(p)

            try:
                doc = Document(template_path)
            except Exception as e:
                raise ValueError(f"Не удалось открыть шаблон: {e}")

            (table, tpl_row, col_filter, col_S, col_num, col_speed,
             col_fact_left, col_fact_right) = find_table_and_template_row(doc)
            if table is None:
                raise ValueError("В шаблоне не найдена строка со спецификатором {num}. Проверьте таблицу Test 11.")

            missing = []
            if col_filter is None: missing.append("{filter_num}")
            if col_S is None: missing.append("{S}")
            if col_num is None: missing.append("{num}")
            if col_speed is None: missing.append("{avg_speed}")
            if col_fact_left is None: missing.append("{fact}")
            if missing:
                raise ValueError("В строке-шаблоне отсутствуют столбцы: " + ", ".join(missing))

            header_row_idx = find_header_row(table, tpl_row)
            if header_row_idx is None:
                raise ValueError("Не найден заголовок секции «Помещение …» выше строки-шаблона.")

            col_yes, col_no = find_yes_no_columns(table, tpl_row)

            first_yes_no = None
            for c in (col_yes, col_no):
                if c is not None:
                    first_yes_no = c if first_yes_no is None else min(first_yes_no, c)

            col_criterion = None
            if first_yes_no is not None and first_yes_no > 0:
                col_criterion = first_yes_no - 1  # слева от ДА/НЕТ

            data_tr_template = deepcopy(table.rows[tpl_row]._tr)

            mean_row_idx = find_mean_row_near(table, tpl_row, room_points[0])
            has_mean = (mean_row_idx is not None)
            mean_tr_template = deepcopy(table.rows[mean_row_idx]._tr) if has_mean else None

            def merge_block(start_row, n_rows, filter_label, S_value: Decimal):
                filter_cells = collect_column_cells(table, start_row, n_rows, col_filter)
                s_cells = collect_column_cells(table, start_row, n_rows, col_S)
                factL_cells = collect_column_cells(table, start_row, n_rows, col_fact_left)
                vmerge_cells(filter_cells, str(filter_label))
                vmerge_cells(s_cells, fmt_area(S_value))
                if factL_cells:
                    vmerge_cells(factL_cells, cell_text_all_runs(factL_cells[0]))

            # --- room 0 ---
            replace_room_in_row(
                table.rows[header_row_idx],
                make_room_value(room_nums[0], room_names[0]),
                room_klasses[0],
            )

            if room_points[0] > 1:
                last_tr = table.rows[tpl_row]._tr
                last_tr = insert_copies_after_tr(last_tr, data_tr_template, times=room_points[0] - 1)

            block_fact_value = fmt_flow(room_S[0] * avgs[0][0] * SEC_PER_HOUR)
            for i in range(room_points[0]):
                row = table.rows[tpl_row + i]
                write_if_cell(row, col_num, str(i + 1))
                write_if_cell(row, col_speed, fmt_speed(speeds[0][0][i]))
                write_if_cell(row, col_fact_left, block_fact_value if i == 0 else "")
                if col_fact_right is not None:
                    write_if_cell(row, col_fact_right, "")

            if has_mean:
                if tpl_row + room_points[0] >= len(table.rows):
                    last_tr = table.rows[len(table.rows) - 1]._tr
                    last_tr = insert_copies_after_tr(last_tr, mean_tr_template, times=1)
                write_if_cell(table.rows[tpl_row + room_points[0]], col_speed, fmt_speed(avgs[0][0]))

            block_height_for_room0 = room_points[0] + (1 if has_mean else 0)
            merge_block(tpl_row, block_height_for_room0, 1, room_S[0])

            room_block_starts: List[List[int]] = [[tpl_row]]
            current_end_row = tpl_row + block_height_for_room0 - 1

            for f in range(2, room_filters[0] + 1):
                last_tr = table.rows[current_end_row]._tr
                for _ in range(room_points[0]):
                    last_tr = insert_copies_after_tr(last_tr, data_tr_template, times=1)
                if has_mean:
                    last_tr = insert_copies_after_tr(last_tr, mean_tr_template, times=1)

                block_start = current_end_row + 1
                current_end_row = block_start + room_points[0] + (1 if has_mean else 0) - 1

                block_fact_value = fmt_flow(room_S[0] * avgs[0][f - 1] * SEC_PER_HOUR)
                for i in range(room_points[0]):
                    row = table.rows[block_start + i]
                    write_if_cell(row, col_num, str(i + 1))
                    write_if_cell(row, col_speed, fmt_speed(speeds[0][f - 1][i]))
                    write_if_cell(row, col_fact_left, block_fact_value if i == 0 else "")
                    if col_fact_right is not None:
                        write_if_cell(row, col_fact_right, "")
                if has_mean:
                    write_if_cell(table.rows[block_start + room_points[0]], col_speed, fmt_speed(avgs[0][f - 1]))

                merge_block(block_start, room_points[0] + (1 if has_mean else 0), f, room_S[0])
                room_block_starts[0].append(block_start)

            header_tr_template = deepcopy(table.rows[header_row_idx]._tr)

            for r in range(1, n_rooms):
                last_tr = table.rows[current_end_row]._tr
                last_tr = insert_copies_after_tr(last_tr, header_tr_template, times=1)
                room_header_row = current_end_row + 1
                current_end_row = room_header_row

                replace_room_in_row(
                    table.rows[room_header_row],
                    make_room_value(room_nums[r], room_names[r]),
                    room_klasses[r],
                )

                starts_this_room: List[int] = []
                for f in range(1, room_filters[r] + 1):
                    for _ in range(room_points[r]):
                        last_tr = insert_copies_after_tr(last_tr, data_tr_template, times=1)
                    if has_mean:
                        last_tr = insert_copies_after_tr(last_tr, mean_tr_template, times=1)

                    block_start = current_end_row + 1
                    current_end_row = block_start + room_points[r] + (1 if has_mean else 0) - 1

                    block_fact_value = fmt_flow(room_S[r] * avgs[r][f - 1] * SEC_PER_HOUR)
                    for i in range(room_points[r]):
                        row = table.rows[block_start + i]
                        write_if_cell(row, col_num, str(i + 1))
                        write_if_cell(row, col_speed, fmt_speed(speeds[r][f - 1][i]))
                        write_if_cell(row, col_fact_left, block_fact_value if i == 0 else "")
                        if col_fact_right is not None:
                            write_if_cell(row, col_fact_right, "")
                    if has_mean:
                        write_if_cell(table.rows[block_start + room_points[r]], col_speed, fmt_speed(avgs[r][f - 1]))

                    merge_block(block_start, room_points[r] + (1 if has_mean else 0), f, room_S[r])
                    starts_this_room.append(block_start)

                room_block_starts.append(starts_this_room)

            def block_height_fn(room_idx: int) -> int:
                return room_points[room_idx] + (1 if has_mean else 0)

            for r, starts in enumerate(room_block_starts):
                if not starts:
                    continue
                first_start = starts[0]
                last_start = starts[-1]
                seg_start = first_start
                seg_rows = (last_start + block_height_fn(r)) - seg_start

                if col_fact_right is not None:
                    cellsR = collect_column_cells(table, seg_start, seg_rows, col_fact_right)
                    if cellsR:
                        vmerge_cells(cellsR, fmt_flow(room_sum_flow[r]))

                if col_criterion is not None and col_criterion not in (col_yes, col_no):
                    cellsC = collect_column_cells(table, seg_start, seg_rows, col_criterion)
                    if cellsC:
                        top_val = criterion_texts[r] if (r < len(criterion_texts) and criterion_texts[r]) else cell_text_all_runs(cellsC[0])
                        vmerge_cells(cellsC, top_val)

            out_path = (CLI.auto_save or "").strip()
            if not out_path:
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                default_name = f"Расчет_воздуха_{ts}.docx"
                p, _ = QFileDialog.getSaveFileName(self, "Сохранить отчёт", default_name, "Word Document (*.docx)")
                if not p:
                    return
                out_path = p

            try:
                Path(out_path).parent.mkdir(parents=True, exist_ok=True)
                doc.save(out_path)
            except PermissionError:
                QMessageBox.critical(self, "Файл занят", "Документ открыт в Word. Закройте и повторите.")
                return

            if not CLI.auto_save:
                QMessageBox.information(self, "Готово", f"Файл сохранён:\n{out_path}")

            if CLI.auto_close:
                self.close()

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"{e}")

# ---------- CLI parse ----------
def _parse_args(argv):
    p = argparse.ArgumentParser(description="Калькулятор Test 11 (расход приточного воздуха)")
    p.add_argument("--template", default="", help="Путь к DOCX с таблицей Test 11")
    p.add_argument("--rooms", type=int, default=0, help="Количество помещений (предзаполнение)")
    p.add_argument("--nums", default="", help="Номера помещений через ;")
    p.add_argument("--names", default="", help="Названия помещений через ;")
    p.add_argument("--klasses", default="", help="Класс чистоты через ;")
    p.add_argument("--areas", default="", help="Площади через ; (точка или запятая)")
    p.add_argument("--filters", default="", help="Кол-во фильтров на помещение через ;")
    p.add_argument("--points", default="", help="Кол-во точек на помещение через ;")
    p.add_argument("--airflows", default="", help="Проектные расходы, м3/ч, через ; (по помещениям)")
    p.add_argument("--auto-save", default="", help="Путь авто-сохранения (без диалога)")
    p.add_argument("--auto-close", action="store_true", help="Закрыть окно после сохранения")
    args = p.parse_args(argv)

    CLI.template = args.template
    CLI.rooms = max(0, int(args.rooms or 0))
    CLI.auto_save = args.auto_save
    CLI.auto_close = bool(args.auto_close)

    def split_list(s):
        return [x.strip() for x in s.split(";")] if s else []

    CLI.nums = split_list(args.nums)
    CLI.names = split_list(args.names)
    CLI.klasses = split_list(args.klasses)
    CLI.areas = split_list(args.areas)
    CLI.filters = split_list(args.filters)
    CLI.points = split_list(args.points)
    CLI.airflows = split_list(args.airflows)

def main():
    app = QApplication(sys.argv)
    wnd = MainWindow()
    wnd.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    _parse_args(sys.argv[1:])
    main()
