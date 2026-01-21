# risk_table5.py
from __future__ import annotations

import re
from copy import deepcopy
from typing import Dict, List, Tuple, Optional

from openpyxl import load_workbook

from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =============================================================================
# Нормализация / сравнение тестов (очень важно для совпадений)
# =============================================================================

def _norm_basic(s: str) -> str:
    s = "" if s is None else str(s)
    s = (
        s.replace("\u00A0", " ")   # nbsp
        .replace("\u202F", " ")   # narrow nbsp
        .replace("\u200B", "")    # zero-width
        .replace("\u00AD", "")    # soft hyphen
    )
    s = re.sub(r"\s+", " ", s).strip().lower()
    s = s.replace("ё", "е")
    return s


def _match_key(s: str) -> str:
    """
    Ключ для сравнения: выкидываем пробелы/пунктуацию.
    Тогда совпадут:
      "обучении/ ознакомлении" и "обучении / ознакомлении"
      с точками/без точек и т.п.
    """
    return re.sub(r"[^0-9a-zа-я]+", "", _norm_basic(s))


def _split_tests_cell(s: str) -> List[str]:
    """
    В Excel в 'Аттестационное испытание' может быть:
    - одно название
    - несколько через перенос строки / ;
    """
    raw = _norm_basic(s)
    if not raw:
        return []

    # но возвращаем ОРИГИНАЛЬНЫЕ части (не lower), поэтому парсим по исходной строке
    src = (s or "").replace("\u00A0", " ").strip()
    parts = re.split(r"[\n;]+", src)
    out: List[str] = []
    for p in parts:
        p = p.strip(" \t\r,")
        if p:
            out.append(p)
    return out


def _selected_tests_to_names(selected_tests: List[str]) -> List[str]:
    """
    selected_tests приходит из UI иногда с префиксами:
      '01. Тест 11.1 ...'
    Нам нужна "чистая" часть (как в Excel).
    """
    names: List[str] = []
    for t in selected_tests or []:
        t0 = (t or "").strip()
        if not t0:
            continue

        # убрать "01."
        t0 = re.sub(r"^\s*\d+\s*\.\s*", "", t0)

        # убрать "Тест 11.1" / "Тест 11" и т.п.
        t0 = re.sub(r"^\s*тест\s*\d+(?:\.\d+)?\.?\s*", "", t0, flags=re.IGNORECASE)

        names.append(t0.strip())
    return names


def _is_match(selected_name: str, test_from_excel: str) -> bool:
    """
    Матч "вхождение ключа" в любую сторону.
    """
    sk = _match_key(selected_name)
    tk = _match_key(test_from_excel)
    if not sk or not tk:
        return False
    return (sk in tk) or (tk in sk)


# =============================================================================
# Excel: чтение строк
# =============================================================================

def _build_header_map(ws) -> Dict[str, int]:
    """
    Мапа "нормализованный заголовок" -> индекс колонки (1-based).
    Считаем, что заголовок в 1-й строке.
    """
    header_map: Dict[str, int] = {}
    for c in range(1, ws.max_column + 1):
        v = ws.cell(row=1, column=c).value
        if v is None:
            continue
        header_map[_norm_basic(str(v))] = c
    return header_map


def _h(header_map: Dict[str, int], *variants: str) -> int | None:
    for v in variants:
        k = _norm_basic(v)
        if k in header_map:
            return header_map[k]
    return None


def get_risk_rows(xlsx_path: str, selected_tests: List[str]) -> List[Dict[str, str]]:
    """
    Возвращает строки для Таблицы 5:
    - берём только те строки Excel, где в "Аттестационное испытание" есть хотя бы 1 выбранный тест
    - ВНУТРИ строки оставляем только выбранные тесты (в порядке выбора пользователем)
    - одинаковые строки (по всем полям кроме тестов) схлопываем, объединяя тесты
    """
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb.active

    header_map = _build_header_map(ws)

    cols = {
        "risk": _h(header_map, "Риск"),
        "cause": _h(header_map, "Возможная причина"),
        "prob_l": _h(header_map, "Вероятность_оценка", "Вероятность оценка"),
        "prob_s": _h(header_map, "Вероятность_балл", "Вероятность балл"),
        "sev_l": _h(header_map, "Тяжесть_оценка", "Тяжесть оценка"),
        "sev_s": _h(header_map, "Тяжесть_балл", "Тяжесть балл"),
        "det_l": _h(header_map, "Необнаружение_оценка", "Необнаружение оценка"),
        "det_s": _h(header_map, "Необнаружение_балл", "Необнаружение балл"),
        "level_l": _h(header_map, "Уровень_риска", "Уровень риска"),
        "rpn": _h(header_map, "ПЧР", "RPN"),
        "tests": _h(header_map, "Аттестационное испытание", "Квалификационное испытание"),
    }

    if any(v is None for v in cols.values()):
        raise ValueError(f"Не найдены все нужные колонки в Excel (ожидается строка 1). Найдено: {cols}")

    sel_names = _selected_tests_to_names(selected_tests)
    sel_names = [s for s in sel_names if s.strip()]  # сохранить порядок выбора

    def sval(v) -> str:
        return "" if v is None else str(v).strip()

    def ival(v) -> str:
        if v is None:
            return ""
        s = str(v).strip()
        # 2.0 -> 2
        if re.fullmatch(r"\d+\.0", s):
            s = s[:-2]
        return s

    raw_rows: List[Dict[str, str | List[str]]] = []

    for r in range(2, ws.max_row + 1):
        risk = sval(ws.cell(row=r, column=cols["risk"]).value)
        cause = sval(ws.cell(row=r, column=cols["cause"]).value)
        tests_raw = ws.cell(row=r, column=cols["tests"]).value
        tests_list = _split_tests_cell("" if tests_raw is None else str(tests_raw))

        if not (risk or cause or tests_list):
            continue

        # ОСТАВЛЯЕМ ТОЛЬКО выбранные тесты (и в ПОРЯДКЕ выбора)
        matched_tests: List[str] = []
        if sel_names:
            for sname in sel_names:
                for t in tests_list:
                    if _is_match(sname, t):
                        if t not in matched_tests:
                            matched_tests.append(t)
            if not matched_tests:
                continue
        else:
            matched_tests = tests_list

        row = {
            "risk": risk,
            "cause": cause,
            "prob_letter": sval(ws.cell(row=r, column=cols["prob_l"]).value),
            "prob_score": ival(ws.cell(row=r, column=cols["prob_s"]).value),
            "sev_letter": sval(ws.cell(row=r, column=cols["sev_l"]).value),
            "sev_score": ival(ws.cell(row=r, column=cols["sev_s"]).value),
            "det_letter": sval(ws.cell(row=r, column=cols["det_l"]).value),
            "det_score": ival(ws.cell(row=r, column=cols["det_s"]).value),
            "level_letter": sval(ws.cell(row=r, column=cols["level_l"]).value),
            "rpn": ival(ws.cell(row=r, column=cols["rpn"]).value),
            "tests": matched_tests,
        }
        raw_rows.append(row)

    # СХЛОПНУТЬ дубли по всем полям, кроме tests
    merged: List[Dict[str, str | List[str]]] = []
    index: Dict[Tuple, int] = {}

    for rr in raw_rows:
        k = (
            _match_key(rr["risk"]),
            _match_key(rr["cause"]),
            rr["prob_letter"],
            rr["prob_score"],
            rr["sev_letter"],
            rr["sev_score"],
            rr["det_letter"],
            rr["det_score"],
            rr["level_letter"],
            rr["rpn"],
        )
        if k in index:
            ex = merged[index[k]]
            ex_tests = ex["tests"]  # type: ignore[assignment]
            for t in rr["tests"]:   # type: ignore[operator]
                if t not in ex_tests:
                    ex_tests.append(t)
        else:
            index[k] = len(merged)
            merged.append(rr)

    # Привести к финальному виду (tests -> list[str] оставляем, дальше в docx вставке сделаем буллеты)
    return [r for r in merged]  # type: ignore[return-value]


# =============================================================================
# DOCX: вставка Таблицы 5 по плейсхолдерам, НЕ ТРОГАЯ ШАПКУ
# =============================================================================

def _find_table5_and_template_row(doc: DocxDocument) -> Tuple[Table, int]:
    """
    Ищем строку-шаблон по маркеру <<T5_RISK>>.
    """
    for tbl in doc.tables:
        for ri, row in enumerate(tbl.rows):
            if any("<<T5_RISK>>" in c.text for c in row.cells):
                return tbl, ri
    raise ValueError("Не найдена Таблица 5: нет строки с <<T5_RISK>>.")


def _remove_rows_from(table: Table, start_idx: int) -> None:
    """
    Удаляем строки начиная со start_idx и до конца.
    """
    while len(table.rows) > start_idx:
        tr = table.rows[start_idx]._tr
        table._tbl.remove(tr)


def _set_para_text_keep_runs(p, text: str) -> None:
    """
    Сменить текст в первом run и очистить остальные,
    чтобы не ломать стиль абзаца.
    """
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _set_cell_lines(cell: _Cell, lines: List[str]) -> None:
    """
    Записать список строк как отдельные абзацы в ячейку.
    """
    if not lines:
        lines = [""]

    while len(cell.paragraphs) < len(lines):
        cell.add_paragraph("")

    for i, txt in enumerate(lines):
        _set_para_text_keep_runs(cell.paragraphs[i], txt)

    # остальные абзацы очистим
    for j in range(len(lines), len(cell.paragraphs)):
        _set_para_text_keep_runs(cell.paragraphs[j], "")


def _set_diag_cell(cell: _Cell, top_left: str, bottom_right: str,
                   bottom_space_before_pt: float = 36.0, after_pt: float = 3.0) -> None:
    """
    Диагональная ячейка: 2 абзаца:
      1) слева (буква)
      2) справа снизу (балл)
    Не трогаем границы/диагональ — они в XML ячейки.
    """
    while len(cell.paragraphs) < 2:
        cell.add_paragraph("")

    p_top = cell.paragraphs[0]
    p_bot = cell.paragraphs[1]

    _set_para_text_keep_runs(p_top, top_left or "")
    p_top.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_top.paragraph_format.space_before = Pt(0)
    p_top.paragraph_format.space_after = Pt(after_pt)

    _set_para_text_keep_runs(p_bot, bottom_right or "")
    p_bot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_bot.paragraph_format.space_before = Pt(bottom_space_before_pt)
    p_bot.paragraph_format.space_after = Pt(after_pt)

    for p in cell.paragraphs[2:]:
        _set_para_text_keep_runs(p, "")


def insert_table5_into_doc(doc: DocxDocument, risk_rows: List[Dict[str, str]]) -> None:
    """
    Заполняем Таблицу 5 по строкам из get_risk_rows():
    - Шапку не трогаем
    - Строку с плейсхолдерами удаляем
    - Вставляем новые строки, копируя формат строки-шаблона
    - Риск объединяем по вертикали, если одинаковый подряд
    """
    table, tpl_row_idx = _find_table5_and_template_row(doc)

    tpl_tr = deepcopy(table.rows[tpl_row_idx]._tr)

    # удаляем плейсхолдер-строку и всё ниже (чтобы заглушки не оставались)
    _remove_rows_from(table, tpl_row_idx)

    for rr in (risk_rows or []):
        table._tbl.append(deepcopy(tpl_tr))
        new_row = table.rows[-1]
        cells = new_row.cells

        # ВАЖНО: в шаблоне ожидается 7 визуальных колонок:
        # 0 риск | 1 причина | 2 prob | 3 sev | 4 det | 5 level | 6 tests
        _set_cell_lines(cells[0], [str(rr.get("risk", "") or "")])
        _set_cell_lines(cells[1], [str(rr.get("cause", "") or "")])

        _set_diag_cell(cells[2], str(rr.get("prob_letter", "") or ""), str(rr.get("prob_score", "") or ""))
        _set_diag_cell(cells[3], str(rr.get("sev_letter", "") or ""), str(rr.get("sev_score", "") or ""))
        _set_diag_cell(cells[4], str(rr.get("det_letter", "") or ""), str(rr.get("det_score", "") or ""))
        _set_diag_cell(cells[5], str(rr.get("level_letter", "") or ""), str(rr.get("rpn", "") or ""))

        tests_list = rr.get("tests")  # может быть list[str]
        if isinstance(tests_list, list):
            lines = [f"•→{t}" for t in tests_list if str(t).strip()]
        else:
            # если вдруг пришло строкой
            raw = str(tests_list or "").strip()
            lines = [f"•→{x.strip()}" for x in raw.splitlines() if x.strip()] if raw else []

        _set_cell_lines(cells[6], lines if lines else [""])

    # Объединяем одинаковые риски подряд (как в образце)
    if not risk_rows:
        return

    risk_keys = [_match_key(str(r.get("risk", ""))) for r in risk_rows]

    i = 0
    while i < len(risk_rows):
        j = i + 1
        while j < len(risk_rows) and risk_keys[j] == risk_keys[i]:
            j += 1

        if j - i > 1:
            top_cell = table.rows[tpl_row_idx + i].cells[0]
            bottom_cell = table.rows[tpl_row_idx + (j - 1)].cells[0]
            top_cell.merge(bottom_cell)
            _set_cell_lines(top_cell, [str(risk_rows[i].get("risk", "") or "")])

        i = j






def _p_text(p_elm) -> str:
    """Текст абзаца из XML (w:t)."""
    out = []
    for t in p_elm.iter():
        if t.tag == qn("w:t") and t.text:
            out.append(t.text)
    return "".join(out)


def _norm_ws(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip()


def _create_page_break_p() -> OxmlElement:
    """Абзац с разрывом страницы."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _find_caption_and_first_table5(doc: DocxDocument) -> Tuple[Optional[int], Optional[int], Optional[OxmlElement]]:
    """
    Ищем в body:
      - абзац 'Таблица 5' (его индекс в body)
      - первую таблицу сразу после него (индекс в body + сам XML tbl)
    """
    body = doc.element.body
    children = list(body)

    cap_body_idx = None
    for i, ch in enumerate(children):
        if ch.tag == qn("w:p"):
            txt = _norm_ws(_p_text(ch))
            if re.match(r"(?i)^таблица\s*5\b", txt):
                cap_body_idx = i
                break

    if cap_body_idx is None:
        return None, None, None

    tbl_body_idx = None
    tbl_xml = None
    for j in range(cap_body_idx + 1, len(children)):
        ch = children[j]
        if ch.tag == qn("w:tbl"):
            tbl_body_idx = j
            tbl_xml = ch
            break

    return cap_body_idx, tbl_body_idx, tbl_xml


def split_table5_like_example(
    doc: DocxDocument,
    *,
    header_rows: int = 2,
    first_page_data_rows: int = 4,
    next_page_data_rows: int = 6,
) -> None:
    """
    Делит Таблицу 5 "как в 6_OQ — копия":
      - В первой части оставляем шапку (2 строки) + N строк данных
      - Затем: разрыв страницы + 'Продолжение таблицы 5' + таблица ТОЛЬКО со строками данных (без шапки)
      - Если данных много — повторяем блоки (каждый блок по next_page_data_rows)

    Важно: python-docx не умеет реально мерить «влезло/не влезло в страницу»,
    поэтому режем детерминированно по числу строк (как в примере).
    """

    body = doc.element.body
    children = list(body)

    cap_body_idx, tbl_body_idx, tbl_xml = _find_caption_and_first_table5(doc)
    if tbl_body_idx is None or tbl_xml is None:
        return

    # стиль подписи "Таблица 5" (чтобы "Продолжение..." было таким же)
    cap_style = None
    try:
        # сопоставим body абзац с doc.paragraphs (по _p)
        p_map = {p._p: p for p in doc.paragraphs}
        cap_p = p_map.get(children[cap_body_idx], None) if cap_body_idx is not None else None
        cap_style = cap_p.style if cap_p is not None else None
    except Exception:
        cap_style = None

    # защита от повторного запуска: если уже есть "Продолжение таблицы 5" после таблицы — выходим
    for k in range(tbl_body_idx + 1, min(tbl_body_idx + 8, len(children))):
        if children[k].tag == qn("w:p") and re.match(r"(?i)^продолжение\s+таблицы\s*5\b", _norm_ws(_p_text(children[k]))):
            return

    # Сохраним "базовую" таблицу (с tblPr/tblGrid и т.п.)
    base_tbl_xml = deepcopy(tbl_xml)

    # tr_list текущей таблицы
    tr_list: List[OxmlElement] = list(tbl_xml.iterchildren(qn("w:tr")))
    if len(tr_list) <= header_rows:
        return

    header_trs = tr_list[:header_rows]
    data_trs = tr_list[header_rows:]

    if len(data_trs) <= first_page_data_rows:
        return  # всё влезло в первую часть

    # чанки данных
    chunks: List[List[OxmlElement]] = []
    chunks.append(data_trs[:first_page_data_rows])

    rest = data_trs[first_page_data_rows:]
    while rest:
        chunks.append(rest[:next_page_data_rows])
        rest = rest[next_page_data_rows:]

    # 1) Обрезаем первую таблицу: header + первый chunk
    keep_first = set(header_trs + chunks[0])
    for tr in tr_list[::-1]:
        if tr not in keep_first:
            tbl_xml.remove(tr)

    # 2) Вставляем продолжения после первой таблицы
    insert_after_idx = tbl_body_idx

    for part in chunks[1:]:
        # (a) разрыв страницы
        insert_after_idx += 1
        body.insert(insert_after_idx, _create_page_break_p())

        # (b) абзац "Продолжение таблицы 5" (вставляем через add_paragraph и переносим XML)
        p = doc.add_paragraph("Продолжение таблицы 5")
        if cap_style is not None:
            try:
                p.style = cap_style
            except Exception:
                pass
        try:
            p.paragraph_format.keep_with_next = True
        except Exception:
            pass

        p_elm = p._p
        body.remove(p_elm)
        insert_after_idx += 1
        body.insert(insert_after_idx, p_elm)

        # (c) новая таблица: КАК В ПРИМЕРЕ — БЕЗ ШАПКИ, только data rows
        new_tbl = deepcopy(base_tbl_xml)

        # удалить все строки из клона
        for tr in list(new_tbl.iterchildren(qn("w:tr"))):
            new_tbl.remove(tr)

        # добавить строки данных
        for tr in part:
            new_tbl.append(deepcopy(tr))

        insert_after_idx += 1
        body.insert(insert_after_idx, new_tbl)
